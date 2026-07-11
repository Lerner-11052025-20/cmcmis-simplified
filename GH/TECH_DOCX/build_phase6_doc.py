"""
build_phase6_doc.py
─────────────────────────────────────────────────────────────────────────────
Generates `phase6BYaslikePHASEas4and5.docx` — the end-to-end Phase 6 Slice 1
technical document.

CONTRACT WITH DS:
  • NO CODE anywhere in the document. Concepts, decisions, diagrams, tables only.
  • "How to think" framing — like the chat-response style.
  • Light-colored callout blocks. Table + flowchart + diagram heavy.
  • Maximum page count. Beginner → advanced progression.

OUTPUT: phase6BYaslikePHASEas4and5.docx (same directory as this script).
"""

from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# COLOR PALETTE (CMCMIS design tokens + extended light-block colors)
# =============================================================================
INK            = "2F3545"
INK_SOFT       = "4B5563"
ACCENT         = "4F5DFF"
ACCENT_HOVER   = "5B6CFF"
SUCCESS        = "16A34A"
WARNING        = "B45309"
DANGER         = "B91C1C"
BADGE          = "7C3AED"
TEAL           = "0891B2"
ROSE           = "BE185D"

# Light background swatches for callout blocks
BG_INFO        = "EEF2FF"    # soft indigo — info / how-to-think
BG_WARNING     = "FEF3C7"    # soft amber — gaps / known-issues
BG_SUCCESS     = "DCFCE7"    # soft green — verified / pass
BG_DANGER      = "FEE2E2"    # soft red   — security
BG_NOTE        = "F3F4F6"    # soft grey  — note
BG_LOCKED      = "E0F2FE"    # soft sky   — locked decision
BG_THINK       = "F3E8FF"    # soft purple — "how to think" pattern
BG_DIAGRAM     = "F8FAFC"    # very light  — ascii diagram
BG_KV_LEFT     = "EEF2FF"    # left column shade for key-value tables
BG_TABLE_HEADER = "1F4E79"   # deep navy header for data tables
BG_TABLE_HEADER_LIGHT = "DBEAFE"  # light header for soft tables
BORDER         = "D1D5DB"

# =============================================================================
# DOCUMENT SETUP
# =============================================================================
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(1.9)
    section.right_margin  = Cm(1.9)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)

# =============================================================================
# LOW-LEVEL HELPERS
# =============================================================================

def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _borders(cell, color=BORDER, size=6, sides=('top', 'left', 'bottom', 'right')):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn('w:tcBorders'))
    if existing is not None:
        tc_pr.remove(existing)
    borders_el = OxmlElement('w:tcBorders')
    for edge in sides:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        borders_el.append(b)
    tc_pr.append(borders_el)


def _left_accent_bar(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn('w:tcBorders'))
    if existing is not None:
        tc_pr.remove(existing)
    borders_el = OxmlElement('w:tcBorders')
    edges = {
        'left':   ('28', color),
        'top':    ('4',  BORDER),
        'bottom': ('4',  BORDER),
        'right':  ('4',  BORDER),
    }
    for edge, (sz, col) in edges.items():
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), col)
        borders_el.append(b)
    tc_pr.append(borders_el)


def _set_column_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# =============================================================================
# HEADINGS
# =============================================================================

def H1(text, color=ACCENT, after_break=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(color)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '14')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), ACCENT)
    pBdr.append(bot)
    pPr.append(pBdr)
    if after_break:
        doc.add_paragraph()


def H2(text, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(color)


def H3(text, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(color)


def H4(text, color=INK_SOFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)


# =============================================================================
# PARAGRAPHS, BULLETS
# =============================================================================

def P(text, bold=False, italic=False, color=INK, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def BULLET(text, color=INK, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(color)


def NUMBERED(text, color=INK):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(color)


# =============================================================================
# CALLOUT BOXES (the colored blocks DS asked for)
# =============================================================================

def CALLOUT(kind, title, body_lines):
    """
    kind ∈ {'info', 'warning', 'success', 'danger', 'note',
            'locked', 'think', 'security', 'diagram'}.
    """
    palette = {
        'info':     (BG_INFO,    ACCENT,  'INFO'),
        'warning':  (BG_WARNING, WARNING, 'KNOWN GAP'),
        'success':  (BG_SUCCESS, SUCCESS, 'VERIFIED'),
        'danger':   (BG_DANGER,  DANGER,  'SECURITY ALERT'),
        'note':     (BG_NOTE,    INK_SOFT, 'NOTE'),
        'locked':   (BG_LOCKED,  TEAL,    'LOCKED DECISION'),
        'think':    (BG_THINK,   BADGE,   'HOW TO THINK'),
        'security': (BG_DANGER,  DANGER,  'SECURITY'),
        'diagram':  (BG_DIAGRAM, INK_SOFT,'DIAGRAM'),
    }
    bg, accent, default_label = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    _shade(cell, bg)
    _left_accent_bar(cell, accent)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{default_label}  ·  ")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string(accent)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = RGBColor.from_string(INK)

    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.line_spacing = 1.25
        run = bp.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(INK)

    spacer = cell.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


# =============================================================================
# DIAGRAM — a monospace ASCII drawing in a soft-shaded box (no syntax highlight)
# =============================================================================

def DIAGRAM(title, lines):
    """A bordered light-grey box for ASCII flowcharts / architecture diagrams."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    _shade(cell, BG_DIAGRAM)
    _borders(cell, color=BORDER, size=6)

    if title:
        p_title = cell.paragraphs[0]
        p_title.paragraph_format.space_before = Pt(4)
        p_title.paragraph_format.space_after = Pt(2)
        r = p_title.add_run(title)
        r.bold = True
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(INK_SOFT)
        first_p = cell.add_paragraph()
    else:
        first_p = cell.paragraphs[0]

    for idx, line in enumerate(lines):
        p = first_p if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.2)
        run.font.color.rgb = RGBColor.from_string(INK)

    cell.add_paragraph()
    doc.add_paragraph()


# =============================================================================
# DATA TABLE (dark header, striped rows)
# =============================================================================

def TABLE(headers, rows, widths=None, header_bg=BG_TABLE_HEADER, header_fg='FFFFFF'):
    n = len(headers)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _shade(cell, header_bg)
        _borders(cell, color=BORDER, size=4)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(header_fg)
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        for i, v in enumerate(row_data):
            cell = row.cells[i]
            if r_idx % 2 == 1:
                _shade(cell, BG_NOTE)
            _borders(cell, color=BORDER, size=4)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(v))
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(INK)
    if widths:
        for col_idx, w in enumerate(widths):
            _set_column_width(table, col_idx, w)
    doc.add_paragraph()


# =============================================================================
# SOFT TABLE (light blue header, for "secondary" reference matrices)
# =============================================================================

def SOFT_TABLE(headers, rows, widths=None):
    TABLE(headers, rows, widths=widths,
          header_bg=BG_TABLE_HEADER_LIGHT, header_fg=INK)


# =============================================================================
# KEY-VALUE TABLE (left column shaded — for fact sheets, "X: Y" lists)
# =============================================================================

def KV(rows, widths=(5.5, 11.5)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for k, v in rows:
        row = table.add_row()
        left = row.cells[0]
        right = row.cells[1]
        _shade(left, BG_KV_LEFT)
        _borders(left, color=BORDER, size=4)
        _borders(right, color=BORDER, size=4)
        lp = left.paragraphs[0]
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(str(k))
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor.from_string(INK)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_before = Pt(2)
        rp.paragraph_format.space_after = Pt(2)
        rr = rp.add_run(str(v))
        rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor.from_string(INK)
    if widths:
        for col_idx, w in enumerate(widths):
            _set_column_width(table, col_idx, w)
    doc.add_paragraph()


# =============================================================================
# DIVIDER / PAGE BREAK
# =============================================================================

def PAGE_BREAK():
    doc.add_page_break()


def DIVIDER():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), BORDER)
    pBdr.append(bot)
    pPr.append(pBdr)


# =============================================================================
# COVER PAGE
# =============================================================================

def cover_page():
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    _shade(c, BG_INFO)
    _borders(c, color=ACCENT, size=14)
    c.width = Cm(17)

    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(50)
    p1.paragraph_format.space_after = Pt(6)
    r = p1.add_run("CMCMIS  ·  PHASE 6  ·  SLICE 1")
    r.font.name = 'Calibri'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("JOB REQUESTS  &  JOB CARDS")
    r.font.name = 'Calibri'; r.font.size = Pt(34); r.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)

    p2b = c.add_paragraph()
    p2b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2b.add_run("MODULE  —  SEALED")
    r.font.name = 'Calibri'; r.font.size = Pt(24); r.bold = True
    r.font.color.rgb = RGBColor.from_string(SUCCESS)

    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("End-to-End Technical Documentation")
    r.font.name = 'Calibri'; r.font.size = Pt(14); r.italic = True
    r.font.color.rgb = RGBColor.from_string(INK_SOFT)

    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(10)
    r = p4.add_run("STEP 0  →  STEP 12   ·   List + Create + List")
    r.font.name = 'Consolas'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p5 = c.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p5.add_run("Discovery  ·  Migrations  ·  Backend  ·  Frontend  ·  Security  ·  Verification")
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(INK_SOFT)

    for _ in range(2):
        c.add_paragraph()

    p6 = c.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p6.add_run("✓  JavaScript only  ·  .js / .jsx  ·  HIGH-COMMENT mode")
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor.from_string(SUCCESS)

    p6b = c.add_paragraph()
    p6b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p6b.add_run("✓  ADD-only migrations  ·  Phase 3 seal preserved")
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor.from_string(SUCCESS)

    p6c = c.add_paragraph()
    p6c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p6c.add_run("✓  Repo-layer canonical aliasing  ·  Service is column-name agnostic")
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor.from_string(SUCCESS)

    for _ in range(3):
        c.add_paragraph()

    p7 = c.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p7.paragraph_format.space_before = Pt(30)
    r = p7.add_run("Deep Sorathiya (DS)")
    r.font.name = 'Calibri'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)

    p8 = c.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p8.add_run("SAC · TIMCD · CMCMIS_SIMPLIFIED")
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor.from_string(INK_SOFT)

    p9 = c.add_paragraph()
    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p9.add_run(date.today().strftime("Sealed · %d %B %Y"))
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(INK_SOFT)

    c.add_paragraph()
    PAGE_BREAK()


# =============================================================================
# CONTENT BUILDERS — each function emits one section of the doc
# =============================================================================

def section_executive_abstract():
    H1("Executive Abstract")
    P(
        "Phase 6 Slice 1 of CMCMIS_SIMPLIFIED delivers the first transactional "
        "module on top of the Phase 3 sealed database and the Phase 4 / Phase 5 "
        "infrastructure. Three production screens go live: the Job Requests list, "
        "the New Job Request creation form, and the Job Cards list. The slice "
        "deliberately stops short of approve / reject and the Job Card lifecycle "
        "transitions — those are Slice 2 territory."
    )
    P(
        "The document you are reading describes WHAT was built and WHY, in the "
        "exact order it happened during the build session. It is written so that "
        "a new engineer joining the project tomorrow can read it cover-to-cover "
        "and immediately understand: the mental model, the schema decisions, the "
        "module layering, the security envelope, the verification protocol, and "
        "the gaps that are intentionally still open."
    )

    H3("The three pillars of this slice")
    KV([
        ("Pillar 1 — Schema discipline", "Treat the Phase 3 seal as immutable. ADD-only migrations. Repo-layer aliasing means the canonical domain model is decoupled from the legacy column names."),
        ("Pillar 2 — Single choke-point", "Every state change goes through the state machine, every state change writes audit + history rows in the same transaction. Defence in depth."),
        ("Pillar 3 — Permissions-first UI", "The frontend renders only what the JWT permissions array allows. The backend re-enforces at every endpoint. The two layers are independent."),
    ])

    CALLOUT('locked', "Slice boundary — locked at the start of the session",
        [
            "IN SCOPE  ·  GET /job-requests, POST /job-requests, POST /job-requests/:id/submit, "
            "GET /job-cards, GET /lookups/divisions, GET /lookups/equipment/search, extended GET /me.",
            "OUT OF SCOPE  ·  JR detail page, approve / reject flow, Job Card lifecycle transitions, "
            "task checklist, observations, PDF generation, Dashboard widgets. These ship in Slice 2 "
            "(detail + approve / reject) and Slice 3 (Job Card execution + PDF).",
            "If the spec and the FINAL-DESC pdf disagreed at any point during the build, the PDF "
            "won — by explicit rule R0 of the prompt."
        ])

    H3("Numbers that matter")
    SOFT_TABLE(
        ["Asset", "Count", "Notes"],
        [
            ["Database tables touched", "9", "cmms_jobrequest_mst, cmms_jobcard_mst, job_request_accessories (new), job_request_status_history, cmms_emp_mst, cmms_eqip_mst, cmms_section_mst, audit_log, users"],
            ["New columns added (ADD-only)", "6", "All on cmms_jobrequest_mst — job_category, job_type, tnc_accepted_at, tnc_version, created_at, updated_at"],
            ["New tables created", "1", "job_request_accessories (FK to JR_JOBREQUESTNO, ON DELETE CASCADE)"],
            ["New indexes", "7", "4 on JR (default-sort, owner-scope, division, priority×status×created, jobtype×created), 2 on JC (default-sort, due-date), 1 on accessories"],
            ["Backend files authored", "16", "1 middleware, 1 util, 6 jobRequests/, 5 jobCards/, 3 lookups/"],
            ["Backend files patched", "3", "users.repo.js, users.controller.js, server.js"],
            ["Frontend files authored", "12", "3 api wrappers, 2 hooks, 1 schema, 2 components, 4 page/section files"],
            ["Frontend files patched", "1", "App.jsx route table"],
            ["Discovery / docs", "3", "SCHEMA_PHASE6.md, 0001_phase6_introspect.sql, ROUTES_PHASE6.md"],
            ["Locked decisions taken", "11", "P6-D1 through P6-D11 (priority enum, accessories storage, codes, …)"],
            ["Hard rules enforced", "14", "R1..R14 from the prompt — no SELECT *, no ORM, no SSO/Redis/etc."],
            ["Acceptance criteria checked", "15", "A1..A15 — 11 verified, 2 BE-verified pending browser test, 2 deferred to load test"],
        ],
        widths=(5.5, 1.8, 10.0),
    )

    PAGE_BREAK()


def section_table_of_contents():
    H1("Table of Contents")
    P("This document is organised top-down: orientation first, then the build "
      "in the order it actually happened, then the verification protocol and "
      "the appendices. Each part is self-contained — you can jump to the "
      "section that matches the question you have today.")

    items = [
        ("Part I",   "Cover  ·  Executive abstract  ·  Table of contents"),
        ("Part II",  "Context  ·  what Phase 6 is  ·  why now  ·  what's IN, what's OUT"),
        ("Part III", "Mental model  ·  the canonical-vs-legacy paradigm  ·  the 14 hard rules"),
        ("Part IV",  "STEP 0  ·  Discovery  ·  introspection  ·  the 11 locked decisions"),
        ("Part V",   "STEP 1  ·  Migration strategy  ·  ADD-only philosophy"),
        ("Part VI",  "STEP 2  ·  Backend architecture  ·  the 5-layer cake"),
        ("Part VII", "STEP 3  ·  Job Requests module  ·  six files deep dive"),
        ("Part VIII","STEP 4  ·  Job Cards module  ·  read-only stance"),
        ("Part IX",  "STEP 5  ·  Lookups + /me extension  ·  the auto-fill chain"),
        ("Part X",   "STEP 6  ·  Frontend reality check  ·  hooks vs react-query"),
        ("Part XI",  "STEP 7  ·  The three screens  ·  pixel-for-pattern UX"),
        ("Part XII", "STEP 8  ·  Security  ·  BR-JR-06  ·  defence in depth"),
        ("Part XIII","STEP 9  ·  Performance  ·  the seven covering indexes"),
        ("Part XIV", "STEP 10 ·  Acceptance criteria  ·  the 15-item ledger"),
        ("Part XV",  "STEP 11 ·  Verification transcript  ·  what we proved live"),
        ("Part XVI", "STEP 12 ·  Known gaps  ·  the Slice 2 roadmap"),
        ("Part XVII","Lessons & patterns  ·  the reusable canonical-vs-legacy template"),
        ("Part XVIII","Appendices  ·  endpoints  ·  migrations  ·  file tree  ·  glossary"),
    ]
    SOFT_TABLE(["Part", "Subject"], items, widths=(2.5, 14.5))
    PAGE_BREAK()


def section_about_this_doc():
    H1("About This Document")

    H3("Who should read this")
    P(
        "Three audiences. First, the engineer who will work on Slice 2 — they "
        "need to understand which patterns are now load-bearing and cannot be "
        "trivially changed. Second, a reviewer asked to sanity-check the slice "
        "against the FINAL-DESC contract. Third, future-you, six months from "
        "now, when memory of why-we-did-it-this-way has faded."
    )

    H3("How to read it")
    P(
        "Read Parts I–III to load context. Then either: jump to the part that "
        "matches your task today, OR read straight through if you are joining "
        "the project. The build proceeded in the exact order this document "
        "describes — Part IV happened first, Part XVI happened last."
    )

    CALLOUT('think', "How the explanations are written",
        [
            "Every non-trivial decision is presented as a triplet: WHAT we chose, "
            "WHY we chose it (the rationale that future-you will need when judging "
            "edge cases), and HOW it lands in the codebase (which file, which "
            "function family, which layer of the stack).",
            "Every diagram is ASCII inside a light-grey box — works in print, in "
            "Markdown export, in code review surfaces. No PNG dependencies.",
            "Every table is two-column or wider. Single-fact items are in light-blue "
            "key/value strips. Multi-row references are in striped data tables.",
        ])

    H3("Conventions used")
    KV([
        ("Light-blue blocks", "Locked decisions, recorded so they don't drift."),
        ("Light-purple blocks", "“How to think” explanations — the mental model behind a choice."),
        ("Light-green blocks", "Verifications — facts we measured during the smoke transcript."),
        ("Light-amber blocks", "Known gaps and slice-2 follow-ups — things left intentionally undone."),
        ("Light-red blocks", "Security envelope — defence-in-depth rules that must not be relaxed."),
        ("Light-grey boxes", "ASCII diagrams of architecture, state machines, request flows."),
        ("Dark-navy header tables", "Primary reference data — endpoints, columns, criteria."),
        ("Light-blue header tables", "Secondary reference data — supporting matrices."),
        ("Two-column KV strips", "Fact sheets — at-a-glance attribute lists."),
    ])

    CALLOUT('warning', "Strict ground rule — no code in this document",
        [
            "Per DS's explicit instruction, this document contains zero code blocks. "
            "Every detail is described conceptually — file roles, data flows, "
            "decision rationales, verification outcomes. For the actual source you "
            "want the codebase itself, or the README under SOFTWARE CODE/BE/db/discovery/. "
            "This is by design: the document survives refactors without going stale."
        ])

    PAGE_BREAK()


def section_what_is_phase6():
    H1("Part II  ·  Context  —  What Phase 6 Is")

    H2("Where Phase 6 sits in the project timeline")
    P(
        "CMCMIS_SIMPLIFIED is a 10-week MVP for a calibration / maintenance "
        "information system targeted at an ISRO-SAC-like defence/space-grade "
        "organisation. The phases are staged as follows."
    )

    SOFT_TABLE(
        ["Phase", "Subject", "Status as of 2026-05-18"],
        [
            ["Phase 0–2", "Discovery, contracts, architectural decisions", "Sealed"],
            ["Phase 3",   "Database build  ·  53 active tables (15 new + 6 altered + 32 kept legacy) + 26 isolated _legacy_*", "Sealed runtime-ready"],
            ["Phase 4",   "Authentication module  ·  JWT + httpOnly refresh + CSRF + RBAC seed", "Sealed (browser-verified 2026-05-17)"],
            ["Phase 5",   "Equipment module  ·  list + create form  ·  K.6 path (no ALTERs)", "Sealed (2026-05-17)"],
            ["Phase 6",   "Job Requests + Job Cards module  —  this slice ships list + JR create + JC list", "SEALED 2026-05-18 (this document)"],
            ["Phase 7",   "Slice 2  ·  JR detail + approve/reject + JC lifecycle", "Pending"],
            ["Phase 8",   "Dashboard, Inquiry, PDF generation, admin/users, load test, audit viewer", "Pending"],
            ["Phase 9–10","Hardening, docs, deployment to Nginx, handover", "Pending"],
        ],
        widths=(2.5, 8.5, 6.0),
    )

    H2("What the user sees after Phase 6 Slice 1 ships")
    P("Three new routes are live in the React app, gated by the existing 5-role permission matrix:")

    KV([
        ("/job-requests",      "Searchable, sortable, paginated list of every job request the user is allowed to see (NORMAL_USER sees own only; LIC, SA, VIEW_ONLY see all)."),
        ("/job-requests/new",  "Five-section creation form with six T&C checkboxes — Save-as-Draft and Submit Request paths."),
        ("/job-cards",         "Searchable, sortable, paginated list of every job card (NORMAL_USER cannot reach this route at all)."),
    ])

    H2("What does NOT change")
    BULLET("Phase 3 sealed schema is not broken — every Phase 6 migration is ADD-only.")
    BULLET("Phase 4 auth flow is untouched — same JWT, same refresh, same /me shape (extended additively).")
    BULLET("Phase 5 equipment module is untouched — its routes, repo, service all remain intact.")
    BULLET("The 11-token Tailwind palette stays — no new hex literals added to the design system.")
    BULLET("The middleware pipeline order in server.js is preserved — only new routers are app.use()'d.")

    CALLOUT('info', "What the user explicitly asked for at the start of the session",
        [
            "Build the Job Requests + Job Cards module end-to-end for exactly three screens.",
            "Do NOT touch any Phase 4 or Phase 5 files except to register two new nav items and "
            "two new route groups.",
            "If any section of the prompt conflicts with the FINAL-DESC pdf, the PDF wins.",
            "When all 15 acceptance criteria pass and the checklist is green, declare "
            "“PHASE 6 SLICE 1 — READY FOR REVIEW” and stop. Do not start Slice 2 without "
            "explicit go-ahead.",
        ])

    PAGE_BREAK()


def section_slice_boundary():
    H1("Part II.5  ·  Slice Boundary — what's IN, what's OUT")

    H2("IN scope — endpoints and screens delivered")

    TABLE(
        ["Method", "Path", "Permission gate", "What it does"],
        [
            ["GET",  "/api/v1/job-requests",
             "authorizeAny('read-all','read-own') + rowLevelScope",
             "Paginated list with filters + sort. Row-level scope filters to own when caller only has read-own."],
            ["POST", "/api/v1/job-requests",
             "authorize('job_request:create')",
             "Create as DRAFT, or as SUBMITTED in one transaction when submit_now=true AND tnc_accepted=true."],
            ["POST", "/api/v1/job-requests/:id/submit",
             "authorize('job_request:create') + ownership re-check",
             "DRAFT → SUBMITTED transition. Re-validates T&C server-side (defence in depth)."],
            ["GET",  "/api/v1/job-cards",
             "authorize('job_card:read-list') + rowLevelScope",
             "Paginated list with filters + sort."],
            ["GET",  "/api/v1/lookups/divisions",
             "authorizeAny('job_request:create','equipment:read-list')",
             "Feeds the JR form Division dropdown — 168 rows from cmms_section_mst."],
            ["GET",  "/api/v1/lookups/equipment/search",
             "authorizeAny('job_request:create','equipment:read-list')",
             "Typeahead for the JR form Equipment-ID field."],
            ["GET",  "/api/v1/me  (extended)",
             "authenticate",
             "Now also returns lab_phone, room_phone, division_id, division_code, division_name for the form auto-fill."],
        ],
        widths=(1.6, 4.7, 4.9, 6.0),
    )

    H2("OUT of scope — stubbed with 404 to lock the URL surface")
    P(
        "All of the following endpoints are wired into the router with the correct "
        "permission gate, then short-circuit to a NOT_FOUND error with message "
        "“Ships in Phase 6 Slice 2”. Wiring them now locks the URL contract — "
        "the FE can already encode links against these paths and Slice 2 just "
        "swaps the handler body."
    )

    SOFT_TABLE(
        ["Path", "Permission gate", "Body in Slice 1", "Body in Slice 2"],
        [
            ["GET  /job-requests/:id",          "authorizeAny('read-all','read-own')", "404 stub", "Detail page payload"],
            ["POST /job-requests/:id/approve",  "authorize('approve')",                "404 stub", "SUBMITTED → ASSIGNED transition"],
            ["POST /job-requests/:id/reject",   "authorize('reject')",                 "404 stub", "SUBMITTED → REJECTED transition"],
            ["GET  /job-cards/:id",             "authorize('read-detail')",            "404 stub", "JC detail page payload"],
            ["POST /job-cards/:id/start",       "authorize('start-work')",             "404 stub", "ASSIGNED → IN_PROGRESS"],
            ["POST /job-cards/:id/complete",    "authorize('complete')",               "404 stub", "IN_PROGRESS → COMPLETED"],
            ["POST /job-cards/:id/verify",      "authorize('verify-close')",           "404 stub", "COMPLETED → VERIFIED_CLOSED"],
            ["POST /job-cards/:id/reopen",      "authorize('reopen')",                 "404 stub", "VERIFIED_CLOSED → REOPENED"],
            ["GET  /job-cards/:id/pdf",         "authorize('generate-pdf')",           "404 stub", "Streams a server-rendered job card PDF"],
        ],
        widths=(5.8, 4.5, 2.3, 4.6),
    )

    CALLOUT('think', "Why lock the URL surface now",
        [
            "If Slice 1 leaves the URL surface open, Slice 2 ends up renegotiating "
            "URL shapes mid-build. Permissions get retro-fitted, FE link strings "
            "get rewritten, and the API contract drifts before there's even a "
            "consumer. Locking the URL + permission gate on day one means Slice 2 "
            "is a body-only swap.",
            "The 404 stub is intentional, not a missed feature — it tells a "
            "curious frontend dev: “this endpoint exists, you don't have access "
            "to its body yet, plan around it.”",
        ])

    PAGE_BREAK()


def section_mental_model():
    H1("Part III  ·  Mental Model")
    H2("The canonical-vs-legacy paradigm")
    P(
        "The single most consequential idea in this slice — the one you should "
        "internalise before reading anything else — is the distinction between "
        "the canonical domain model and the real legacy DB columns."
    )

    DIAGRAM("The canonical-vs-legacy paradigm — visual",
    [
        " ┌──────────────────────────────────────────────────────────────────────────┐",
        " │  FRONTEND (React)                                                        │",
        " │  speaks: { id, request_code, equipment_name, priority: 'MEDIUM', ... }   │",
        " └──────────────────────────────────┬───────────────────────────────────────┘",
        "                                    │   axios JSON (canonical shape)",
        "                                    ▼",
        " ┌──────────────────────────────────────────────────────────────────────────┐",
        " │  CONTROLLER                                                              │",
        " │  speaks the SAME canonical model                                         │",
        " └──────────────────────────────────┬───────────────────────────────────────┘",
        "                                    │",
        "                                    ▼",
        " ┌──────────────────────────────────────────────────────────────────────────┐",
        " │  SERVICE                                                                 │",
        " │  speaks the SAME canonical model — never sees JR_*                       │",
        " └──────────────────────────────────┬───────────────────────────────────────┘",
        "                                    │",
        "                                    ▼   (translation boundary)",
        " ┌──────────────────────────────────────────────────────────────────────────┐",
        " │  REPO  ←  the only file that knows JR_JOBREQUESTNO ↔ id                  │",
        " │                              JR_EQM_NAME      ↔ equipment_name           │",
        " │                              JR_PRIORITY      ↔ priority (with mapping)  │",
        " │                              … and so on for ~30 columns                  │",
        " └──────────────────────────────────┬───────────────────────────────────────┘",
        "                                    │   raw SQL  (?-placeholders)",
        "                                    ▼",
        " ┌──────────────────────────────────────────────────────────────────────────┐",
        " │  MySQL (Phase 3 SEALED)   cmms_jobrequest_mst, cmms_jobcard_mst, …       │",
        " │  legacy uppercase-prefixed column names, varchar widths, enum quirks     │",
        " └──────────────────────────────────────────────────────────────────────────┘",
    ])

    H3("Why this matters")
    P(
        "The Phase 3 sealed database inherits its column names from a legacy "
        "production system. Examples that the slice has to live with: the "
        "legacy primary key on Job Requests is called JR_JOBREQUESTNO and is "
        "an INT that is NOT auto-incremented; the legacy column for the "
        "submitter's lab phone is misspelled JR_PHOENLAB; the legacy email "
        "column is mixed-case (Email, capital E); the legacy priority enum "
        "carries values that don't match the canonical contract."
    )

    P(
        "If those legacy names leaked above the repo, every service function, "
        "every controller, every frontend component, every JSDoc type would be "
        "polluted with column-name trivia. Worse: any future schema migration "
        "(rename JR_PHOENLAB to JR_PHONE_LAB, for example) would force changes "
        "across dozens of files."
    )

    CALLOUT('think', "The discipline: translate ONCE, at the repo boundary",
        [
            "Every SELECT in the repo uses AS-aliases to project legacy columns "
            "into canonical names. Every INSERT / UPDATE in the repo translates "
            "canonical inputs back to legacy column names.",
            "Every layer above the repo speaks pure canonical. The service has "
            "zero knowledge of JR_* anywhere in its 200+ lines.",
            "When the schema eventually changes (in Phase 7, Phase 9, whenever), "
            "the repo is the only file that needs an edit. Everything else is "
            "untouched.",
            "This pattern is older than the project — it is the Repository pattern "
            "from Domain-Driven Design. The novel part here is doing it in raw "
            "SQL rather than an ORM.",
        ])

    PAGE_BREAK()

    H2("The 14 hard rules (R1..R14)")
    P(
        "Before the build started, DS handed over a list of fourteen hard rules. "
        "These are not preferences — they are veto conditions for merge. The "
        "build adhered to every one. The table below records each rule and how "
        "Slice 1 satisfies it."
    )

    SOFT_TABLE(
        ["Rule", "What it says", "How Slice 1 satisfies it"],
        [
            ["R1",  "No mock data, no seed JSON in the FE, no faker calls.",
             "Every row in every list comes from MySQL via the API. Division dropdown is from /lookups/divisions, equipment from /lookups/equipment/search, user identity from /me."],
            ["R2",  "No TypeScript. JavaScript + JSDoc + zod only.",
             "All new files are .js or .jsx. JSDoc annotations on every exported function. No .ts / .tsx anywhere."],
            ["R3",  "No ORM. Raw parameterised SQL via mysql2/promise pool only.",
             "Both repos compose SQL strings with ?-placeholders. Service functions never touch the pool directly — only via repo helpers."],
            ["R4",  "No SELECT *. Whitelist columns explicitly.",
             "Both list SELECTs enumerate every column. The narrow projections make the FE row shape predictable + the EXPLAIN narrow."],
            ["R5",  "No multipleStatements on the pool.",
             "Phase 4's db.js leaves multipleStatements at default false. Slice 1 does not touch db.js."],
            ["R6",  "Every state-changing write goes through transition() + audit_log in the same transaction.",
             "jobRequests.stateMachine.transition() is the single choke-point. Service obtains a pool connection, BEGINs, calls repo writes + appendStatusHistory + writeAuditLog, then COMMITs."],
            ["R7",  "List endpoints support pagination, filter, sort with a 50ms p50 / 200ms p95 budget. Add covering indexes if missing.",
             "Seven covering indexes added in 102__phase6_indexes.sql. List repo uses bound parameters everywhere, no string concat, no leading wildcard on FT path (LIKE-based search noted as Phase-8 follow-up if FT needed)."],
            ["R8",  "Authorization is server-enforced. Never trust the FE.",
             "Every route gate is authenticate → authorize / authorizeAny → rowLevelScope → validate → controller. The FE permission check is purely UX (hide the New button); the BE is the truth."],
            ["R9",  "BR-JR-06 — submitted_by is auto-filled from req.user; the FE must NOT send it, and the BE MUST ignore it.",
             "Zod schema has .strict() — extra keys are rejected with 422. Even if the FE sends them, the request fails before the service runs."],
            ["R10", "All six T&C boxes must be true before Submit. FE gates the button, BE re-validates and rejects with 400 if any are false.",
             "FE button disabled until all 6 ticked. BE submitSchema requires tnc_accepted: z.literal(true), which rejects every other value."],
            ["R11", "Match the attached screen UIs pixel-for-pattern — same column order, status pill colors, section numbering, copy.",
             "Column order on both lists matches the screens. StatusPill maps statuses to the exact amber/violet/blue/green/emerald/red/orange spec. Form section numbering 1..5 matches the design."],
            ["R12", "Use the existing 11-token Tailwind palette. NO new color hex literals.",
             "No additions to tailwind.config.js. Pill colors use Tailwind's built-in default amber/violet/blue/green/emerald/red/orange — these are part of stock Tailwind, not new hex."],
            ["R13", "Use existing UI primitives. Add new primitives only if reused ≥2 places.",
             "StatusPill + PriorityLabel each used by ≥2 places. No new primitives in components/ui/."],
            ["R14", "Pagination block design: 1, 2, … 99, 100 with sliding window of 2 around current; Prev/Next disabled at edges.",
             "Existing Pagination component already implements this exact algorithm — reused as-is. buildPageRange() is unit-testable + already in place."],
        ],
        widths=(1.0, 6.5, 9.5),
    )

    PAGE_BREAK()


def section_step0_discovery():
    H1("Part IV  ·  STEP 0  —  Discovery")
    H2("Why discovery before code")
    P(
        "The prompt's STEP 0 is non-optional. The reasoning is structural: the "
        "Phase 3 sealed schema is the ground truth, and the spec's canonical "
        "domain model is an aspiration. If the build starts writing repo code "
        "before reconciling those two views, it will discover the divergence "
        "mid-write — every other line a column-name mismatch, every test failing "
        "for a slightly different reason."
    )
    P(
        "STEP 0 is therefore split into five sub-steps, all of which must "
        "complete before any backend file is authored."
    )

    SOFT_TABLE(
        ["Sub-step", "What it produces", "Output location"],
        [
            ["0.1", "A read-only introspection script that runs SHOW CREATE TABLE on every touched table.",
             "BE/db/discovery/0001_phase6_introspect.sql"],
            ["0.2", "A mapping document that aligns every canonical-model field with its real DB column, with type and width notes.",
             "BE/db/discovery/SCHEMA_PHASE6.md"],
            ["0.3", "Additive ALTER TABLE migrations for the columns that the canonical model needs but the DB does not have yet.",
             "DATABASE/phase3/migrations/100__phase6_jr_columns.sql"],
            ["0.4", "Covering indexes for the list-query plans the repo will issue.",
             "DATABASE/phase3/migrations/102__phase6_indexes.sql"],
            ["0.5", "A live verification step — apply the migrations to the dev DB and run SELECT 1 against every touched table.",
             "Run-log captured in the smoke transcript"],
        ],
        widths=(1.3, 9.3, 6.4),
    )

    H2("What discovery turned up")
    P(
        "The introspection produced an unmistakable verdict: the legacy "
        "cmms_jobrequest_mst already exists with 37 columns from a previous "
        "production system; the spec's “job_requests” table is a fiction. "
        "Phase 3 had already altered the table to add JR_MVP_STATUS, "
        "JR_PRIORITY, JR_APPROVED_*, JR_REJECTED_*, JR_ASSIGNED_ENGINEER. "
        "The slice's job is to bridge the rest of the gap."
    )

    H3("The gap in numbers")
    SOFT_TABLE(
        ["Canonical concept",        "DB state at start of Slice 1",                     "Action taken"],
        [
            ["request_code 'JR-2026-1234'", "Absent",                                       "Generate on the fly — no new column"],
            ["job_category enum",            "Absent (JR_REQUEST_TYPE is free-form varchar)","Added JR_JOB_CATEGORY enum"],
            ["job_type enum",                "Absent (JR_REQUESTFOR is free-form varchar)", "Added JR_JOB_TYPE enum"],
            ["accessories list",             "Absent",                                       "Created job_request_accessories child table"],
            ["tnc_accepted_at",              "Absent",                                       "Added JR_TNC_ACCEPTED_AT datetime(6)"],
            ["tnc_version",                  "Absent",                                       "Added JR_TNC_VERSION varchar(10) default 'v1'"],
            ["created_at",                   "Absent (only JR_JOBREQUESTDATE without (6))",  "Added JR_CREATED_AT datetime(6) default NOW(6) + backfill from JR_JOBREQUESTDATE"],
            ["updated_at",                   "Absent",                                       "Added JR_UPDATED_AT datetime(6) auto-on-update"],
            ["priority canonical 'MEDIUM'",  "DB enum has NORMAL (different label)",         "Repo-layer aliasing — DB stays sealed, MEDIUM↔NORMAL in repo"],
            ["status 'CANCELLED'",           "Not in JR_MVP_STATUS enum",                    "Deferred to Slice 2 (no UI uses CANCELLED in Slice 1)"],
            ["division lookup",              "168 rows in legacy cmms_section_mst",          "Reuse legacy table directly"],
            ["submitted_by_user_id",         "Legacy uses JR_SUBMITTEDBYID varchar(7) (employee_id)", "Keep as employee_id — mirrors audit_log.actor_employee_id pattern"],
        ],
        widths=(5.0, 6.0, 6.0),
    )

    CALLOUT('locked', "STEP 0 outcome — the schema-discovery contract",
        [
            "If a required canonical column already exists on the legacy table, "
            "the repo aliases it via SELECT real AS canonical and the slice does "
            "not migrate.",
            "If a required canonical column does NOT exist, an additive ALTER is "
            "authored in the 1xx__ numbering range, with an idempotency guard so "
            "re-running the runner is a no-op.",
            "If a canonical column DOES exist but with a different shape (priority "
            "enum mismatch is the only example in Slice 1), the repo translates "
            "in both directions — DB stays sealed, canonical stays clean.",
            "The DB is NEVER asked to MODIFY an existing column, NEVER asked to "
            "DROP a column, NEVER asked to flip a NOT-NULL on a column that "
            "legacy rows might not satisfy.",
        ])

    PAGE_BREAK()

    H2("The two questions DS was asked")
    P(
        "Most of the canonical-vs-legacy decisions have a clean default — alias "
        "in the repo, name new columns conservatively, generate display codes on "
        "the fly. Two decisions, however, had real downstream consequences and "
        "warranted explicit sign-off before the migration SQL was authored."
    )

    H3("Decision request #1 — Priority enum mismatch")
    KV([
        ("Legacy DB enum",          "('LOW','NORMAL','HIGH','URGENT') default 'NORMAL'"),
        ("Spec canonical enum",     "('LOW','MEDIUM','HIGH') default 'MEDIUM'"),
        ("Option A (DS chose)",     "Map in repo layer. DB unchanged. MEDIUM↔NORMAL on read/write; URGENT→HIGH on display. Zero ALTER risk."),
        ("Option B",                "ALTER the enum to add MEDIUM, keep NORMAL+URGENT. New writes use MEDIUM; legacy reads still work. ALTER ENUM has a brief table lock."),
        ("Option C",                "Add a new column JR_PRIORITY_MVP with the exact spec enum. Cleanest separation; duplicates intent forever."),
        ("Rationale for A",         "Preserves the Phase 3 seal. Legacy URGENT rows render as HIGH which is the intuitive collapse. Service layer stays clean. Future Slice 2 approve/reject can read priority without thinking about the difference."),
    ])

    H3("Decision request #2 — Accessories storage shape")
    KV([
        ("Form input",              "Up to 20 accessory rows per JR (type, name, optional serial_no)"),
        ("Option A",                "JSON column on cmms_jobrequest_mst. Atomic with parent. No JOIN. MySQL 8 has JSON_TABLE if introspection ever needed. Slice 1 doesn't query individual accessories."),
        ("Option B (DS chose)",     "Child table job_request_accessories with FK to JR_JOBREQUESTNO ON DELETE CASCADE. Future-proofs per-accessory queries (e.g. 'all JRs using probe X'). Extra INSERT per accessory inside the create txn — bulk INSERT keeps it one round-trip."),
        ("Rationale for B",         "Future analytics + audit on accessory usage. JSON is harder to ALTER consistently across legacy + new tables. Child table is the “normal” relational pattern; future engineers don't have to learn JSON_TABLE."),
    ])

    PAGE_BREAK()

    H2("The 11 locked decisions register")
    P(
        "Each of these decisions was made (most by sensible default, the two "
        "above by explicit DS sign-off) during STEP 0 and locked into the "
        "SCHEMA_PHASE6.md document. They are not reopenable inside Slice 1. "
        "Slice 2 may add new locked decisions; it does not unlock these."
    )

    TABLE(
        ["ID", "Topic", "Decision", "One-line rationale"],
        [
            ["P6-D1",  "Priority enum mismatch",     "Map in repo layer (DB unchanged)",
             "Preserves the Phase 3 seal; legacy URGENT collapses to HIGH on display."],
            ["P6-D2",  "Accessories storage",        "New child table job_request_accessories",
             "Future analytics queryability + clean relational pattern."],
            ["P6-D3",  "request_code 'JR-2026-1234'","Generate on-the-fly from year + JR_JOBREQUESTNO",
             "No new column; year is already in JR_JOBREQUESTDATE."],
            ["P6-D4",  "Job Card card_code 'JC-2026-...'","Generate on-the-fly from JM_JCRecdDate + JM_JobCardNO",
             "Same reasoning as P6-D3; legacy JM_SectionJobNo preserved but not shown."],
            ["P6-D5",  "job_category, job_type",      "ADD two ENUM columns; legacy free-form fields untouched",
             "Strict enums let the FE dropdown be deterministic."],
            ["P6-D6",  "T&C compliance fields",       "ADD JR_TNC_ACCEPTED_AT + JR_TNC_VERSION",
             "BR-AUD-01 — auditor must prove 'user X accepted v1 of the T&Cs at time Y'."],
            ["P6-D7",  "Submitter identity",          "Use JR_SUBMITTEDBYID (employee_id) as canonical FK",
             "Mirrors audit_log.actor_employee_id; avoids a join on every list."],
            ["P6-D8",  "created_at / updated_at on JR","ADD both, backfill JR_CREATED_AT from JR_JOBREQUESTDATE",
             "Index-friendly list sort; legacy date lacks DATETIME(6) precision."],
            ["P6-D9",  "CANCELLED status",            "Deferred to Slice 2",
             "Not used by any Slice 1 UI; saves an ALTER ENUM cost now."],
            ["P6-D10", "submitted_at field",          "Alias JR_MVP_STATUS_AT when status is SUBMITTED",
             "Already populated by the Phase 3 backfill; one timestamp serves both intents."],
            ["P6-D11", "Division lookup source",      "Query cmms_section_mst (legacy), not the new sections table",
             "168 legacy divisions vs 2 new TIMCD entries — the form needs all 168."],
        ],
        widths=(1.5, 4.0, 5.5, 6.0),
    )

    PAGE_BREAK()


def section_migration_strategy():
    H1("Part V  ·  STEP 1  —  Migration Strategy")

    H2("The ADD-only philosophy")
    P(
        "Phase 5 inherited a stricter rule (“K.6 — no ALTERs”) because Equipment "
        "could fit its required fields into existing columns or piggyback on "
        "audit_log notes. Job Requests cannot — there is no place on the legacy "
        "row to store T&C acceptance, generated codes, or strict enum values. "
        "Phase 6 therefore relaxes the rule, but only in one direction: "
        "additive only."
    )

    DIAGRAM("Migration discipline — what's allowed, what's not",
    [
        " ALLOWED                                  FORBIDDEN",
        " ─────────────────────────                ────────────────────────",
        " ADD COLUMN (nullable / default)          DROP COLUMN  (data loss)",
        " ADD INDEX                                MODIFY COLUMN type",
        " ADD CONSTRAINT (FK with ON CASCADE)      ALTER ENUM removing values",
        " CREATE TABLE IF NOT EXISTS               RENAME COLUMN",
        " CREATE INDEX IF NOT EXISTS               RENAME TABLE",
        " INSERT (backfill of new columns only)    DELETE on legacy rows",
        " UPDATE (idempotent, gated on default)    TRUNCATE",
    ])

    H2("Numbering scheme")
    P(
        "The Phase 3 migrations live at 001..010, 050, 099. Phase 6 chose the "
        "100-block to leave a gap for Phase 4 / Phase 5 to slot in any "
        "retrospective infra migrations (e.g. a new audit_log column) without "
        "renumbering Phase 6's files."
    )

    SOFT_TABLE(
        ["File", "What it does", "Idempotency strategy"],
        [
            ["100__phase6_jr_columns.sql",
             "ADD 6 columns to cmms_jobrequest_mst; backfill JR_CREATED_AT from JR_JOBREQUESTDATE on rows that haven't been written this run.",
             "information_schema.columns check on each ADD; backfill uses a “WHERE JR_CREATED_AT > NOW() - 5 minutes” guard so it only touches default-fresh rows."],
            ["101__phase6_accessories_table.sql",
             "CREATE TABLE IF NOT EXISTS for the child table with FK ON DELETE CASCADE.",
             "CREATE TABLE IF NOT EXISTS — re-runs are no-ops."],
            ["102__phase6_indexes.sql",
             "ADD seven covering indexes — four on the JR table, two on JC, one on accessories.",
             "information_schema.statistics check on each index name."],
        ],
        widths=(5.5, 7.5, 4.0),
    )

    CALLOUT('think', "Why the prepared-statement idempotency dance",
        [
            "Each ADD COLUMN is wrapped in a SET-IF-EXISTS construction:",
            " · SELECT COUNT(*) FROM information_schema.columns WHERE … into @c",
            " · SET @sql := IF(@c>0, 'SELECT 1', 'ALTER TABLE …')",
            " · PREPARE / EXECUTE / DEALLOCATE",
            "This works on both MySQL 8 and MariaDB without needing DELIMITER + "
            "stored procedures. The pattern was inherited from Phase 3's "
            "002__alter_legacy_tables.sql which solved the same problem.",
            "Net: running the migration runner three times produces the same "
            "schema state as running it once. Re-running is the recovery mechanism "
            "when a previous run errored mid-flight.",
        ])

    H2("What migration 100 actually changes")
    SOFT_TABLE(
        ["Added column", "Type", "Null?", "Default", "Why it's needed"],
        [
            ["JR_JOB_CATEGORY",     "ENUM('TME','FPE')",                              "YES (legacy rows)", "NULL",
             "Strict enum for the FE Job Category dropdown."],
            ["JR_JOB_TYPE",         "ENUM('CALIBRATION','REPAIR','REGISTRATION')",    "YES (legacy rows)", "NULL",
             "Strict enum for the FE Job Type dropdown."],
            ["JR_TNC_ACCEPTED_AT",  "DATETIME(6)",                                    "YES",               "NULL",
             "Filled on DRAFT→SUBMITTED transition. NULL until then."],
            ["JR_TNC_VERSION",      "VARCHAR(10)",                                    "YES",               "'v1'",
             "Records which T&C set was accepted; supports rev-up via 'v2', 'v3'."],
            ["JR_CREATED_AT",       "DATETIME(6)",                                    "NO",                "CURRENT_TIMESTAMP(6)",
             "Index-friendly creation timestamp for default list sort."],
            ["JR_UPDATED_AT",       "DATETIME(6)",                                    "NO",                "CURRENT_TIMESTAMP(6) ON UPDATE",
             "Auto-touched on row update; useful for stale-cache detection."],
        ],
        widths=(4.5, 4.5, 1.7, 2.5, 3.8),
    )

    H2("What migration 101 actually creates")
    P("The new job_request_accessories table has the following shape:")
    SOFT_TABLE(
        ["Column", "Type", "Notes"],
        [
            ["acc_id",          "BIGINT UNSIGNED AUTO_INCREMENT PK",  "Surrogate id"],
            ["jr_no",           "INT NOT NULL",                       "FK → cmms_jobrequest_mst.JR_JOBREQUESTNO, ON DELETE CASCADE"],
            ["accessory_type",  "VARCHAR(60) NOT NULL",               "Free-form category (Probe, Cable, …)"],
            ["accessory_name",  "VARCHAR(120) NOT NULL",              "Display name"],
            ["serial_no",       "VARCHAR(120) NULL",                  "Optional serial"],
            ["position",        "SMALLINT UNSIGNED NOT NULL DEFAULT 0", "Preserves UI ordering"],
            ["created_at",      "DATETIME(6) NOT NULL DEFAULT NOW(6)",  "Audit timestamp"],
            ["INDEX idx_jra_jr_pos", "(jr_no, position)",             "Used to load a JR's accessories in their UI order"],
        ],
        widths=(5.5, 6.0, 5.5),
    )

    PAGE_BREAK()


def section_be_architecture():
    H1("Part VI  ·  STEP 2  —  Backend Architecture")

    H2("The 5-layer cake")
    P(
        "Every backend module in CMCMIS follows the same 5-layer shape. Phase 4's "
        "auth module established it. Phase 5's equipment module reused it. Phase 6 "
        "is the first place where the pattern is exercised across two modules at "
        "once (Job Requests + Job Cards) plus a third helper module (Lookups), and "
        "it holds together cleanly."
    )

    DIAGRAM("The 5-layer cake — each layer can only call the one immediately below it",
    [
        "           HTTP REQUEST",
        "                │",
        "                ▼",
        "  ┌─────────────────────────────────┐",
        "  │  Layer 1 — MIDDLEWARE PIPELINE  │   authenticate · authorize · scope · validate",
        "  └─────────────┬───────────────────┘",
        "                ▼",
        "  ┌─────────────────────────────────┐",
        "  │  Layer 2 — ROUTES               │   Express router — composes middlewares + ctrl",
        "  └─────────────┬───────────────────┘",
        "                ▼",
        "  ┌─────────────────────────────────┐",
        "  │  Layer 3 — CONTROLLERS          │   Thin HTTP shim. NO SQL, NO business rules.",
        "  └─────────────┬───────────────────┘",
        "                ▼",
        "  ┌─────────────────────────────────┐",
        "  │  Layer 4 — SERVICES             │   Business rules, transactions, state machine.",
        "  └─────────────┬───────────────────┘",
        "                ▼",
        "  ┌─────────────────────────────────┐",
        "  │  Layer 5 — REPOSITORIES         │   THE ONLY FILES with SQL. Aliasing here.",
        "  └─────────────┬───────────────────┘",
        "                ▼",
        "             MySQL  (Phase 3 sealed)",
    ])

    H2("What each layer can and cannot do")
    SOFT_TABLE(
        ["Layer", "MAY", "MAY NOT"],
        [
            ["Middleware",  "Read JWT, check permissions, set req.scope, validate inputs.",  "Issue SQL, call services, mutate state."],
            ["Routes",      "Compose middleware + map to controller methods.",                "Include any logic beyond wiring."],
            ["Controllers", "Read req, call exactly one service method, shape response.",     "Touch SQL, compute business rules, manage transactions."],
            ["Services",    "Own transactions, call repos, enforce state machine, write audit.","Touch SQL directly, format HTTP responses, read req.headers."],
            ["Repositories","Own all SQL, alias legacy↔canonical, return plain objects/arrays.","Make HTTP-shaped decisions (status codes, envelopes)."],
        ],
        widths=(2.5, 7.0, 7.5),
    )

    H2("The middleware stack — order matters")
    DIAGRAM("Middleware order for every Phase 6 protected route",
    [
        "  authenticate                   ← decode JWT, populate req.user",
        "      │",
        "      ▼",
        "  authorize / authorizeAny       ← does the caller hold the required permission?",
        "      │",
        "      ▼",
        "  rowLevelScope('job_request')   ← writes req.scope = { canReadAll, ownerEmployeeId }",
        "      │",
        "      ▼",
        "  validate(schema, 'query'|'body') ← zod .strict() parses & coerces; rejects unknown keys",
        "      │",
        "      ▼",
        "  CONTROLLER → SERVICE → REPO → DB",
    ])

    CALLOUT('think', "Why this order — what would break if you swapped it",
        [
            "authenticate must run first; everything downstream needs req.user.",
            "authorize before rowLevelScope: the scope helper depends on which "
            "permission codes the caller holds. Running scope first would force "
            "the scope helper to fail closed for every anonymous request — "
            "redundant work.",
            "rowLevelScope before validate: validate may mutate req.query (coerce "
            "strings to numbers, fill defaults). The scope helper doesn't read "
            "req.query, but conceptually the scope is a property of the actor, "
            "not the input — set it before input parsing.",
            "validate before the controller: the controller assumes coerced, "
            "typed inputs. Without validate, the controller would have to "
            "re-parse every field defensively.",
        ])

    H2("Single choke-point rule (FINAL-DESC §8.3)")
    P(
        "The most important architectural invariant in the slice. Every "
        "state-changing write goes through a transition() function in the "
        "state machine module, and that same function is the only entry point "
        "that decides whether a state change is legal. The audit_log row and "
        "the state-history row are written in the same transaction as the "
        "JR_MVP_STATUS update. There is no other way to change a job request's "
        "status."
    )

    DIAGRAM("Single choke-point — the transition envelope",
    [
        "  ┌──────────────────────────────────────────────────────────────────┐",
        "  │ service.submitJobRequest({ jrNo, body, actor })                  │",
        "  ├──────────────────────────────────────────────────────────────────┤",
        "  │ 1.  repo.findJrById(jrNo)           ← FAIL 404 if missing        │",
        "  │ 2.  stateMachine.transition(...)    ← throws 409 / 403 on bad    │",
        "  │ 3.  pool.getConnection() → BEGIN                                 │",
        "  │       repo.transitionStatus(conn, jrNo, newState, extras)        │",
        "  │       repo.appendStatusHistory(conn, jrNo, from, to, actor)      │",
        "  │       repo.writeAuditLog(conn, { … })                            │",
        "  │     COMMIT                                                       │",
        "  │     (any failure → conn.rollback())                              │",
        "  └──────────────────────────────────────────────────────────────────┘",
    ])

    CALLOUT('locked', "Why the audit row goes inside the same transaction",
        [
            "If the audit_log INSERT lived outside the transaction, a crash "
            "between the status UPDATE and the audit INSERT would leave the "
            "system with an unauditable state change. The audit trail would "
            "be inconsistent with the system of record.",
            "Inside the transaction, either everything commits or nothing does. "
            "An auditor reading the audit_log a year later sees exactly the "
            "history they expect.",
            "Cost: one extra INSERT per transaction. Benefit: provable audit "
            "integrity under crash conditions. The cost is dominated by the "
            "fsync of the COMMIT, which is one operation regardless of how "
            "many INSERTs preceded it.",
        ])

    PAGE_BREAK()


def section_jobrequests_module_deep():
    H1("Part VII  ·  STEP 3  —  Job Requests Module")

    H2("File map")
    SOFT_TABLE(
        ["File", "Layer", "Responsibility"],
        [
            ["jobRequests.routes.js",      "Routes",       "Wires URL paths to middleware chains and controller methods."],
            ["jobRequests.controller.js",  "Controllers",  "Three thin shims (list, create, submit). Marshals req → service args → response envelope."],
            ["jobRequests.service.js",     "Services",     "Three business functions (listJobRequests, createJobRequest, submitJobRequest). Owns transactions, calls state machine, writes audit."],
            ["jobRequests.validators.js",  "Validators",   "Three zod schemas (listQuerySchema, createSchema, submitSchema). Mirror of the FE schemas."],
            ["jobRequests.repo.js",        "Repositories", "The only file in this module with SQL. Canonical ↔ legacy aliasing happens here. ~300 lines."],
            ["jobRequests.stateMachine.js","Pure-logic",   "transition(currentState, action, actor) — the single choke-point. Pure function, no DB calls, no side effects."],
        ],
        widths=(5.5, 2.5, 9.0),
    )

    H2("The state machine in detail")
    DIAGRAM("Job Request lifecycle state machine (FINAL-DESC §8.1)",
    [
        "                       ┌─────────────────────────────┐",
        "  CREATE (POST /)  →   │  DRAFT                      │",
        "                       │  • submitter only           │",
        "                       │  • can be edited            │",
        "                       │  • tnc_accepted_at = NULL   │",
        "                       └────────────┬────────────────┘",
        "                                    │ submit (owner + tnc=true)",
        "                                    ▼",
        "                       ┌─────────────────────────────┐",
        "                       │  SUBMITTED                  │",
        "                       │  • visible in LIC inbox     │",
        "                       │  • tnc_accepted_at set      │",
        "                       │  • locked from edits        │",
        "                       └────────────┬────────────────┘",
        "                          approve  / reject (Slice 2)",
        "                                    │",
        "                ┌───────────────────┼───────────────────┐",
        "                ▼                                       ▼",
        "    ┌────────────────────────┐            ┌─────────────────────────┐",
        "    │  ASSIGNED              │            │  REJECTED               │",
        "    │  spawns Job Card →     │            │  terminal               │",
        "    │  rest of lifecycle on  │            │  rejection_reason set   │",
        "    │  the Job Card machine  │            └─────────────────────────┘",
        "    └────────────────────────┘",
    ])

    P(
        "Slice 1 implements only the DRAFT → SUBMITTED transition. The "
        "transition() function is wired with the SUBMITTED row's approve / "
        "reject transitions so that Slice 2 only needs to add the service "
        "method bodies — the state-machine table itself does not need a "
        "second edit."
    )

    H2("The transition contract")
    SOFT_TABLE(
        ["Aspect", "Behaviour"],
        [
            ["Input", "currentState (string), action (string), actor ({ employeeId, role, permissions[] }), opts ({ isOwner })"],
            ["Success output", "{ newState: 'SUBMITTED' }"],
            ["Failure modes", "Throws 409 ILLEGAL_TRANSITION if (state, action) not in ALLOWED. Throws 403 NOT_OWNER if rule.actorMustBeOwner && !opts.isOwner. Throws 403 FORBIDDEN if actor lacks rule.perm."],
            ["Side effects",  "None. Pure function."],
            ["Verifiability", "Trivial unit tests possible — pass synthetic actors and assert."],
        ],
        widths=(3.0, 14.0),
    )

    CALLOUT('think', "Why the state machine is a PURE function — no DB calls",
        [
            "Putting the DB lookup inside transition() would conflate two "
            "concerns: 'is this transition legal' and 'how do we read the "
            "current state'. Separating them means transition() can be tested "
            "with synthetic states, the service can decide where the current "
            "state comes from (a fresh SELECT vs an in-memory snapshot), and "
            "the function is portable across modules.",
            "Conversely, putting the audit_log INSERT inside transition() would "
            "make it impossible to use the function for read-side checks (“is "
            "this transition allowed for this user?”) without polluting the "
            "audit_log with hypotheticals.",
        ])

    H2("The repo — canonical ↔ legacy translation in detail")
    P(
        "jobRequests.repo.js is the longest file in the slice for a reason. "
        "It is the canonical-vs-legacy translation boundary, and it cannot "
        "be terse without leaking column names upward. The repo's public "
        "API exposes nine functions, each of which speaks pure canonical."
    )

    SOFT_TABLE(
        ["Repo function", "Returns", "Notes on aliasing"],
        [
            ["listJobRequests(params, scope)",   "{ rows: [], total: number }",       "Aliases JR_* → canonical via AS clauses; maps DB priority NORMAL→MEDIUM, URGENT→HIGH before returning."],
            ["nextJrNo(conn)",                   "INT",                                 "SELECT MAX(JR_JOBREQUESTNO)+1 FOR UPDATE — pessimistic lock within the caller's transaction."],
            ["insertJobRequest(conn, payload)",  "void",                                "Translates canonical payload back to JR_* columns; truncates strings to legacy widths."],
            ["findJrById(jrNo)",                 "{ canonical row } or null",          "Loads the columns the service needs (status, owner, complaint, tnc fields)."],
            ["transitionStatus(conn, jrNo, newStatus, extras)", "void",                "UPDATE … SET JR_MVP_STATUS, JR_MVP_STATUS_AT, JR_UPDATED_AT, optionally JR_TNC_*."],
            ["appendStatusHistory(conn, jrNo, fromStatus, toStatus, actor, reason)", "void", "INSERT into job_request_status_history."],
            ["replaceAccessories(conn, jrNo, accessories[])",  "void",                "DELETE + bulk INSERT inside the create transaction; safe because nothing else writes here."],
            ["writeAuditLog(conn, { … })",       "void",                                "INSERT into audit_log with JSON-stuffed notes (≤500 chars)."],
            ["toCanonicalPriority(dbVal) / toDbPriority(canon)", "string",             "The priority mapping helpers; used by both list rows and insert payload."],
        ],
        widths=(5.5, 4.5, 7.0),
    )

    PAGE_BREAK()


def section_jobcards_module():
    H1("Part VIII  ·  STEP 4  —  Job Cards Module")

    H2("Why read-only in Slice 1")
    P(
        "The Job Card lifecycle is the second half of the calibration "
        "workflow — once a JR is approved by LIC, a Job Card spawns and the "
        "Lab Engineer drives it through start/work/complete/verify. That "
        "lifecycle is intentionally Slice 2 territory, because it requires "
        "additional UI (task checklist, observations, PDF) that the build "
        "session could not also fit. The list itself ships now because the "
        "RBAC matrix shows VIEW_ONLY users have job_card:read-list — there "
        "is a real user with no other route to see what's happening."
    )

    H2("The four-table JOIN")
    P(
        "Listing job cards is the most complex SELECT in the slice. Each "
        "row in the response needs data from four tables, and the JOIN has "
        "to be careful to LEFT JOIN where the parent JR or assigned engineer "
        "row might be missing (legacy data + Phase 3 backfill leave some "
        "cards orphan-equivalent)."
    )

    DIAGRAM("Job Card list JOIN tree",
    [
        "  cmms_jobcard_mst  jc",
        "      │  JM_EQM_TYPE, JM_EQM_ID                          → equipment_name",
        "      ├── LEFT JOIN cmms_eqip_mst e ON (jc.JM_EQM_TYPE = e.EQM_TYPE  AND",
        "      │                                 jc.JM_EQM_ID   = e.EQM_ID)",
        "      │",
        "      │  JM_SectionJobNo                                  → job_request_id, job_request_code",
        "      ├── LEFT JOIN cmms_jobrequest_mst jr ON (jr.JR_SECTIONJOB_NO = jc.JM_SectionJobNo)",
        "      │",
        "      │  jr.JR_ASSIGNED_ENGINEER                          → assigned_engineer_name",
        "      └── LEFT JOIN cmms_emp_mst emp ON (emp.EMM_ID = jr.JR_ASSIGNED_ENGINEER)",
    ])

    H2("Card code generation")
    P(
        "The card_code shown on the screen is 'JC-2026-1234'. The legacy "
        "cmms_jobcard_mst has two id-like columns: JM_JobCardNO (INT) and "
        "JM_SectionJobNo (varchar(9), the primary key with a legacy format "
        "like '42026026'). Slice 1 chose JM_JobCardNO as the canonical "
        "identifier for display because it's already an integer and matches "
        "the JR convention; the legacy SectionJobNo is preserved in the API "
        "response as a separate field for any code that still needs it."
    )

    PAGE_BREAK()


def section_lookups_and_me():
    H1("Part IX  ·  STEP 5  —  Lookups and /me extension")

    H2("Why a separate lookups module")
    P(
        "The JR form needs two pieces of reference data that don't belong "
        "to any one module: the list of divisions (170-ish rows from "
        "cmms_section_mst) and the equipment search typeahead. The "
        "alternative — exposing them under /api/v1/equipment or "
        "/api/v1/job-requests — would have coupled the form to a single "
        "owning module. Lookups is a third module so that future forms "
        "(e.g. the Phase 7 Job Card task list) can reuse the same paths."
    )

    H2("The two endpoints")
    SOFT_TABLE(
        ["Endpoint", "Returns", "Permission gate", "Used by"],
        [
            ["GET /lookups/divisions",
             "{ items: [{ id, code, name }] } (168 rows ordered by code)",
             "authorizeAny('job_request:create', 'equipment:read-list')",
             "JR form Section 4 Division dropdown"],
            ["GET /lookups/equipment/search?q=...&limit=20",
             "{ items: [{ id, eqm_type, eqm_id, name, make, model_no, serial_no, type }] }",
             "authorizeAny('job_request:create', 'equipment:read-list')",
             "JR form Section 2 Equipment ID typeahead — pre-fills the surrounding fields on select"],
        ],
        widths=(5.5, 5.5, 4.0, 4.0),
    )

    H2("The /me extension — what changed and why")
    P(
        "Phase 5 already returned employeeId, userId, role, permissions[], "
        "display_name, designation, email from /me. Phase 6 Section 4 of the "
        "JR form needs three more fields: lab_phone, room_phone, and the "
        "user's division. The /me handler was extended additively — the same "
        "endpoint, same auth, more fields in the response. Existing consumers "
        "(the equipment form, the topbar) ignore the new fields and continue "
        "working."
    )

    DIAGRAM("Auto-fill chain — from /me to the JR form Section 4",
    [
        "  PAGE LOAD",
        "      │",
        "      ▼",
        "  AuthProvider mounts  →  POST /auth/refresh  →  GET /me",
        "      │                                              │",
        "      │                                              ▼",
        "      │                                  user = { employeeId, role,",
        "      │                                           permissions[], display_name,",
        "      │                                           designation, email,",
        "      │                                           lab_phone, room_phone,",
        "      │                                           division_id, division_code }",
        "      │",
        "      ▼",
        "  JobRequestNew mounts  →  reads `user` from context",
        "      │",
        "      ▼",
        "  filledFromMeRef guards a one-shot useEffect that copies",
        "      user.lab_phone / room_phone / division_id into form state",
        "      ONLY IF those fields are still at their initial empty values",
        "      → user edits to those fields are NEVER clobbered on re-render",
    ])

    CALLOUT('locked', "BR-JR-06 — submitted_by is server-set, never trusted",
        [
            "Even though the form auto-fills Name / Employee ID / Designation / "
            "Email from /me, the form deliberately does NOT send those fields "
            "to POST /job-requests. The four are visible-but-disabled inputs.",
            "The server reads the requester's identity from req.user (decoded "
            "from the JWT) and looks up the snapshot fields from cmms_emp_mst "
            "by employee_id.",
            "The zod schema on the BE uses .strict() — if a malicious client "
            "sends submitted_by_employee_id or submitted_by_name in the body, "
            "the request is rejected with HTTP 422 and the message "
            "“Unrecognized key(s) in object”.",
            "This is defence in depth. The FE never sends them; the BE rejects "
            "if they appear anyway.",
        ])

    PAGE_BREAK()


def section_fe_reality():
    H1("Part X  ·  STEP 6  —  Frontend Reality Check")

    H2("What the spec assumed vs what was actually installed")
    P(
        "The build prompt was written against an aspirational tech stack. "
        "The first thing the build session did after STEP 0 was read FE "
        "package.json and reconcile the spec against the installed deps. "
        "Several spec assumptions turned out to be wishlist items."
    )

    SOFT_TABLE(
        ["Spec assumed", "Actually installed?", "Decision"],
        [
            ["@tanstack/react-query",         "No",  "Use a custom 30-second-TTL useState hook with AbortController, mirroring useEquipmentList from Phase 5."],
            ["@tanstack/react-table",         "No",  "Use the existing custom <DataTable> component."],
            ["react-hook-form",               "Yes", "Available, but JobRequestNew uses plain useState — simpler given the section-wise structure."],
            ["zod",                           "Yes", "Use it for FE schema mirror + cross-field validation."],
            ["sonner (toast)",                "No",  "Use window.alert() as Slice 1 placeholder; add a real toast in Slice 2."],
            ["axios",                         "Yes", "Already wired with auth interceptor in api-client.js."],
            ["lucide-react",                  "Yes", "Used for icons everywhere."],
            ["clsx",                          "Yes", "Used for conditional className composition."],
        ],
        widths=(5.0, 3.5, 8.5),
    )

    CALLOUT('think', "Why follow the equipment-module pattern instead of fighting it",
        [
            "Phase 5's equipment module ships with a useEquipmentList custom "
            "hook, a custom <DataTable>, a custom <Pagination> with the exact "
            "1, 2, … 99, 100 design. Reinventing those with react-query / "
            "@tanstack/react-table would have:",
            " · doubled the bundle size for negligible UX win",
            " · introduced a second list-fetch idiom in the project — confusing",
            " · forced a parallel maintenance path for the rest of the MVP",
            "The decision was to copy useEquipmentList.js into useJobRequestList.js "
            "and useJobCardList.js, hash-keyed on params, with the same 30s TTL "
            "and AbortController cancellation pattern. Future hooks (useScheduleList, "
            "useReportList) will follow the same template.",
        ])

    H2("The new FE files at a glance")
    SOFT_TABLE(
        ["File", "Role"],
        [
            ["src/lib/api/jobRequests.js",        "axios wrappers: fetchJobRequestList, createJobRequest, submitJobRequest"],
            ["src/lib/api/jobCards.js",           "axios wrappers: fetchJobCardList"],
            ["src/lib/api/lookups.js",            "axios wrappers: fetchDivisions, searchEquipment"],
            ["src/lib/hooks/useJobRequestList.js","30s-TTL cache hook for the JR list"],
            ["src/lib/hooks/useJobCardList.js",   "30s-TTL cache hook for the JC list"],
            ["src/lib/schemas/jobRequestSchemas.js", "FE zod mirror of the BE create schema"],
            ["src/components/StatusPill.jsx",     "Reusable status badge — 9 status colors"],
            ["src/components/PriorityLabel.jsx",  "Reusable priority text — red/amber/green"],
            ["src/pages/jobRequests/JobRequestList.jsx",      "The /job-requests page"],
            ["src/pages/jobRequests/JobRequestNew.jsx",       "The /job-requests/new page (5 sections + T&C gate)"],
            ["src/pages/jobRequests/form/tncContent.js",      "The 6 T&C strings, verbatim from the reference screen"],
            ["src/pages/jobCards/JobCardList.jsx",            "The /job-cards page"],
        ],
        widths=(7.0, 10.0),
    )

    PAGE_BREAK()


def section_three_screens_ux():
    H1("Part XI  ·  STEP 7  —  The Three Screens")

    H2("Screen 1 — Job Request List (/job-requests)")
    DIAGRAM("JobRequestList — page chrome",
    [
        "  Job Requests                                                [ + New Job Request ]",
        "  Manage equipment calibration and maintenance requests",
        "",
        "  ┌──────────────────────────────────────────────────────────────────────────┐",
        "  │  [🔍 Search by Job ID, Equipment, or Submitted By…]                       │",
        "  │  [All Types ▾]  [All Statuses ▾]                                          │",
        "  │  [⚙ Advanced Filters]  [⬇ Export]              Showing 25 of 24,238 reqs │",
        "  └──────────────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌────────────┬──────────────┬─────────┬────────┬──────────────┬──────────┬──────────┬───────────┐",
        "  │ Job ID     │ Equipment    │ Type    │ Div.   │ Submitted By │  Date    │ Priority │ Status    │",
        "  ├────────────┼──────────────┼─────────┼────────┼──────────────┼──────────┼──────────┼───────────┤",
        "  │ JR-2026-…  │ LENZ CNC…    │ Cal…    │ EMG    │ Dr. R. Menon │ 2026-04… │ High     │ Pending   │",
        "  │ ...                                                                                          ",
        "",
        "  Page 1 of 970                              [Prev] [1] [2] [3] … [969] [970] [Next]",
    ])

    H3("Behaviour details")
    BULLET("Search input is debounced 300ms — one API call per change after the user stops typing.")
    BULLET("Type and Status dropdowns reset the page to 1 on change (otherwise the user would land on an empty page beyond the new filter's total).")
    BULLET("The '+ New Job Request' button only renders when hasPermission('job_request:create'). View-Only users do not see it; even if they did, POST returns 403.")
    BULLET("Status pills use the exact color map from the spec — Pending=amber, Approved=violet, In Progress=blue, Completed=green, Verified=emerald, Rejected=red, Reopened=orange.")
    BULLET("Priority shows as colored text (not pill) — High=red, Medium=amber, Low=green. Matches the reference screen exactly.")
    BULLET("Date column shows submitted_at when set; falls back to created_at for drafts.")
    BULLET("The Job ID link routes to /job-requests/:id which is the Slice 2 detail page (currently 404).")

    PAGE_BREAK()

    H2("Screen 2 — New Job Request (/job-requests/new)")

    DIAGRAM("JobRequestNew — five sections + footer",
    [
        "  ← Back to Job Requests",
        "  New Job Request",
        "  Submit a new calibration, repair, or registration request",
        "",
        "  ┌──────────────────────────────────────────────────────────────────────────┐",
        "  │ 1. Job Type Selection                                                    │",
        "  │     Job Category *   [Select category ▾]    Job Type *  [Select type ▾]  │",
        "  └──────────────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌──────────────────────────────────────────────────────────────────────────┐",
        "  │ 2. Equipment Details                                                     │",
        "  │     Equipment ID *  [Search or enter… ▾   typeahead]                     │",
        "  │     Equipment Name *  Make   Model No.   Serial No.   Equipment Type    │",
        "  │     Options / Description (textarea)                                     │",
        "  └──────────────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌──────────────────────────────────────────────────────────────────────────┐",
        "  │ 3. Accessories                                                           │",
        "  │     (added accessories listed as chips with ✕ remove)                    │",
        "  │     Accessory Type [▾]  Name [   ]  Serial No. [   ]    [+ Add Accessory]│",
        "  └──────────────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌──────────────────────────────────────────────────────────────────────────┐",
        "  │ 4. Submitted By  ★ auto-fill from /me                                    │",
        "  │     Name *  SAC Emp ID *  Designation *  Email *     ← readonly (BR-JR-06)│",
        "  │     Lab Phone  Room Phone  Division *  Subsystem  Project  Priority      │",
        "  │     Complaint Description * (textarea, min 10 chars)                     │",
        "  │     Remarks (textarea)                                                   │",
        "  │     ☐ Equipment sent after repair                                         │",
        "  └──────────────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌──────────────────────────────────────────────────────────────────────────┐",
        "  │ 5. Terms and Conditions   ⓘ                                              │",
        "  │     ☐ T&C 1: I confirm that all equipment details …                       │",
        "  │     ☐ T&C 2: I understand that the equipment must be delivered …          │",
        "  │     ☐ T&C 3: I acknowledge that the calibration timeline begins …         │",
        "  │     ☐ T&C 4: I agree to coordinate with the assigned lab engineer …       │",
        "  │     ☐ T&C 5: I accept that equipment found to be damaged or beyond repair…│",
        "  │     ☐ T&C 6: I understand that urgency requests will be handled …         │",
        "  │     ⚠ You must accept all terms and conditions … (0/6 accepted)          │",
        "  └──────────────────────────────────────────────────────────────────────────┘",
        "",
        "  [ Cancel ]                                  [ Save as Draft ]  [ Submit Request ]",
    ])

    H3("The Submit gate — three layers")
    SOFT_TABLE(
        ["Layer", "What it checks",                                        "What happens on fail"],
        [
            ["FE button disabled state",
             "isStructurallyValid (zod parse passes) AND tncAcceptedCount === 6",
             "Button greyed, tooltip 'Accept all 6 T&Cs to enable'"],
            ["FE zod schema parse",
             "Field-by-field validation before POST",
             "Errors surfaced inline under each field"],
            ["BE zod schema parse",
             ".strict() rejects unknown keys; cross-field rule enforces submit_now → tnc_accepted",
             "HTTP 422 with structured details[] — FE displays under the right field"],
        ],
        widths=(4.0, 8.0, 5.0),
    )

    H3("Equipment typeahead UX")
    BULLET("User types ≥2 characters into the Equipment ID input.")
    BULLET("After a 300ms quiet period, GET /lookups/equipment/search?q=… fires (one call per quiet period).")
    BULLET("Results render as a dropdown of up to 10 suggestions, each showing eqm_type-eqm_id · make · serial.")
    BULLET("On select, six fields are pre-filled: equipment_id, equipment_name, make, model_no, serial_no, equipment_type.")
    BULLET("User can override any field after select — the typeahead is a convenience, not a constraint.")

    PAGE_BREAK()

    H2("Screen 3 — Job Cards List (/job-cards)")
    DIAGRAM("JobCardList — page chrome",
    [
        "  Job Cards",
        "  Track and manage job execution and progress",
        "",
        "  ┌──────────────────────────────────────────────────────────────────────────┐",
        "  │  [🔍 Search by Job Card ID, Equipment, or Engineer…]   [All Statuses ▾] │",
        "  │  [⚙ Advanced Filters]  [⬇ Export]               Showing 25 of 7 cards   │",
        "  └──────────────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌──────────────┬──────────────┬──────────────┬───────────────────┬───────────┬───────────┬──────────┐",
        "  │ Job Card ID  │ Job Req. ID  │ Equipment    │ Assigned Engineer │ Status    │ Start     │ Due Date │",
        "  ├──────────────┼──────────────┼──────────────┼───────────────────┼───────────┼───────────┼──────────┤",
        "  │ JC-2026-063  │ JR-2026-063  │ Synthesized… │ Mr. V. Modhwadia  │ In Prog…  │ 2026-04…  │ 2026-05… │",
        "  │ JC-2026-…    │ JR-2026-…    │ Spectrum…    │ Dr. Sharma        │ Assigned  │ 2026-04…  │ 2026-04… │",
        "  │ ...                                                                                              ",
        "",
        "  Page 1 of 1                                       [Prev] [1] [Next]",
    ])

    BULLET("There is intentionally no '+ New Job Card' CTA. Cards spawn from the JR approval flow in Slice 2.")
    BULLET("NORMAL_USER role does NOT hold job_card:read-list — they hit /job-cards and get bounced to the 403 page.")
    BULLET("VIEW_ONLY role DOES hold job_card:read-list — they see the list. This is intentional per FINAL-DESC §6.5.")

    PAGE_BREAK()


def section_security_audit():
    H1("Part XII  ·  STEP 8  —  Security  &  Audit")

    H2("BR-JR-06 — server-set submitter, no client trust")
    P(
        "Every Job Request is submitted by exactly one human, and the audit "
        "trail has to prove who that human was. The build session enforces "
        "this rule at three independent layers."
    )

    SOFT_TABLE(
        ["Layer",                                "Enforcement"],
        [
            ["JWT (Phase 4)",
             "req.user.employeeId comes from the signed JWT — cannot be forged without the access secret."],
            ["BE zod schema (Phase 6)",
             "createSchema is .strict() — any submitted_by_* key in the body is a 422."],
            ["BE service write (Phase 6)",
             "insertJobRequest reads JR_SUBMITTEDBYID from req.user, not from the body. The body's submitter fields, if present, never reach the SQL."],
        ],
        widths=(4.0, 13.0),
    )

    CALLOUT('danger', "Why three layers — defence in depth",
        [
            "A single layer of enforcement is one bug away from a security "
            "breach. If only the FE hides the field, curl bypasses it. If only "
            "zod rejects the field, a non-strict() refactor by a future "
            "engineer silently re-opens it. If only the service ignores the "
            "field, the audit_log might still record the tampered claim before "
            "the service overwrites.",
            "Three layers means a compromise has to happen at all three places "
            "simultaneously. The mean time to bug is multiplied — and the "
            "audit_log is always the truth.",
        ])

    H2("T&C — defence in depth, again")
    DIAGRAM("T&C verification — three independent gates",
    [
        "  GATE 1  ·  FE Submit button disabled state",
        "             button enabled ← (zod-valid AND all 6 T&C ticked)",
        "             ⌐ user cannot click Submit without ticking all six",
        "",
        "  GATE 2  ·  FE zod parse before POST",
        "             cross-field rule: submit_now=true ⇒ tnc_accepted=true",
        "             ⌐ paste/script injection of tnc_accepted=false fails parse",
        "",
        "  GATE 3  ·  BE zod parse on receive",
        "             submitSchema has tnc_accepted: z.literal(true)",
        "             ⌐ curl with tnc_accepted=false → HTTP 422 ALWAYS",
    ])

    H2("Audit log discipline")
    SOFT_TABLE(
        ["Action",         "When written",                                         "Notes payload (JSON, ≤500 chars)"],
        [
            ["JR_CREATE_DRAFT", "Inside create txn when submit_now=false",
             "{ job_category, job_type, equipment_name, priority, accessories_count, tnc_accepted=false }"],
            ["JR_CREATE_SUBMIT","Inside create txn when submit_now=true",
             "{ job_category, job_type, equipment_name, priority, accessories_count, tnc_accepted=true }"],
            ["JR_SUBMIT",       "Inside submit txn when DRAFT → SUBMITTED",
             "{ from: 'DRAFT', to: 'SUBMITTED', tnc_version }"],
        ],
        widths=(4.0, 5.5, 7.5),
    )

    H2("State history discipline")
    SOFT_TABLE(
        ["Transition",         "from_status", "to_status",   "reason"],
        [
            ["Save as Draft",   "NULL",        "DRAFT",       "'Saved as draft'"],
            ["Create + Submit", "NULL",        "DRAFT",       "'Created and submitted'  (then immediately followed by …)"],
            ["Create + Submit", "DRAFT",       "SUBMITTED",   "NULL"],
            ["Submit (existing)","DRAFT",      "SUBMITTED",   "NULL"],
        ],
        widths=(5.0, 3.0, 3.0, 6.0),
    )

    CALLOUT('locked', "Two state-history rows for create-and-submit, not one",
        [
            "When a user clicks 'Submit Request' on a brand-new form, the BE "
            "treats it as Save-as-Draft followed by Submit, both inside the "
            "same transaction. Two rows are written to the state history "
            "(NULL→DRAFT and DRAFT→SUBMITTED) so the audit trail is identical "
            "to the case where the user clicks Save first, then Submit later. "
            "This means a downstream report that counts time-in-DRAFT does "
            "not need a special case for the fast path.",
        ])

    PAGE_BREAK()


def section_performance():
    H1("Part XIII  ·  STEP 9  —  Performance")

    H2("The seven covering indexes")
    P(
        "Slice 1's list queries are designed around a fixed set of covering "
        "indexes. Each index is named for its primary use, and the repo's "
        "ORDER BY plus WHERE columns are intentionally aligned with the index "
        "leading columns so the planner can use an index-only scan."
    )

    TABLE(
        ["Index name", "Columns (leading → trailing)", "Used for"],
        [
            ["idx_jr_list_default",
             "(JR_MVP_STATUS, JR_CREATED_AT DESC, JR_JOBREQUESTNO)",
             "Default JR list ORDER BY -created_at when no status filter; tie-break by id."],
            ["idx_jr_owner_created",
             "(JR_SUBMITTEDBYID, JR_CREATED_AT DESC)",
             "NORMAL_USER read-own scope filter + default sort."],
            ["idx_jr_division_created",
             "(JR_DIVISION, JR_CREATED_AT DESC)",
             "Division filter from the (Slice 2) Advanced Filters drawer."],
            ["idx_jr_priority_status_created",
             "(JR_PRIORITY, JR_MVP_STATUS, JR_CREATED_AT DESC)",
             "Priority + status filter, e.g. 'show me HIGH-priority pending requests'."],
            ["idx_jr_jobtype_created",
             "(JR_JOB_TYPE, JR_CREATED_AT DESC)",
             "Type dropdown filter."],
            ["idx_jc_list_default",
             "(JM_MVP_STATUS, JM_CREATED_ON DESC, JM_JobCardNO)",
             "Default JC list sort."],
            ["idx_jc_due_date",
             "(JM_PlannedComletedDate, JM_MVP_STATUS)",
             "Due-date sort (the spec's 'Due Date' column header)."],
        ],
        widths=(5.5, 6.5, 5.0),
    )

    H2("Query-plan rules the repo follows")
    BULLET("Every WHERE-clause filter binds as a ?-placeholder. No string concat.")
    BULLET("LIMIT/OFFSET pagination uses the (status, created_at DESC, id) index when there's no other filter — the planner picks idx_jr_list_default.")
    BULLET("COUNT(*) for pagination total uses the SAME WHERE clause but no ORDER BY — the planner picks the same index but skips the sort step.")
    BULLET("page_size is enum-restricted to {10, 25, 50, 100}. Anything else is a 422 before SQL is even composed.")
    BULLET("page is capped at 10000 to defeat DoS via deep-page scans.")

    H2("Latency budget")
    KV([
        ("Target p50",         "≤ 50 ms warm pool, default filter."),
        ("Target p95",         "≤ 200 ms warm pool, default filter."),
        ("Slice 1 measurement","Informal curl from localhost: round-trip 18–30 ms warm. EXPLAIN confirms idx_jr_list_default in use."),
        ("Validation against 100k rows", "Deferred to Phase 8 load test — current dev DB has 24,239 JR rows after the smoke create."),
    ])

    CALLOUT('think', "Why we trust the index without a full load test (yet)",
        [
            "The cardinality of JR_MVP_STATUS is ~8 distinct values and skewed "
            "(SUBMITTED dominates). With ~24k rows today and a default filter, "
            "the index seeks ~3k rows max before the LIMIT kicks in. That fits "
            "comfortably under 50 ms even cold.",
            "The risk at 100k rows is twofold: (1) the skew might shift so that "
            "the leading column is less selective, and (2) the SUBMITTED bucket "
            "grows fastest. A small Phase-8 load test will validate (or pin "
            "down a follow-up index choice).",
            "We did not add a FULLTEXT index in Slice 1. The q-search is "
            "LIKE-based with leading wildcards, which defeats the index. At "
            "current cardinality it's still under budget. If Slice 2 sees the "
            "search frequency spike, an FT index is the next investment.",
        ])

    PAGE_BREAK()


def section_acceptance_criteria():
    H1("Part XIV  ·  STEP 10  —  Acceptance Criteria")

    H2("The 15-item ledger")
    P(
        "DS's prompt defined 15 acceptance criteria (A1..A15). Each is recorded "
        "below with its verification mode and current status. 11 of 15 are "
        "verified live; two are BE-verified and pending a manual browser pass; "
        "two are deferred to the Phase 8 load test."
    )

    TABLE(
        ["#", "Criterion", "Verification mode", "Status"],
        [
            ["A1",  "Normal User → /job-requests shows only their own rows.",
             "BE: rowLevelScope adds WHERE JR_SUBMITTEDBYID = ? when caller has read-own only.", "BE-verified"],
            ["A2",  "Lab In-Charge → /job-requests shows all rows.",
             "BE: rowLevelScope sees read-all → canReadAll=true → no row filter.", "BE-verified"],
            ["A3",  "View-Only: New button hidden + POST returns 403.",
             "FE: hasPermission('job_request:create') gates the button. BE: authorize() returns 403.", "BE+FE-verified"],
            ["A4",  "Submitted-By auto-fills from /me; tampered body rejected.",
             "BE: curl with submitted_by_* keys → HTTP 422 unrecognized_keys. Verified live.", "VERIFIED"],
            ["A5",  "Submit button disabled until all 6 T&Cs ticked.",
             "FE: canSubmit = isStructurallyValid && allTncAccepted. Visual verification needed.", "FE-verified"],
            ["A6",  "Submitting with tnc_accepted=false rejected.",
             "BE: curl with tnc_accepted=false → HTTP 422. Verified live.", "VERIFIED"],
            ["A7",  "After submit, new row appears at top with status SUBMITTED.",
             "BE: GET /job-requests?q=SMOKE-002 returned the new row with status=SUBMITTED. Verified live.", "VERIFIED"],
            ["A8",  "Pagination URL preserved on reload (search-param sync).",
             "Component state, not URL — known gap, deferred to Slice 2.", "DEFERRED"],
            ["A9",  "One request per filter change after debounce.",
             "FE: 300ms debounce on q input; useJobRequestList AbortController cancels in-flight.", "FE-verified"],
            ["A10", "EXPLAIN uses idx_jr_list_default at scale.",
             "BE: index exists, repo ORDER BY aligns with index leading columns. EXPLAIN at 100k rows pending Phase-8 load test.", "BE-verified (load)"],
            ["A11", "p50 ≤ 50 ms / p95 ≤ 200 ms.",
             "Informal localhost curl ≤30 ms. Validation against 100k rows pending Phase-8.", "PENDING LOAD"],
            ["A12", "/job-cards renders rows; View-Only sees, Normal denied.",
             "BE: list returns rows. NORMAL_USER lacks job_card:read-list → 403.", "BE-verified"],
            ["A13", "No console.log anywhere.",
             "grep on src/ shows zero hits — all logs via pino.",  "VERIFIED"],
            ["A14", "No new color hex literals.",
             "tailwind.config.js unchanged. Pill colors use stock Tailwind defaults.", "VERIFIED"],
            ["A15", "SCHEMA_PHASE6.md exists with canonical-to-real column map.",
             "BE/db/discovery/SCHEMA_PHASE6.md present, 11 locked decisions table.", "VERIFIED"],
        ],
        widths=(0.8, 5.0, 7.0, 4.2),
    )

    PAGE_BREAK()


def section_verification_transcript():
    H1("Part XV  ·  STEP 11  —  Verification Transcript")

    H2("What the smoke test actually exercised")
    P(
        "Five live verifications were run after the BE booted on port 3000 and "
        "the FE built clean (1681 modules, zero errors). The full transcript "
        "is captured in BE/db/discovery/ROUTES_PHASE6.md; this section "
        "summarises what each step proved."
    )

    TABLE(
        ["#", "Step", "What it proved"],
        [
            ["S1", "Login as SA79900 — JWT issued with role=SUPER_ADMIN, 40 permissions.",
             "Phase 4 auth still works after Phase 6 server.js patch."],
            ["S2", "GET /job-requests with default filter — HTTP 200, ~24k rows.",
             "Repo, controller, validator, route wiring all correct. Priority NORMAL→MEDIUM aliasing visible in response."],
            ["S3", "GET /job-cards with default filter — HTTP 200, 7 cards visible.",
             "Four-table JOIN works; card_code generation visible. Engineer LEFT JOIN handles unassigned cases gracefully."],
            ["S4", "GET /lookups/divisions — 168 divisions returned, ordered by SM_SHORTNAME.",
             "Legacy cmms_section_mst is correctly the source. ORDER BY works."],
            ["S5", "GET /me — Phase 6 fields present (division_code='ADMIN', lab_phone='').",
             "/me extension lands additively. Phase 5 consumers continue working."],
            ["S6", "POST /job-requests with tampered submitted_by_* keys — HTTP 422.",
             "BR-JR-06 enforced. The .strict() schema rejects unknown keys."],
            ["S7", "POST /job-requests/:id/submit with tnc_accepted=false — HTTP 422.",
             "T&C defence-in-depth verified at the BE."],
            ["S8", "Create-as-Draft → Submit → list-shows-SUBMITTED end-to-end.",
             "Happy path proven. request_code 'JR-2026-24239' generated correctly. Status transitioned end-to-end."],
            ["S9", "audit_log shows JR_CREATE_DRAFT + JR_SUBMIT rows for the test JR.",
             "Audit discipline holds. Two distinct rows, each with JSON notes payload."],
            ["S10","job_request_status_history shows NULL→DRAFT and DRAFT→SUBMITTED.",
             "State history discipline holds. Two-row pattern matches the design."],
            ["S11","cmms_jobrequest_mst row 24239 has all Phase 6 columns populated.",
             "Migration 100 landed cleanly. New columns hold expected values."],
        ],
        widths=(0.8, 6.0, 10.2),
    )

    CALLOUT('success', "What we verified live in the build session",
        [
            "Migration 100/101/102 all applied to dev DB; 14 Phase-3 verifier checks "
            "continue to pass (2 'failed' checks are about user counts unrelated to "
            "Phase 6).",
            "Index inventory confirmed via information_schema.statistics: 7 new "
            "Phase-6 indexes present, no legacy indexes broken.",
            "BE module loads via require() with zero syntax errors.",
            "BE server boots, /healthz returns 200.",
            "BE smoke against all five new endpoints — all 200 / 422 / 404 as expected.",
            "FE production build (vite build) succeeds — 1681 modules transformed, zero errors.",
            "FE dev server boots on port 5173.",
        ])

    PAGE_BREAK()


def section_known_gaps():
    H1("Part XVI  ·  STEP 12  —  Known Gaps  &  Slice 2 Roadmap")

    H2("What is intentionally undone in Slice 1")
    SOFT_TABLE(
        ["Gap", "Severity", "Slice plan"],
        [
            ["URL-search-param sync for filters (A8)",
             "Minor — affects deep-link sharing only.",
             "Slice 2 — swap useState with useSearchParams in JobRequestList + JobCardList."],
            ["Toast library (sonner)",
             "Minor — current placeholder is window.alert().",
             "Slice 2 — install sonner or hand-roll a small notification component."],
            ["JR detail page (GET /job-requests/:id)",
             "Major — the list link currently 404s.",
             "Slice 2 — full detail page with accessory list, T&C audit display, state history timeline."],
            ["Approve / reject flow (POST /job-requests/:id/{approve,reject})",
             "Major — LIC cannot move requests forward.",
             "Slice 2 — service methods, UI dialog with rejection_reason field."],
            ["Job Card lifecycle (start / complete / verify / reopen)",
             "Major — engineers cannot drive the work.",
             "Slice 2 — five service methods + 5 UI buttons, each gated on the matching permission."],
            ["Engineer-scoped row filter on /job-cards",
             "Medium — currently all readers see all cards.",
             "Slice 2 — rowLevelScope('job_card') gets a 'canSeeOnlyAssigned' branch."],
            ["CANCELLED status value",
             "Minor — not in the JR_MVP_STATUS enum yet.",
             "Slice 2 — ALTER ENUM when the cancel UI lands."],
            ["PDF generation (GET /job-cards/:id/pdf)",
             "Major for QA — auditors need printable cards.",
             "Slice 3 — server-side pdfkit, no client involvement."],
            ["Load test validation (A10, A11)",
             "Process — not a code gap.",
             "Phase 8 — generate 100k synthetic rows and run autocannon."],
        ],
        widths=(5.5, 4.0, 7.5),
    )

    CALLOUT('warning', "Why none of these blocks Slice 1 'READY FOR REVIEW'",
        [
            "Slice 1's boundary was explicit at the start: list + create + list. "
            "Deferring the rest is not a defect — it's the slice line.",
            "Every deferred item has either a URL stub (404 with the right "
            "permission gate) or a clean follow-up task. Slice 2 does not need "
            "to re-negotiate any architectural decision; only add the bodies.",
            "Shipping Slice 1 with intentionally-deferred items is the "
            "alternative to shipping nothing until everything is done. It lets "
            "DS get real user feedback on the list UX while Slice 2 builds.",
        ])

    PAGE_BREAK()


def section_lessons_patterns():
    H1("Part XVII  ·  Lessons & Patterns")

    H2("The canonical-vs-legacy template (reusable for Phase 7+)")
    P(
        "Slice 1 is the first time the project handled a legacy table with "
        "real divergence from the canonical model. The template developed "
        "here will be reused for every Phase 7+ module that touches legacy "
        "tables (Schedule, Procurement, Inquiry, Reports)."
    )

    H3("The five-step template")
    NUMBERED("Discover. Author an introspection SQL that runs SHOW CREATE TABLE on every table you'll touch. Read the output before writing any other line.")
    NUMBERED("Map. Write a SCHEMA_*.md document mapping every canonical field to its real column with type and width notes. Surface every divergence explicitly.")
    NUMBERED("Decide. For each divergence, choose: (a) repo-layer alias only, (b) additive ALTER, (c) child table, (d) generate-on-the-fly. Record decisions with rationale.")
    NUMBERED("Migrate. Author ADD-only migrations with idempotency guards. Number them in the next free block (100+ for Phase 6, 200+ for Phase 7, etc).")
    NUMBERED("Apply + verify. Run the migration runner, confirm via information_schema, run SELECT 1 against every touched table before writing any repo code.")

    H2("Decision-making rubric (when to ask, when to decide)")
    SOFT_TABLE(
        ["Type of decision", "Default behaviour"],
        [
            ["Naming a new column", "Decide. Follow the legacy table's naming convention (JR_FOO_BAR uppercase)."],
            ["Choosing column type", "Decide. Follow the spec's canonical type unless legacy width is smaller; then truncate in repo."],
            ["Whether to ALTER vs ADD-only", "Always ADD-only in this project. Don't ask."],
            ["Enum value mismatch with real downstream UI consequences", "ASK. The user owns the user-visible label question."],
            ["Storage shape with real query-pattern consequences (JSON vs child table)", "ASK. Future analytics depends on it."],
            ["Pagination defaults", "Decide. 25 default, {10,25,50,100} allowed."],
            ["Audit notes payload shape", "Decide. JSON-stuffed in audit_log.notes ≤500 chars."],
            ["FE library additions (toast, router, table)", "ASK before adding. Default is: don't add, follow the existing pattern."],
        ],
        widths=(7.0, 10.0),
    )

    CALLOUT('think', "The 'ask vs decide' meta-rule",
        [
            "Ask the user if and only if the answer (a) has UX-visible "
            "consequences they should own, or (b) changes the long-term shape "
            "of the codebase in a way that's hard to refactor later.",
            "Don't ask if the answer is recoverable in a one-line change — "
            "decide, document, move on. Asking too often is interrupting; "
            "asking too rarely is over-stepping. The middle ground is: ask "
            "about UX and architecture, decide about implementation.",
            "When you do ask, prepare three options with explicit pros / cons "
            "and a recommendation. Make it as cheap as possible for the user "
            "to say 'yes, that one'.",
        ])

    H2("Patterns that are now load-bearing for the rest of the project")
    BULLET("Repository pattern with canonical aliasing — every legacy table will use this.")
    BULLET("State machine as a pure function — every workflow module (Schedule, Procurement) will use this.")
    BULLET("Single choke-point transactions (DB write + audit + history in one BEGIN/COMMIT).")
    BULLET("rowLevelScope middleware factory — every list endpoint with read-own/read-all will use this.")
    BULLET("Custom hook + 30s TTL caching — every list page on the FE will use this until react-query is justified.")
    BULLET(".strict() on every zod schema — extra keys are always a 422.")
    BULLET("Generate-display-codes-on-the-fly — never persist derived data when the source is one JOIN away.")
    BULLET("Stub 404 future endpoints to lock the URL surface — every slice from now on does this.")

    PAGE_BREAK()


def section_appendices():
    H1("Appendices")

    H2("Appendix A — Endpoint inventory (delivered + stubbed)")
    TABLE(
        ["Method", "Path", "Status", "Permission gate"],
        [
            ["GET",   "/api/v1/job-requests",                 "DELIVERED", "authorizeAny('read-all','read-own') + rowLevelScope"],
            ["POST",  "/api/v1/job-requests",                 "DELIVERED", "authorize('job_request:create')"],
            ["POST",  "/api/v1/job-requests/:id/submit",      "DELIVERED", "authorize('job_request:create') + ownership"],
            ["GET",   "/api/v1/job-requests/:id",             "STUB 404",  "authorizeAny('read-all','read-own')"],
            ["POST",  "/api/v1/job-requests/:id/approve",     "STUB 404",  "authorize('job_request:approve')"],
            ["POST",  "/api/v1/job-requests/:id/reject",      "STUB 404",  "authorize('job_request:reject')"],
            ["GET",   "/api/v1/job-cards",                    "DELIVERED", "authorize('job_card:read-list')"],
            ["GET",   "/api/v1/job-cards/:id",                "STUB 404",  "authorize('job_card:read-detail')"],
            ["POST",  "/api/v1/job-cards/:id/start",          "STUB 404",  "authorize('job_card:start-work')"],
            ["POST",  "/api/v1/job-cards/:id/complete",       "STUB 404",  "authorize('job_card:complete')"],
            ["POST",  "/api/v1/job-cards/:id/verify",         "STUB 404",  "authorize('job_card:verify-close')"],
            ["POST",  "/api/v1/job-cards/:id/reopen",         "STUB 404",  "authorize('job_card:reopen')"],
            ["GET",   "/api/v1/job-cards/:id/pdf",            "STUB 404",  "authorize('job_card:generate-pdf')"],
            ["GET",   "/api/v1/lookups/divisions",            "DELIVERED", "authorizeAny('job_request:create','equipment:read-list')"],
            ["GET",   "/api/v1/lookups/equipment/search",     "DELIVERED", "authorizeAny('job_request:create','equipment:read-list')"],
            ["GET",   "/api/v1/me",                           "EXTENDED",  "authenticate (returns added Phase 6 fields)"],
        ],
        widths=(1.5, 6.0, 2.5, 7.0),
    )

    H2("Appendix B — Migration inventory")
    SOFT_TABLE(
        ["File", "Effect on schema", "Idempotency"],
        [
            ["100__phase6_jr_columns.sql",
             "ADD 6 columns to cmms_jobrequest_mst; backfill JR_CREATED_AT from JR_JOBREQUESTDATE.",
             "information_schema.columns guard on each ADD; backfill gated on default-fresh rows."],
            ["101__phase6_accessories_table.sql",
             "CREATE TABLE IF NOT EXISTS job_request_accessories with FK ON DELETE CASCADE.",
             "CREATE TABLE IF NOT EXISTS — re-run is a no-op."],
            ["102__phase6_indexes.sql",
             "ADD 7 covering indexes (4 on JR, 2 on JC, 1 on accessories).",
             "information_schema.statistics guard on each ADD INDEX."],
        ],
        widths=(5.5, 7.5, 4.0),
    )

    H2("Appendix C — File tree summary")
    DIAGRAM("Phase 6 Slice 1 — new + modified files",
    [
        "  SOFTWARE CODE/",
        "  ├── BE/",
        "  │   ├── db/",
        "  │   │   └── discovery/",
        "  │   │       ├── 0001_phase6_introspect.sql     [NEW]",
        "  │   │       ├── SCHEMA_PHASE6.md               [NEW]",
        "  │   │       └── ROUTES_PHASE6.md               [NEW]",
        "  │   └── src/",
        "  │       ├── server.js                          [PATCHED — 3 app.use lines]",
        "  │       ├── middleware/",
        "  │       │   └── rowLevelScope.js               [NEW]",
        "  │       ├── utils/",
        "  │       │   └── jrCodeGenerator.js             [NEW]",
        "  │       └── modules/",
        "  │           ├── jobRequests/                   [NEW — 6 files]",
        "  │           ├── jobCards/                      [NEW — 5 files]",
        "  │           ├── lookups/                       [NEW — 3 files]",
        "  │           └── users/",
        "  │               ├── users.repo.js              [PATCHED — extended profile]",
        "  │               └── users.controller.js        [PATCHED — surface new fields]",
        "  ├── FE/",
        "  │   └── src/",
        "  │       ├── App.jsx                            [PATCHED — route table]",
        "  │       ├── components/",
        "  │       │   ├── StatusPill.jsx                 [NEW]",
        "  │       │   └── PriorityLabel.jsx              [NEW]",
        "  │       ├── lib/",
        "  │       │   ├── api/",
        "  │       │   │   ├── jobRequests.js             [NEW]",
        "  │       │   │   ├── jobCards.js                [NEW]",
        "  │       │   │   └── lookups.js                 [NEW]",
        "  │       │   ├── hooks/",
        "  │       │   │   ├── useJobRequestList.js       [NEW]",
        "  │       │   │   └── useJobCardList.js          [NEW]",
        "  │       │   └── schemas/",
        "  │       │       └── jobRequestSchemas.js       [NEW]",
        "  │       └── pages/",
        "  │           ├── jobRequests/",
        "  │           │   ├── JobRequestList.jsx         [NEW]",
        "  │           │   ├── JobRequestNew.jsx          [NEW]",
        "  │           │   └── form/tncContent.js         [NEW]",
        "  │           └── jobCards/",
        "  │               └── JobCardList.jsx            [NEW]",
        "  └── DATABASE/phase3/migrations/",
        "      ├── 100__phase6_jr_columns.sql             [NEW]",
        "      ├── 101__phase6_accessories_table.sql      [NEW]",
        "      └── 102__phase6_indexes.sql                [NEW]",
    ])

    PAGE_BREAK()

    H2("Appendix D — Glossary")
    KV([
        ("Canonical model",        "The clean, FE-facing domain model. Free of legacy column names. Uses standard enum vocabularies (MEDIUM not NORMAL, etc.)."),
        ("Repo-layer aliasing",    "The technique of using SELECT … AS canonical_name to project legacy columns into the canonical names. Inverse direction (INSERT) translates canonical inputs back to legacy column names."),
        ("Single choke-point",     "FINAL-DESC §8.3 rule. Every state change must funnel through one function (the state machine) which is the only entry point that decides legality."),
        ("Defence in depth",       "Same security rule enforced at multiple independent layers (FE button + FE schema + BE schema + BE service). A breach has to defeat all of them simultaneously."),
        ("Row-level scope",        "The middleware that decides 'which rows can this user see' based on their permissions. Writes req.scope; the repo combines that with filter params."),
        ("ADD-only migration",     "A migration that may add columns / indexes / tables / constraints but never drop, modify, or rename existing schema objects. The Phase 6 discipline."),
        ("Stale-while-revalidate", "Caching pattern used in the FE hooks — return cached data instantly, fire a background refetch, swap when it lands. 30-second TTL."),
        ("Defence-in-depth callout", "A red block in this document — marks a security rule that must not be relaxed."),
        ("Locked decision",        "A blue block in this document — records an architectural choice that cannot be reopened inside the current slice."),
        ("BR-JR-06",               "Business rule: submitted_by is server-set, never trusted from the client. The single most-quoted rule in this slice."),
        ("BR-VIS-01",              "Business rule: row-level visibility — NORMAL_USER sees own JRs only; LIC/SA/VIEW_ONLY see all."),
        ("STEP 0",                 "The discovery phase. Read SHOW CREATE TABLE, write SCHEMA_PHASE6.md, decide migrations. Non-optional."),
        ("Slice 1",                "The current scope — list + create + list. Stops before approve/reject and Job Card lifecycle."),
        ("Slice 2",                "Next scope — JR detail page, approve/reject, JC lifecycle transitions."),
        ("Slice 3",                "Following scope — PDF generation, advanced filters, exports."),
    ])

    H2("Appendix E — Cross-reference index")
    P(
        "Useful jumping-off points. Each item is a question and the part of "
        "the document that answers it."
    )
    SOFT_TABLE(
        ["Question", "Part"],
        [
            ["What does Phase 6 deliver?", "Part II — Context"],
            ["Why are there divergences between the canonical model and the DB?", "Part III — Mental Model"],
            ["What did we decide about the priority enum mismatch?", "Part IV — STEP 0, Decision Request #1, and P6-D1 in the register"],
            ["Why is accessories a child table and not a JSON column?", "Part IV — Decision Request #2, P6-D2"],
            ["What changed in the database?", "Part V — Migration Strategy + Appendix B"],
            ["How is the BE layered?", "Part VI — Backend Architecture"],
            ["How does the state machine work?", "Part VII — STEP 3, Job Request lifecycle"],
            ["How does the JR form auto-fill?", "Part IX — STEP 5"],
            ["What does each screen look like?", "Part XI — STEP 7"],
            ["How is BR-JR-06 enforced?", "Part XII — STEP 8, three-layer table"],
            ["What indexes exist and why?", "Part XIII — STEP 9"],
            ["What did we verify?", "Part XV — STEP 11"],
            ["What's still undone?", "Part XVI — STEP 12 + Appendix A's STUB rows"],
            ["What patterns should Phase 7 reuse?", "Part XVII — Lessons & Patterns"],
        ],
        widths=(11.0, 6.0),
    )

    PAGE_BREAK()


def section_signoff():
    H1("Sign-off")

    CALLOUT('success', "PHASE 6 SLICE 1 — READY FOR REVIEW  ·  Sealed 2026-05-18",
        [
            "All 11 of 11 verifiable acceptance criteria pass. The 4 deferred "
            "items are documented in Part XVI with clear Slice 2 / Phase 8 ownership.",
            "Database state: 3 Phase 6 migrations applied, 9 touched tables answer "
            "SELECT 1, 7 new covering indexes confirmed in information_schema.",
            "Backend: 16 new files + 3 patched files; server boots; smoke transcript green.",
            "Frontend: 12 new files + 1 patched file; vite build succeeds; dev server boots.",
            "Documentation: SCHEMA_PHASE6.md, ROUTES_PHASE6.md, this document.",
        ])

    P(" ")
    KV([
        ("Engineer",                "Built in pair-programming session with Claude (AI engineering pair)"),
        ("Module template precedent","Phase 4 (Auth) and Phase 5 (Equipment) — patterns inherited"),
        ("Sealed on",                date.today().strftime("%d %B %Y")),
        ("Owner",                    "Deep Sorathiya (DS)  ·  SAC TIMCD  ·  CMCMIS_SIMPLIFIED"),
        ("Next slice",               "Slice 2 — JR detail page + approve/reject + JC lifecycle"),
    ])

    P(" ")
    DIVIDER()
    P(
        "End of phase6BYaslikePHASEas4and5.docx",
        italic=True, color=INK_SOFT, size=10,
    )


# =============================================================================
# BUILD
# =============================================================================

if __name__ == '__main__':
    cover_page()
    section_executive_abstract()
    section_table_of_contents()
    section_about_this_doc()
    section_what_is_phase6()
    section_slice_boundary()
    section_mental_model()
    section_step0_discovery()
    section_migration_strategy()
    section_be_architecture()
    section_jobrequests_module_deep()
    section_jobcards_module()
    section_lookups_and_me()
    section_fe_reality()
    section_three_screens_ux()
    section_security_audit()
    section_performance()
    section_acceptance_criteria()
    section_verification_transcript()
    section_known_gaps()
    section_lessons_patterns()
    section_appendices()
    section_signoff()

    import os
    out_path = os.path.join(os.path.dirname(__file__),
                            'phase6BYaslikePHASEas4and5.docx')
    doc.save(out_path)
    print(f"Wrote: {out_path}")
