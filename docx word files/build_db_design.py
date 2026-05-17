"""
FINAL-DB-DESIGN-v1.docx generator.
Phase 3 database design output. Audits 64 legacy tables, classifies them,
designs new cmcm_* tables, and lays out the migration plan.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x0B, 0x2A, 0x5B)
ACCENT = RGBColor(0x1E, 0x6F, 0xA8)
GREEN = RGBColor(0x1B, 0x7F, 0x3A)
RED = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x55, 0x55, 0x55)
ORANGE = RGBColor(0xC8, 0x6E, 0x10)
PURPLE = RGBColor(0x70, 0x3F, 0xA8)
TABLE_HEADER_FILL = "1E6FA8"
SECTION_FILL = "0B2A5B"
LOCKED_FILL = "1B7F3A"
WARN_FILL = "B31B1B"


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size is None:
        size = {1: 22, 2: 17, 3: 14, 4: 12}.get(level, 12)
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_callout(doc, text, fill="EAF3FB", text_color=None):
    if text_color is None:
        text_color = NAVY
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = text_color
    run.bold = True
    return tbl


def add_warning_banner(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, WARN_FILL)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run("WARNING  -  " + text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


def add_locked_banner(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, LOCKED_FILL)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run("LOCKED  -  " + text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


def add_mono_block(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F4F4F4")
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return tbl


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, TABLE_HEADER_FILL)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = tbl.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            v = str(val).strip()
            if v in ("YES", "Y", "KEEP", "LOCKED"):
                run.font.color.rgb = GREEN
                run.bold = True
            elif v in ("NO", "N", "DEPRECATE", "REPLACE"):
                run.font.color.rgb = RED
                run.bold = True
            elif v in ("EXTEND", "REFACTOR", "BRIDGE"):
                run.font.color.rgb = ACCENT
                run.bold = True
            elif v.startswith("DECISION") or v.startswith("BR-VIOLATION"):
                run.font.color.rgb = ORANGE
                run.bold = True
            elif "ORPHAN" in v:
                run.font.color.rgb = PURPLE
                run.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)


def add_page_break(doc):
    doc.add_page_break()


def add_section_banner(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, SECTION_FILL)
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============ COVER ============
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\nFINAL-DB-DESIGN")
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CMCMIS - Phase 3 - Database Design v1\n")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Audit + Classification + New Schema + Migration Plan")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n64 legacy tables (cmcmis_redev) classified.\n"
                  "8 new auth/audit tables designed.\n"
                  "Equipment + Job lifecycle refactor planned.\n"
                  "Zero-drop migration strategy locked.\n")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = GREY

    add_locked_banner(doc,
        "Super Admin seed IDs: SA79900 + AC77777   -   Dated 2026-05-16")

    add_page_break(doc)

    # ============ TOC ============
    add_heading(doc, "Table of Contents", level=1)
    toc = [
        "Section 1  - Inputs Received + Audit Methodology",
        "Section 2  - The 7 Loudest Flags (Read First)",
        "Section 3  - Hot Audit of All 64 Legacy Tables",
        "Section 4  - Cluster 1: Identity & Access (DEEP DESIGN)",
        "Section 5  - Cluster 10: Audit & Logs (DEEP DESIGN)",
        "Section 6  - Cluster 12: Lookups (DEEP DESIGN)",
        "Section 7  - Cluster 3: Equipment Master (REFACTOR)",
        "Section 8  - Cluster 4: Job Lifecycle (REFACTOR)",
        "Section 9  - Phase 2 Clusters (sketches)",
        "Section 10 - Zero-Drop Migration Strategy",
        "Section 11 - 8 Questions Awaiting Your Answers",
        "Section 12 - Updated Naming + Numbering Conventions",
        "Section 13 - Next Steps",
    ]
    for line in toc:
        add_para(doc, "  " + line, size=11)

    add_page_break(doc)

    # =====================================================================
    # SECTION 1
    # =====================================================================
    add_section_banner(doc, "Section 1 - Inputs Received + Audit Methodology")

    add_heading(doc, "1.1 Inputs", level=2)
    add_table(doc, ["Input", "Status"], [
        ["Existing DB schema (cmcmis_schema_only.sql)", "Received"],
        ["Table summary (cmcmis_table_summary.csv)", "Received"],
        ["Key report (cmcmis_schema_key_report.md)", "Received"],
        ["Super Admin seed IDs", "SA79900 + AC77777"],
        ["Table classification by user", "Not provided - derived in this doc"],
    ])

    add_heading(doc, "1.2 Method", level=2)
    add_bullets(doc, [
        "Audit each of the 64 tables individually.",
        "Apply 5-flag system: STRAIGHTFORWARD / DECISION-NEEDED / BR-VIOLATION / ORPHAN / REVISION.",
        "Assign verdict: KEEP / EXTEND / REFACTOR / BRIDGE / DEPRECATE / REPLACE.",
        "Group by 13 entity clusters per the cluster framework.",
        "Prioritise MVP-critical clusters (1, 3, 4, 10, 12).",
        "Design NEW cmcm_* tables alongside legacy cmms_* tables.",
        "Plan zero-drop migration: never DROP in MVP, only quarantine.",
    ])

    add_heading(doc, "1.3 Flag legend", level=2)
    add_table(doc, ["Flag", "Meaning"], [
        ["STRAIGHTFORWARD", "Clean keep / modify, no debate"],
        ["DECISION-NEEDED", "Genuine fork, needs user input"],
        ["BR-VIOLATION", "Contradicts a locked business rule (must fix)"],
        ["ORPHAN-CANDIDATE", "Likely dead but need confirmation"],
        ["REVISION-CANDIDATE", "Schema bug / forces v1.1 to FINAL-DESC"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 2 - LOUDEST FLAGS
    # =====================================================================
    add_section_banner(doc, "Section 2 - The 7 Loudest Flags (Read First)")

    add_warning_banner(doc,
        "These 7 issues must be addressed before any MVP code ships. "
        "Each maps to a locked business rule or NFR.")

    add_table(doc, ["#", "Flag", "Issue", "Fix"], [
        ["1", "BR-VIOLATION",
         "cmms_userrole_mst.USER_PASSWORD VARCHAR(10) - plaintext, cannot fit bcrypt(60)",
         "New cmcm_users.password_hash VARCHAR(60). Force first-login reset for all 565 users."],
        ["2", "BR-VIOLATION",
         "cmms_jobcard_status_hist has NO primary key - violates BR-JC-08 (append-only immutable)",
         "ALTER TABLE add id BIGINT UNSIGNED AUTO_INCREMENT PK."],
        ["3", "BR-VIOLATION",
         "cmms_eqip_mst missing registered_by / verified_by / verified_at (BR-EQP-09)",
         "ADD 3 new columns; backfill registered_by = EQM_CREATED_BY."],
        ["4", "BR-VIOLATION",
         "cmms_eqip_mst status is free-text EQM_DIV_STATUS - violates BR-EQP-10 state machine",
         "ADD status_code FK to cmcm_lookup_values(equipment_status)."],
        ["5", "BR-VIOLATION",
         "No central audit_log; only per-entity _hist tables - violates BR-AUD-01",
         "Build cmcm_audit_log with entity_type, entity_id, before/after JSON."],
        ["6", "DECISION-NEEDED",
         "cmms_role_mst has 23 roles - we have 5",
         "Need user mapping or my proposed mapping confirmation."],
        ["7", "DECISION-NEEDED",
         "cmms_cont_mst referenced by 4+ FKs but NOT in 64-table dump",
         "Confirm: missing from dump OR FK constraints broken in prod?"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 3 - HOT AUDIT (ALL 64)
    # =====================================================================
    add_section_banner(doc, "Section 3 - Hot Audit of All 64 Legacy Tables")

    add_heading(doc, "3.1 Cluster 1 - Identity & Access", level=2)
    add_table(doc,
        ["#", "Table", "Rows", "Cols", "Flag", "Verdict", "Why"],
        [
            ["1", "cmms_emp_mst", "57", "30", "STRAIGHTFORWARD", "KEEP+BRIDGE", "Employee directory; FK target for cmcm_users"],
            ["2", "cmms_role_mst", "23", "7", "DECISION-NEEDED", "REPLACE", "Trim 23 -> 5 roles; mapping pending"],
            ["3", "cmms_userrole_mst", "565", "9", "BR-VIOLATION", "REPLACE", "Plaintext VARCHAR(10) password"],
            ["4", "cmms_section_user_mst", "294", "4", "STRAIGHTFORWARD", "KEEP", "Division-user mapping"],
            ["5", "cmms_accessright_mst", "3221", "11", "BR-VIOLATION", "REPLACE", "Module CRUD bits != resource:action"],
            ["6", "cmms_module_mst", "163", "13", "DECISION-NEEDED", "KEEP-READ", "Menu metadata; permissions drive visibility"],
            ["7", "cmms_designation_mst", "40", "8", "STRAIGHTFORWARD", "KEEP", "Designation lookup"],
        ])

    add_heading(doc, "3.2 Cluster 3 - Equipment Master", level=2)
    add_table(doc,
        ["#", "Table", "Rows", "Cols", "Flag", "Verdict", "Why"],
        [
            ["8", "cmms_eqip_mst", "5704", "47", "BR-VIOLATION + REVISION", "EXTEND", "Missing registered_by, verified_by, status enum"],
            ["9", "cmms_eqip_mst_hist", "519", "49", "STRAIGHTFORWARD", "KEEP", "Equipment change snapshots"],
            ["10", "cmms_eqipinst_identification", "2286", "10", "STRAIGHTFORWARD", "KEEP", "Sub-instruments"],
            ["11", "cmms_ins_accuracy_info", "1501", "5", "STRAIGHTFORWARD", "KEEP", "Accuracy specs"],
            ["12", "cmms_eqip_detail_spec", "3", "4", "ORPHAN + REVISION", "DEPRECATE", "longblob storage; only 3 rows"],
            ["13", "cmms_division_hist", "3676", "7", "STRAIGHTFORWARD", "KEEP", "Division-transfer audit"],
            ["14", "cmms_product_mst", "32", "11", "STRAIGHTFORWARD", "KEEP", "T&ME/F&PE classifier"],
        ])

    add_heading(doc, "3.3 Cluster 4 - Job Lifecycle", level=2)
    add_table(doc,
        ["#", "Table", "Rows", "Cols", "Flag", "Verdict", "Why"],
        [
            ["15", "cmms_jobrequest_mst", "21485", "37", "STRAIGHTFORWARD", "EXTEND", "Add status_code, priority_code, submitted_by_user_id"],
            ["16", "cmms_jobrequest_item_dtl", "7786", "11", "STRAIGHTFORWARD", "KEEP", "Items per request"],
            ["17", "cmms_jobrequest_project_dtl", "19624", "2", "STRAIGHTFORWARD", "KEEP", "M2M to projects"],
            ["18", "cmms_jobcard_mst", "19432", "34", "DECISION-NEEDED", "EXTEND", "Two candidate keys; col-type mismatch"],
            ["19", "cmms_jobcard_mst_history", "22143", "37", "STRAIGHTFORWARD", "KEEP", "Snapshots"],
            ["20", "cmms_jobcard_status_hist", "22214", "5", "BR-VIOLATION", "REFACTOR", "No PK -> add id BIGINT PK"],
            ["21", "cmms_jobcard_request_info", "19432", "27", "STRAIGHTFORWARD", "KEEP", "1:1 jobcard request snapshot"],
            ["22", "cmms_jobcard_request_item_dtl", "11064", "10", "STRAIGHTFORWARD", "KEEP", "Items captured per card"],
            ["23", "cmms_jobcard_request_project_dtl", "22316", "2", "STRAIGHTFORWARD", "KEEP", "M2M projects"],
            ["24", "cmms_jobcard_attendedby_dtl", "27890", "4", "STRAIGHTFORWARD", "KEEP", "Assignees per card"],
            ["25", "cmms_jobcard_eq_used", "38316", "4", "STRAIGHTFORWARD", "KEEP", "Standards used in job"],
        ])

    add_heading(doc, "3.4 Cluster 5 - Calibration (Phase 2 sketch)", level=2)
    add_table(doc,
        ["#", "Table", "Rows", "Cols", "Flag", "Verdict"],
        [
            ["26", "cmms_jobcard_cal_dtl", "9065", "21", "STRAIGHTFORWARD", "KEEP"],
            ["27", "cmms_jobcard_cal_adjustments_dtl", "1831", "6", "STRAIGHTFORWARD", "KEEP"],
            ["28", "cmms_jobcard_cal_observations", "77171", "4", "STRAIGHTFORWARD", "KEEP (biggest table)"],
            ["29", "cmms_cal_jobcard_feedback_spec", "0", "4", "ORPHAN", "DEPRECATE"],
            ["30", "cmms_checklist_mst", "928", "9", "STRAIGHTFORWARD", "KEEP"],
            ["31", "cmms_checklist_tasks", "7536", "2", "STRAIGHTFORWARD", "KEEP"],
            ["32", "cmms_checklist_hist", "811", "7", "STRAIGHTFORWARD", "KEEP"],
            ["33", "cmms_checklist_tasks_hist", "8450", "4", "STRAIGHTFORWARD", "KEEP"],
            ["34", "cmms_task_mst", "1489", "11", "STRAIGHTFORWARD", "KEEP"],
        ])

    add_heading(doc, "3.5 Cluster 6 - Maintenance / Repair (Phase 2 sketch)", level=2)
    add_table(doc,
        ["#", "Table", "Rows", "Cols", "Flag", "Verdict"],
        [
            ["35", "cmms_jobcard_repair_info", "8118", "18", "STRAIGHTFORWARD", "KEEP"],
            ["36", "cmms_jobcard_inspection_info", "2214", "21", "STRAIGHTFORWARD", "KEEP"],
            ["37", "cmms_jobcard_insp_maint_dtl", "0", "6", "ORPHAN", "DEPRECATE"],
            ["38", "cmms_jobcard_awaitinginfo", "7261", "19", "STRAIGHTFORWARD", "KEEP"],
            ["39", "cmms_jobcard_contract_warranty_dtl", "17225", "17", "STRAIGHTFORWARD", "KEEP"],
            ["40", "cmms_jobcard_faulty_category", "8605", "3", "STRAIGHTFORWARD", "KEEP"],
            ["41", "cmms_jobcard_faulty_section", "8131", "3", "STRAIGHTFORWARD", "KEEP"],
            ["42", "cmms_jobcard_spares_equip", "2804", "9", "STRAIGHTFORWARD", "KEEP"],
            ["43", "cmms_fault_mst", "30", "8", "STRAIGHTFORWARD", "KEEP"],
        ])

    add_heading(doc, "3.6 Clusters 2, 7, 8, 12 + Legacy", level=2)
    add_table(doc,
        ["#", "Table", "Rows", "Cols", "Flag", "Verdict", "Cluster"],
        [
            ["44", "cmms_section_mst", "293", "14", "STRAIGHTFORWARD", "KEEP", "2 Org"],
            ["45", "cmms_proj_mst", "182", "7", "STRAIGHTFORWARD", "KEEP", "2 Org"],
            ["46", "cmms_schedule_mst", "6", "8", "STRAIGHTFORWARD", "KEEP (P2)", "7 Sched"],
            ["47", "cmms_schedule_eqip_dtl", "316", "5", "STRAIGHTFORWARD", "KEEP (P2)", "7 Sched"],
            ["48", "cmms_po_mst", "115", "16", "STRAIGHTFORWARD", "KEEP (P2)", "8 Procure"],
            ["49", "cmms_pur_mst", "0", "22", "ORPHAN", "DEPRECATE", "8 Procure"],
            ["50", "cmms_pur_dtl", "0", "13", "ORPHAN", "DEPRECATE", "8 Procure"],
            ["51", "cmms_amc_mst", "0", "11", "ORPHAN + REVISION", "DEPRECATE", "8 Procure"],
            ["52", "cmms_inv_mst", "42", "29", "STRAIGHTFORWARD", "KEEP (P2)", "8 Procure"],
            ["53", "cmms_device_spares_mst", "67", "8", "STRAIGHTFORWARD", "KEEP (P2)", "8 Procure"],
            ["54", "cmms_lineitem_mst", "24", "9", "STRAIGHTFORWARD", "KEEP (P2)", "8 Procure"],
            ["55", "chklistvendor", "238", "3", "DECISION-NEEDED", "KEEP (P2)", "8 Procure"],
            ["56", "cmms_documentno_mst", "151", "6", "STRAIGHTFORWARD", "KEEP", "12 Lookup"],
            ["57", "cmms_parameter_master", "337", "4", "STRAIGHTFORWARD", "KEEP", "12 Lookup"],
            ["58", "cmms_parameter_master_bkp", "4", "4", "ORPHAN", "DEPRECATE", "12 Lookup"],
            ["59", "cmms_parameter_master_incharge", "9", "4", "ORPHAN", "DEPRECATE", "12 Lookup"],
            ["60", "cmms_parameter_master_jun2016", "233", "4", "ORPHAN", "DEPRECATE", "12 Lookup"],
            ["61", "cf001", "6", "7", "ORPHAN", "DEPRECATE", "Legacy"],
            ["62", "cf002", "553", "4", "ORPHAN", "DEPRECATE", "Legacy"],
            ["63", "cf003", "570", "9", "ORPHAN", "DEPRECATE", "Legacy"],
            ["64", "cf004", "3449", "2", "ORPHAN", "DEPRECATE", "Legacy"],
        ])

    add_heading(doc, "3.7 Tally", level=2)
    add_table(doc, ["Verdict", "Count"], [
        ["KEEP (as-is or KEEP+BRIDGE)", "42"],
        ["EXTEND / REFACTOR", "5"],
        ["REPLACE (BR-violation -> new design)", "3"],
        ["DEPRECATE (orphan / unused / backup / dead)", "14"],
        ["TOTAL", "64"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 4 - CLUSTER 1 DEEP
    # =====================================================================
    add_section_banner(doc, "Section 4 - Cluster 1: Identity & Access (DEEP DESIGN)")

    add_heading(doc, "4.1 The 8 New Tables", level=2)
    add_table(doc, ["New Table", "Purpose"], [
        ["cmcm_users", "Auth identity bridging cmms_emp_mst -> modern auth"],
        ["cmcm_roles", "Exactly 5 roles (Super Admin -> View-Only)"],
        ["cmcm_permissions", "resource:action permission codes"],
        ["cmcm_role_permissions", "Many-to-many permission grants per role"],
        ["cmcm_user_roles", "One role per user (PK on user_id alone)"],
        ["cmcm_refresh_tokens", "Hashed refresh tokens, JWT jti claim, rotation tracking"],
        ["cmcm_login_audit", "Every login attempt (success/failure) per BR-AUTH-06"],
        ["cmcm_password_reset_tokens", "Forced first-login reset flow"],
    ])

    add_heading(doc, "4.2 cmcm_users - DDL", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_users (\n"
        "  id                     BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  employee_id            VARCHAR(7)   NOT NULL,\n"
        "  password_hash          VARCHAR(60)  NOT NULL,\n"
        "  password_must_change   TINYINT(1)   NOT NULL DEFAULT 1,\n"
        "  password_changed_at    DATETIME(6)  NULL,\n"
        "  is_active              TINYINT(1)   NOT NULL DEFAULT 1,\n"
        "  is_locked              TINYINT(1)   NOT NULL DEFAULT 0,\n"
        "  failed_login_attempts  INT UNSIGNED NOT NULL DEFAULT 0,\n"
        "  last_login_at          DATETIME(6)  NULL,\n"
        "  last_login_ip          VARCHAR(45)  NULL,\n"
        "  created_at             DATETIME(6)  NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  updated_at             DATETIME(6)  NOT NULL DEFAULT CURRENT_TIMESTAMP(6) ON UPDATE CURRENT_TIMESTAMP(6),\n"
        "  created_by             BIGINT UNSIGNED NULL,\n"
        "  updated_by             BIGINT UNSIGNED NULL,\n"
        "  PRIMARY KEY (id),\n"
        "  UNIQUE KEY uk_employee_id (employee_id),\n"
        "  KEY ix_active_locked (is_active, is_locked),\n"
        "  KEY ix_last_login (last_login_at),\n"
        "  CONSTRAINT fk_cmcm_users_emp FOREIGN KEY (employee_id)\n"
        "    REFERENCES cmms_emp_mst (EMM_ID)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "4.3 cmcm_roles - DDL + Seed", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_roles (\n"
        "  id              INT UNSIGNED NOT NULL,\n"
        "  code            VARCHAR(50)  NOT NULL,\n"
        "  display_name    VARCHAR(100) NOT NULL,\n"
        "  description     VARCHAR(255) NULL,\n"
        "  hierarchy_level TINYINT UNSIGNED NOT NULL,\n"
        "  is_system       TINYINT(1)  NOT NULL DEFAULT 1,\n"
        "  created_at      DATETIME(6) NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  updated_at      DATETIME(6) NOT NULL DEFAULT CURRENT_TIMESTAMP(6) ON UPDATE CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  UNIQUE KEY uk_code (code)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;\n"
        "\n"
        "INSERT INTO cmcm_roles (id, code, display_name, hierarchy_level) VALUES\n"
        "  (1, 'SUPER_ADMIN',   'Super Admin',    1),\n"
        "  (2, 'LAB_INCHARGE',  'Lab In-charge',  2),\n"
        "  (3, 'LAB_ENGINEER',  'Lab Engineer',   3),\n"
        "  (4, 'NORMAL_USER',   'Normal User',    4),\n"
        "  (5, 'VIEW_ONLY',     'View-Only User', 5);")

    add_heading(doc, "4.4 cmcm_permissions - DDL + Seed (sample)", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_permissions (\n"
        "  id          INT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  code        VARCHAR(100) NOT NULL,\n"
        "  resource    VARCHAR(50)  NOT NULL,\n"
        "  action      VARCHAR(50)  NOT NULL,\n"
        "  module      VARCHAR(50)  NOT NULL,\n"
        "  description VARCHAR(255) NULL,\n"
        "  is_system   TINYINT(1)   NOT NULL DEFAULT 1,\n"
        "  created_at  DATETIME(6)  NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  updated_at  DATETIME(6)  NOT NULL DEFAULT CURRENT_TIMESTAMP(6) ON UPDATE CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  UNIQUE KEY uk_code (code),\n"
        "  KEY ix_module (module)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;\n"
        "\n"
        "-- Seed (sample - full list ~35 permissions):\n"
        "INSERT INTO cmcm_permissions (code, resource, action, module, description) VALUES\n"
        "  ('auth:login',                'auth',       'login',         'auth',       'Submit credentials'),\n"
        "  ('me:read',                   'me',         'read',          'auth',       'Read own profile'),\n"
        "  ('user:read-list',            'user',       'read-list',     'users',      'List users (SA only)'),\n"
        "  ('user:role-assign',          'user',       'role-assign',   'users',      'Change a user role (SA only)'),\n"
        "  ('equipment:read-list',       'equipment',  'read-list',     'equipment',  'List equipment'),\n"
        "  ('equipment:create',          'equipment',  'create',        'equipment',  'Register new equipment'),\n"
        "  ('equipment:verify',          'equipment',  'verify',        'equipment',  'Flip PENDING -> ACTIVE'),\n"
        "  ('equipment:delete',          'equipment',  'delete',        'equipment',  'Hard delete (SA only)'),\n"
        "  ('job_request:create',        'job_request','create',        'jobReq',     'Raise a job request'),\n"
        "  ('job_request:approve',       'job_request','approve',       'jobReq',     'Approve a request'),\n"
        "  ('job_card:verify-close',     'job_card',   'verify-close',  'jobCard',    'Verify and close a card'),\n"
        "  ('job_card:generate-pdf',     'job_card',   'generate-pdf',  'jobCard',    'Generate PDF (no storage)'),\n"
        "  ('audit_log:read',            'audit_log',  'read',          'audit',      'Read central audit log'),\n"
        "  ('export:trigger',            'export',     'trigger',       'audit',      'Trigger export (PDF/future Excel)');\n"
        "-- ... plus all rows from FINAL-DESC Section 5.")

    add_heading(doc, "4.5 cmcm_user_roles - DDL (enforces BR-RBAC-02)", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_user_roles (\n"
        "  user_id      BIGINT UNSIGNED NOT NULL,\n"
        "  role_id      INT UNSIGNED    NOT NULL,\n"
        "  assigned_at  DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  assigned_by  BIGINT UNSIGNED NULL,\n"
        "  notes        VARCHAR(255)    NULL,\n"
        "  PRIMARY KEY (user_id),     -- <- PK on user_id alone forces 1 role per user\n"
        "  KEY ix_role (role_id),\n"
        "  CONSTRAINT fk_cmcm_user_roles_user FOREIGN KEY (user_id) REFERENCES cmcm_users (id),\n"
        "  CONSTRAINT fk_cmcm_user_roles_role FOREIGN KEY (role_id) REFERENCES cmcm_roles (id)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "4.6 cmcm_role_permissions - DDL", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_role_permissions (\n"
        "  role_id       INT UNSIGNED    NOT NULL,\n"
        "  permission_id INT UNSIGNED    NOT NULL,\n"
        "  granted_at    DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  granted_by    BIGINT UNSIGNED NULL,\n"
        "  PRIMARY KEY (role_id, permission_id),\n"
        "  KEY ix_perm (permission_id),\n"
        "  CONSTRAINT fk_cmcm_rp_role FOREIGN KEY (role_id)       REFERENCES cmcm_roles (id) ON DELETE CASCADE,\n"
        "  CONSTRAINT fk_cmcm_rp_perm FOREIGN KEY (permission_id) REFERENCES cmcm_permissions (id) ON DELETE CASCADE\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "4.7 cmcm_refresh_tokens - DDL", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_refresh_tokens (\n"
        "  id            BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  user_id       BIGINT UNSIGNED NOT NULL,\n"
        "  jti           VARCHAR(36)     NOT NULL,\n"
        "  token_hash    CHAR(64)        NOT NULL,    -- SHA-256, never raw\n"
        "  expires_at    DATETIME(6)     NOT NULL,    -- 7-day absolute cap\n"
        "  last_used_at  DATETIME(6)     NULL,        -- sliding 60-min idle\n"
        "  revoked_at    DATETIME(6)     NULL,\n"
        "  ip_address    VARCHAR(45)     NULL,\n"
        "  user_agent    VARCHAR(255)    NULL,\n"
        "  created_at    DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  UNIQUE KEY uk_jti (jti),\n"
        "  KEY ix_user_active (user_id, revoked_at),\n"
        "  KEY ix_expires (expires_at),\n"
        "  CONSTRAINT fk_cmcm_rt_user FOREIGN KEY (user_id) REFERENCES cmcm_users (id)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "4.8 cmcm_login_audit - DDL", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_login_audit (\n"
        "  id              BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  employee_id     VARCHAR(7)      NOT NULL,   -- not FK: failed attempts may not match\n"
        "  user_id         BIGINT UNSIGNED NULL,\n"
        "  attempt_at      DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  success         TINYINT(1)      NOT NULL,\n"
        "  failure_reason  ENUM('invalid_credentials','user_locked','user_inactive',\n"
        "                       'user_not_found','password_must_change') NULL,\n"
        "  ip_address      VARCHAR(45)     NULL,\n"
        "  user_agent      VARCHAR(255)    NULL,\n"
        "  PRIMARY KEY (id),\n"
        "  KEY ix_emp_time (employee_id, attempt_at),\n"
        "  KEY ix_success_time (success, attempt_at),\n"
        "  CONSTRAINT fk_cmcm_login_user FOREIGN KEY (user_id) REFERENCES cmcm_users (id)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "4.9 cmcm_password_reset_tokens - DDL", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_password_reset_tokens (\n"
        "  id          BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  user_id     BIGINT UNSIGNED NOT NULL,\n"
        "  token_hash  CHAR(64)        NOT NULL,\n"
        "  expires_at  DATETIME(6)     NOT NULL,    -- typical 1 hour\n"
        "  used_at     DATETIME(6)     NULL,\n"
        "  created_at  DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  KEY ix_user_active (user_id, used_at),\n"
        "  CONSTRAINT fk_cmcm_prt_user FOREIGN KEY (user_id) REFERENCES cmcm_users (id)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "4.10 OLD vs NEW comparison", level=2)
    add_table(doc, ["Concern", "OLD (legacy)", "NEW (cmcm)"], [
        ["User identity", "cmms_userrole_mst", "cmcm_users (FK -> cmms_emp_mst.EMM_ID)"],
        ["Password storage", "USER_PASSWORD VARCHAR(10) plaintext", "password_hash VARCHAR(60) bcryptjs"],
        ["Force reset", "none", "password_must_change column"],
        ["Lockout", "none", "is_locked + failed_login_attempts"],
        ["Last login", "none", "last_login_at + last_login_ip"],
        ["Roles", "cmms_role_mst (23 rows)", "cmcm_roles (exactly 5)"],
        ["One role per user", "not enforced", "PK on cmcm_user_roles.user_id"],
        ["Permissions", "cmms_accessright_mst (module CRUD bits)", "cmcm_permissions (resource:action codes)"],
        ["Permission grant", "role x module x {add,view,edit,delete,print}", "cmcm_role_permissions(role_id, permission_id)"],
        ["Sidebar visibility", "cmms_module_mst + access bits", "permissions drive visibility (BR-RBAC-04)"],
        ["Refresh tokens", "none", "cmcm_refresh_tokens (hashed)"],
        ["Login audit", "none", "cmcm_login_audit"],
        ["Audit log", "per-entity _hist", "central cmcm_audit_log"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 5 - CLUSTER 10
    # =====================================================================
    add_section_banner(doc, "Section 5 - Cluster 10: Audit & Logs (DEEP DESIGN)")

    add_heading(doc, "5.1 cmcm_audit_log - the central log (BR-AUD-01..05)", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_audit_log (\n"
        "  id                 BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  entity_type        VARCHAR(50)     NOT NULL,         -- 'equipment','job_request',...\n"
        "  entity_id          VARCHAR(64)     NOT NULL,         -- supports compound legacy keys\n"
        "  action             ENUM('INSERT','UPDATE','DELETE','STATE_CHANGE','EXPORT',\n"
        "                          'LOGIN','LOGOUT','ROLE_ASSIGN','VERIFY','CONDEMN') NOT NULL,\n"
        "  actor_user_id      BIGINT UNSIGNED NULL,\n"
        "  actor_employee_id  VARCHAR(7)      NULL,             -- denormalised for fast read\n"
        "  actor_role_code    VARCHAR(50)     NULL,             -- role at the time\n"
        "  before_json        JSON            NULL,\n"
        "  after_json         JSON            NULL,\n"
        "  reason             VARCHAR(500)    NULL,             -- mandatory for REJECT/REOPEN/CONDEMN\n"
        "  ip_address         VARCHAR(45)     NULL,\n"
        "  user_agent         VARCHAR(255)    NULL,\n"
        "  request_id         VARCHAR(36)     NULL,             -- correlates with Pino logs\n"
        "  created_at         DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  KEY ix_entity (entity_type, entity_id, created_at),\n"
        "  KEY ix_actor (actor_user_id, created_at),\n"
        "  KEY ix_action (action, created_at),\n"
        "  CONSTRAINT fk_cmcm_audit_actor FOREIGN KEY (actor_user_id) REFERENCES cmcm_users (id)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "5.2 Why legacy *_hist tables stay", level=2)
    add_bullets(doc, [
        "cmcm_audit_log is the master record (BR-AUD-01).",
        "Legacy cmms_*_hist tables are kept as fast per-entity timelines.",
        "Going forward we ALWAYS write to cmcm_audit_log; cmms_*_hist are read-only views.",
        "cmms_jobcard_status_hist gets a NEW PK (id BIGINT AUTO_INCREMENT) - BR-VIOLATION fix #2.",
    ])

    add_heading(doc, "5.3 Fix for cmms_jobcard_status_hist (no PK)", level=2)
    add_mono_block(doc,
        "ALTER TABLE cmms_jobcard_status_hist\n"
        "  ADD COLUMN id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT FIRST,\n"
        "  ADD PRIMARY KEY (id);")

    add_page_break(doc)

    # =====================================================================
    # SECTION 6 - CLUSTER 12
    # =====================================================================
    add_section_banner(doc, "Section 6 - Cluster 12: Lookups (DEEP DESIGN)")

    add_heading(doc, "6.1 cmcm_lookup_values - DDL", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_lookup_values (\n"
        "  id           INT UNSIGNED    NOT NULL AUTO_INCREMENT,\n"
        "  category     VARCHAR(50)     NOT NULL,\n"
        "  code         VARCHAR(50)     NOT NULL,\n"
        "  display_name VARCHAR(100)    NOT NULL,\n"
        "  sort_order   INT             NOT NULL DEFAULT 0,\n"
        "  is_active    TINYINT(1)      NOT NULL DEFAULT 1,\n"
        "  meta_json    JSON            NULL,        -- e.g. colour hex for status pills\n"
        "  created_at   DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  updated_at   DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6) ON UPDATE CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  UNIQUE KEY uk_cat_code (category, code),\n"
        "  KEY ix_category_active (category, is_active)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "6.2 Seed - equipment_status (Cluster 12 + BR-EQP-10)", level=2)
    add_mono_block(doc,
        "INSERT INTO cmcm_lookup_values (category, code, display_name, sort_order) VALUES\n"
        "  ('equipment_status', 'PENDING_VERIFICATION', 'Pending Verification', 1),\n"
        "  ('equipment_status', 'ACTIVE',               'Active',               2),\n"
        "  ('equipment_status', 'UNDER_CALIBRATION',    'Under Calibration',    3),\n"
        "  ('equipment_status', 'UNDER_REPAIR',         'Under Repair',         4),\n"
        "  ('equipment_status', 'OUT_OF_TOLERANCE',     'Out of Tolerance',     5),\n"
        "  ('equipment_status', 'QUARANTINED',          'Quarantined',          6),\n"
        "  ('equipment_status', 'CONDEMNED',            'Condemned',            7),\n"
        "  ('equipment_status', 'RETIRED',              'Retired',              8);")

    add_heading(doc, "6.3 Other categories to seed", level=2)
    add_bullets(doc, [
        "job_request_status: DRAFT, SUBMITTED, ASSIGNED, IN_PROGRESS, COMPLETED, VERIFIED, REJECTED, REOPENED",
        "job_card_status: ASSIGNED, IN_PROGRESS, COMPLETED, VERIFIED_CLOSED, REOPENED",
        "job_type: CALIBRATION, REPAIR, REGISTRATION",
        "priority: LOW, NORMAL, HIGH, CRITICAL",
        "rejection_reason: INSUFFICIENT_INFO, WRONG_EQUIPMENT, OUT_OF_SCOPE, OTHER",
        "system_type: TME, FPE (sourced from cmms_product_mst flags)",
        "Plus: keep cmms_documentno_mst for JC2026-00001 numbering (already 151 rows).",
        "Plus: keep cmms_parameter_master for legacy lookups (337 rows).",
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 7 - CLUSTER 3 EQUIPMENT
    # =====================================================================
    add_section_banner(doc, "Section 7 - Cluster 3: Equipment Master (REFACTOR)")

    add_heading(doc, "7.1 New columns on cmms_eqip_mst", level=2)
    add_mono_block(doc,
        "ALTER TABLE cmms_eqip_mst\n"
        "  ADD COLUMN status_code VARCHAR(40) NOT NULL DEFAULT 'PENDING_VERIFICATION'\n"
        "    AFTER EQM_DIV_STATUS,\n"
        "  ADD COLUMN registered_by_user_id BIGINT UNSIGNED NULL,\n"
        "  ADD COLUMN registered_at         DATETIME(6)     NULL,\n"
        "  ADD COLUMN verified_by_user_id   BIGINT UNSIGNED NULL,\n"
        "  ADD COLUMN verified_at           DATETIME(6)     NULL,\n"
        "  ADD COLUMN equipment_uid BIGINT UNSIGNED\n"
        "    GENERATED ALWAYS AS\n"
        "      (CAST(CONCAT(IF(EQM_TYPE='TME','1','2'),LPAD(EQM_ID,10,'0')) AS UNSIGNED))\n"
        "    STORED UNIQUE,\n"
        "  ADD KEY ix_status (status_code),\n"
        "  ADD KEY ix_registered_at (registered_at),\n"
        "  ADD CONSTRAINT fk_eqip_status FOREIGN KEY (status_code)\n"
        "    REFERENCES cmcm_lookup_values (code), -- (with category filter handled in app)\n"
        "  ADD CONSTRAINT fk_eqip_registered_by FOREIGN KEY (registered_by_user_id)\n"
        "    REFERENCES cmcm_users (id),\n"
        "  ADD CONSTRAINT fk_eqip_verified_by FOREIGN KEY (verified_by_user_id)\n"
        "    REFERENCES cmcm_users (id);")

    add_heading(doc, "7.2 New table - cmcm_equipment_status_history", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_equipment_status_history (\n"
        "  id                BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  equipment_uid     BIGINT UNSIGNED NOT NULL,\n"
        "  from_status       VARCHAR(40)     NULL,\n"
        "  to_status         VARCHAR(40)     NOT NULL,\n"
        "  changed_by_user_id BIGINT UNSIGNED NOT NULL,\n"
        "  reason            VARCHAR(500)    NULL,\n"
        "  changed_at        DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  KEY ix_eq_time (equipment_uid, changed_at),\n"
        "  CONSTRAINT fk_eqsh_eq FOREIGN KEY (equipment_uid) REFERENCES cmms_eqip_mst (equipment_uid),\n"
        "  CONSTRAINT fk_eqsh_user FOREIGN KEY (changed_by_user_id) REFERENCES cmcm_users (id)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "7.3 Equipment cluster final table list", level=2)
    add_table(doc, ["Status", "Table"], [
        ["EXTEND", "cmms_eqip_mst (+6 new columns)"],
        ["KEEP", "cmms_eqip_mst_hist"],
        ["KEEP", "cmms_eqipinst_identification"],
        ["KEEP", "cmms_ins_accuracy_info"],
        ["KEEP", "cmms_division_hist"],
        ["KEEP", "cmms_product_mst (T&ME/F&PE)"],
        ["DEPRECATE", "cmms_eqip_detail_spec (longblob, no-storage violation)"],
        ["NEW", "cmcm_equipment_status_history"],
        ["NEW", "cmcm_equipment_calibration_history (optional view)"],
        ["NEW", "cmcm_equipment_tags (optional)"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 8 - CLUSTER 4 JOB LIFECYCLE
    # =====================================================================
    add_section_banner(doc, "Section 8 - Cluster 4: Job Lifecycle (REFACTOR)")

    add_heading(doc, "8.1 New columns on cmms_jobrequest_mst", level=2)
    add_mono_block(doc,
        "ALTER TABLE cmms_jobrequest_mst\n"
        "  ADD COLUMN status_code             VARCHAR(40) NOT NULL DEFAULT 'DRAFT',\n"
        "  ADD COLUMN priority_code           VARCHAR(40) NULL,\n"
        "  ADD COLUMN submitted_by_user_id    BIGINT UNSIGNED NULL,\n"
        "  ADD COLUMN approved_by_user_id     BIGINT UNSIGNED NULL,\n"
        "  ADD COLUMN approved_at             DATETIME(6)     NULL,\n"
        "  ADD COLUMN rejection_reason_code   VARCHAR(40)     NULL,\n"
        "  ADD COLUMN rejection_reason_text   VARCHAR(500)    NULL,\n"
        "  ADD KEY ix_status (status_code),\n"
        "  ADD KEY ix_submitted_by (submitted_by_user_id),\n"
        "  ADD CONSTRAINT fk_jr_submitted FOREIGN KEY (submitted_by_user_id) REFERENCES cmcm_users (id),\n"
        "  ADD CONSTRAINT fk_jr_approved  FOREIGN KEY (approved_by_user_id)  REFERENCES cmcm_users (id);")

    add_heading(doc, "8.2 New columns on cmms_jobcard_mst", level=2)
    add_mono_block(doc,
        "ALTER TABLE cmms_jobcard_mst\n"
        "  ADD COLUMN status_code               VARCHAR(40) NOT NULL DEFAULT 'ASSIGNED',\n"
        "  ADD COLUMN assigned_engineer_user_id BIGINT UNSIGNED NULL,\n"
        "  ADD COLUMN started_at                DATETIME(6)     NULL,\n"
        "  ADD COLUMN completed_at              DATETIME(6)     NULL,\n"
        "  ADD COLUMN verified_by_user_id       BIGINT UNSIGNED NULL,\n"
        "  ADD COLUMN verified_at               DATETIME(6)     NULL,\n"
        "  ADD COLUMN reopen_reason             VARCHAR(500)    NULL,\n"
        "  ADD KEY ix_status (status_code),\n"
        "  ADD KEY ix_assigned (assigned_engineer_user_id),\n"
        "  ADD CONSTRAINT fk_jc_assignee FOREIGN KEY (assigned_engineer_user_id) REFERENCES cmcm_users (id),\n"
        "  ADD CONSTRAINT fk_jc_verifier FOREIGN KEY (verified_by_user_id)       REFERENCES cmcm_users (id);")

    add_heading(doc, "8.3 New table - cmcm_job_request_status_history", level=2)
    add_mono_block(doc,
        "CREATE TABLE cmcm_job_request_status_history (\n"
        "  id                BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,\n"
        "  job_request_no    INT             NOT NULL,\n"
        "  from_status       VARCHAR(40)     NULL,\n"
        "  to_status         VARCHAR(40)     NOT NULL,\n"
        "  changed_by_user_id BIGINT UNSIGNED NOT NULL,\n"
        "  reason            VARCHAR(500)    NULL,\n"
        "  changed_at        DATETIME(6)     NOT NULL DEFAULT CURRENT_TIMESTAMP(6),\n"
        "  PRIMARY KEY (id),\n"
        "  KEY ix_jr_time (job_request_no, changed_at),\n"
        "  CONSTRAINT fk_jrsh_jr FOREIGN KEY (job_request_no) REFERENCES cmms_jobrequest_mst (JR_JOBREQUESTNO),\n"
        "  CONSTRAINT fk_jrsh_user FOREIGN KEY (changed_by_user_id) REFERENCES cmcm_users (id)\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;")

    add_heading(doc, "8.4 The two decisions you need to call", level=2)
    add_table(doc, ["Decision", "My Recommendation"], [
        ["Equipment compound PK", "Keep (EQM_TYPE, EQM_ID) compound PK. ADD `equipment_uid` BIGINT GENERATED column UNIQUE for new FKs. Zero break."],
        ["Job card canonical id", "Treat JM_SectionJobNo (varchar(9)) as canonical. JM_JobCardNO (int) becomes legacy alt-key."],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 9 - PHASE 2 SKETCHES
    # =====================================================================
    add_section_banner(doc, "Section 9 - Phase 2 Clusters (Sketches)")

    add_heading(doc, "9.1 Cluster 2 - Organisation", level=2)
    add_bullets(doc, [
        "KEEP cmms_section_mst (293 rows) as canonical divisions/labs master.",
        "KEEP cmms_proj_mst (182 rows) as projects.",
        "KEEP cmms_section_user_mst as division-user mapping.",
        "Optional NEW: cmcm_division_role_scope (Phase 2; for row-level visibility tuning).",
    ])

    add_heading(doc, "9.2 Cluster 5 - Calibration", level=2)
    add_bullets(doc, [
        "KEEP cmms_jobcard_cal_dtl, _cal_adjustments_dtl, _cal_observations.",
        "KEEP cmms_checklist_mst, _tasks, _hist, _tasks_hist, cmms_task_mst.",
        "DEPRECATE cmms_cal_jobcard_feedback_spec (0 rows + longblob).",
    ])

    add_heading(doc, "9.3 Cluster 6 - Maintenance", level=2)
    add_bullets(doc, [
        "KEEP cmms_jobcard_repair_info, _inspection_info, _awaitinginfo, _contract_warranty_dtl, _faulty_*, _spares_equip.",
        "KEEP cmms_fault_mst.",
        "DEPRECATE cmms_jobcard_insp_maint_dtl (0 rows).",
    ])

    add_heading(doc, "9.4 Cluster 7 - Scheduling", level=2)
    add_bullets(doc, [
        "KEEP cmms_schedule_mst (6 rows) + cmms_schedule_eqip_dtl (316 rows) for Phase 2.",
        "NEW: cmcm_schedule_notifications (when notifications go live).",
    ])

    add_heading(doc, "9.5 Cluster 8 - Procurement", level=2)
    add_bullets(doc, [
        "KEEP cmms_po_mst (115 rows), cmms_inv_mst (42), cmms_device_spares_mst (67), cmms_lineitem_mst (24).",
        "DEPRECATE cmms_pur_mst (0), cmms_pur_dtl (0), cmms_amc_mst (0 + col-type swap bug).",
        "Need to confirm: cmms_cont_mst (vendors) - missing from dump but referenced by 4+ FKs.",
        "Probable NEW (Phase 2): cmcm_vendors if cmms_cont_mst is dead.",
    ])

    add_heading(doc, "9.6 Cluster 9 - Documents", level=2)
    add_bullets(doc, [
        "No new tables; PDF is generate-and-download only (constraint #5).",
        "DEPRECATE longblob columns in cmms_eqip_detail_spec + cmms_cal_jobcard_feedback_spec.",
    ])

    add_heading(doc, "9.7 Cluster 11 - Notifications", level=2)
    add_bullets(doc, [
        "All NEW in Phase 2 (no legacy).",
        "NEW: cmcm_notifications, cmcm_notification_templates.",
        "Channels deferred (user will instruct).",
    ])

    add_heading(doc, "9.8 Cluster 13 - Reporting", level=2)
    add_bullets(doc, [
        "All NEW in Phase 2 (no legacy).",
        "NEW: cmcm_report_definitions, cmcm_report_exports (audit-only).",
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 10 - MIGRATION
    # =====================================================================
    add_section_banner(doc, "Section 10 - Zero-Drop Migration Strategy")

    add_callout(doc,
        "Rule: in MVP we NEVER DROP a table. Orphans get quarantined "
        "(suffix _deprecated_2026) and stay readable.",
        fill="FFF4E5"
    )

    add_mono_block(doc,
        "STEP 1   Create all NEW cmcm_* tables (no data risk).\n"
        "STEP 2   Seed 5 roles + ~35 permissions + role_permissions matrix.\n"
        "STEP 3   Seed Super Admins (from env SUPER_ADMIN_EMPLOYEE_IDS=SA79900,AC77777):\n"
        "         INSERT cmcm_users (employee_id='SA79900', password_hash=bcrypt(random()),\n"
        "                             password_must_change=1, ...)\n"
        "         INSERT cmcm_users (employee_id='AC77777', ...)\n"
        "         INSERT cmcm_user_roles (user_id, role_id=1 SUPER_ADMIN) for both.\n"
        "         Write 'BOOTSTRAP' rows to cmcm_audit_log.\n"
        "STEP 4   Migrate cmms_userrole_mst (565 rows) -> cmcm_users:\n"
        "         For each row:\n"
        "           - INSERT cmcm_users with bcrypt(random()) placeholder\n"
        "           - password_must_change = 1 (forces reset on first login)\n"
        "           - Map cmms_userrole_mst.USER_ROLE -> 5-role bucket via approved mapping\n"
        "           - INSERT cmcm_user_roles\n"
        "STEP 5   ADD COLUMNS on cmms_eqip_mst (status_code, registered_by_user_id, ...)\n"
        "         Backfill:\n"
        "           UPDATE cmms_eqip_mst SET status_code = CASE\n"
        "             WHEN LOWER(EQM_DIV_STATUS) LIKE '%condemn%' THEN 'CONDEMNED'\n"
        "             WHEN LOWER(EQM_DIV_STATUS) LIKE '%quarant%' THEN 'QUARANTINED'\n"
        "             WHEN LOWER(EQM_DIV_STATUS) LIKE '%retired%' THEN 'RETIRED'\n"
        "             WHEN EQM_DIV_STATUS IS NULL OR EQM_DIV_STATUS = '' THEN 'ACTIVE'\n"
        "             ELSE 'ACTIVE'\n"
        "           END;\n"
        "           UPDATE cmms_eqip_mst SET registered_by_user_id = (\n"
        "             SELECT id FROM cmcm_users WHERE employee_id = EQM_CREATED_BY\n"
        "           );\n"
        "STEP 6   ADD COLUMNS on cmms_jobrequest_mst + cmms_jobcard_mst.\n"
        "STEP 7   Fix cmms_jobcard_status_hist (no PK):\n"
        "           ALTER TABLE cmms_jobcard_status_hist\n"
        "             ADD COLUMN id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT FIRST,\n"
        "             ADD PRIMARY KEY (id);\n"
        "STEP 8   Quarantine 14 ORPHAN tables (suffix _deprecated_2026):\n"
        "           RENAME TABLE cmms_parameter_master_bkp TO cmms_parameter_master_bkp_deprecated_2026;\n"
        "           ... (repeat for all 14)\n"
        "STEP 9   Run smoke test: login as SA79900, force pw change, view equipment list,\n"
        "         create a fake equipment in PENDING_VERIFICATION, verify it.\n"
        "STEP 10  Commit. Move to Phase 0 (repo scaffold) when Auth tables are live.")

    add_page_break(doc)

    # =====================================================================
    # SECTION 11 - QUESTIONS
    # =====================================================================
    add_section_banner(doc, "Section 11 - 8 Questions Awaiting Your Answers")

    add_table(doc, ["Q#", "Question", "Why It Matters"], [
        ["Q1", "Can you share the 23 role names from cmms_role_mst, or shall I propose a name-pattern mapping (e.g. names containing 'Admin' -> SUPER_ADMIN) and you confirm?",
         "Blocks user role assignment for 565 existing users."],
        ["Q2", "Does cmms_cont_mst (vendors/contacts) actually exist? It's referenced by 4+ FKs but not in the dump.",
         "Determines vendors data source for Phase 2."],
        ["Q3", "Equipment compound PK (EQM_TYPE, EQM_ID): keep + add equipment_uid surrogate (my recommendation), or migrate to single int id?",
         "Affects 12+ FK chains."],
        ["Q4", "Confirm JM_SectionJobNo is the canonical job card id (not JM_JobCardNO)?",
         "PK referenced by 20+ FKs; appears on printed cards."],
        ["Q5", "Confirm cf001/cf002/cf003/cf004 (4 unbound legacy tables) DEPRECATE?",
         "They have data (6/553/570/3449 rows) but no FKs."],
        ["Q6", "Plaintext password migration: confirm force first-login reset for all 565 users (no decoding plaintexts)?",
         "Security best practice; one-time user friction."],
        ["Q7", "Is cmms_section_mst (293 rows) the canonical 'division/lab' master for BR-VIS row-level scoping?",
         "Drives Normal/Engineer visibility scope."],
        ["Q8", "Should new numbers continue using cmms_documentno_mst-style sequences (e.g. JC2026-00001)?",
         "Affects display + reports + audit."],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 12 - CONVENTIONS
    # =====================================================================
    add_section_banner(doc, "Section 12 - Updated Naming + Numbering Conventions")

    add_table(doc, ["Convention", "Rule"], [
        ["Legacy tables", "Keep cmms_* names (don't rename to avoid breaking legacy queries)"],
        ["New tables", "Prefix cmcm_* (distinguishes new from old at a glance)"],
        ["Quarantined tables", "Suffix _deprecated_2026 (move out of cmms_ to make orphan obvious)"],
        ["Primary keys", "BIGINT UNSIGNED AUTO_INCREMENT id for new tables; legacy compound PKs stay"],
        ["FK columns", "Singular_table_id pattern (e.g. role_id, permission_id)"],
        ["Timestamps", "created_at, updated_at; DATETIME(6) precision (microsecond)"],
        ["Audit fields", "created_by, updated_by; BIGINT UNSIGNED FK to cmcm_users.id"],
        ["Status fields", "VARCHAR(40) + FK to cmcm_lookup_values (NOT raw ENUM in new tables)"],
        ["Soft-delete", "is_deleted TINYINT(1) + deleted_at; reserved for Phase 2 if needed"],
        ["JSON columns", "Use sparingly; only for audit before/after diffs and meta_json on lookups"],
        ["Doc numbering", "Continue cmms_documentno_mst pattern (e.g. JC2026-00001)"],
        ["Money", "DECIMAL(18,4) - never FLOAT"],
        ["Charset / Engine", "utf8mb4_0900_ai_ci + InnoDB (matches existing)"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 13 - NEXT STEPS
    # =====================================================================
    add_section_banner(doc, "Section 13 - Next Steps")

    add_heading(doc, "13.1 What I will do next (after your answers)", level=2)
    add_bullets(doc, [
        "Lock the 23-role mapping (Q1) so all 565 users can be migrated cleanly.",
        "Confirm cmms_cont_mst (Q2) before designing Cluster 8 (Procurement).",
        "Lock equipment PK strategy (Q3) before writing equipment repository code.",
        "Confirm job-card canonical id (Q4) before writing job card repository code.",
        "Confirm cf00x deprecation (Q5).",
        "Write the full seed SQL: roles, permissions, role_permissions matrix, super admins.",
        "Write the ALTER TABLE migration SQL for status_code + audit columns on legacy tables.",
        "Write the cmms_jobcard_status_hist PK fix migration.",
        "Lock the migration SQL into db/migrations/ folder (numbered files).",
    ])

    add_heading(doc, "13.2 The cluster build order for MVP weeks 3-10", level=2)
    add_table(doc, ["Order", "Cluster", "Why first"], [
        ["1", "Cluster 1 - Identity & Access", "Nothing else works without login + RBAC"],
        ["2", "Cluster 10 - Audit & Logs", "Every write needs to log; build the log writer first"],
        ["3", "Cluster 12 - Lookups", "Lookup data is referenced by every other cluster"],
        ["4", "Cluster 2 - Organisation", "Divisions/projects referenced by equipment + jobs"],
        ["5", "Cluster 3 - Equipment Master", "MVP core; everything else depends on this"],
        ["6", "Cluster 4 - Job Lifecycle", "MVP core; the operational heart of the system"],
        ["7", "Cluster 5 + 6 - Calibration + Maintenance details", "Job Card depth"],
        ["8", "Cluster 9 - Documents (PDF gen)", "Closes the calibration loop"],
        ["9", "Phase 2 Clusters", "After MVP"],
    ])

    add_locked_banner(doc,
        "END OF FINAL-DB-DESIGN v1  -  awaiting answers to Q1-Q8 before locking v2.")

    return doc


if __name__ == "__main__":
    out = build()
    out_path = r"e:\SOFTWAREs By DS\cmcmis-simplified\Documents\FINAL-DB-DESIGN-v1.docx"
    out.save(out_path)
    print(f"Saved: {out_path}")
