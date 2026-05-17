"""
FINAL-DESC-CMCMIS.docx generator.
The gold-standard final project description, with every decision locked.
Builds on the v3 master plan and absorbs all D1-D11 + C1-C5 confirmations,
plus 6 stack additions and the JS+JSDoc+Zod / Raw-SQL+Repository / Nginx-prod /
feature-FE-folders / pdfkit / Pino / RQ / Zustand stack.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
#                  STYLING HELPERS
# ============================================================

NAVY = RGBColor(0x0B, 0x2A, 0x5B)
ACCENT = RGBColor(0x1E, 0x6F, 0xA8)
GREEN = RGBColor(0x1B, 0x7F, 0x3A)
RED = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x55, 0x55, 0x55)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TABLE_HEADER_FILL = "1E6FA8"
SECTION_FILL = "0B2A5B"
LOCKED_FILL = "1B7F3A"


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size is None:
        size = {1: 22, 2: 17, 3: 14, 4: 12}.get(level, 12)
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_callout(doc, text, fill="EAF3FB", text_color=None):
    if text_color is None:
        text_color = NAVY
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = text_color
    run.bold = True
    return tbl


def add_locked_banner(doc, text):
    """Green LOCKED banner for confirmed decisions."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, LOCKED_FILL)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run("LOCKED  -  " + text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


def add_mono_block(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F4F4F4")
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return tbl


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, TABLE_HEADER_FILL)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = tbl.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            v = str(val).strip()
            if v in ("YES", "Y", "LOCKED", "✓"):
                run.font.color.rgb = GREEN
                run.bold = True
            elif v in ("NO", "N", "X"):
                run.font.color.rgb = RED
                run.bold = True
            elif v.startswith("MUST"):
                run.font.color.rgb = NAVY
                run.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)


def add_page_break(doc):
    doc.add_page_break()


def add_section_banner(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, SECTION_FILL)
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


# ============================================================
#                       BUILD
# ============================================================

def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------------- COVER ----------------
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\nFINAL-DESC-CMCMIS")
    r.bold = True; r.font.size = Pt(38); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("The Final Project Description\n")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Computerized Maintenance & Calibration\nManagement Information System")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\nEvery Decision Locked  |  Every Module Mapped  |  Every Rule Written\n")
    r.italic = True; r.font.size = Pt(13); r.font.color.rgb = GREY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "\nFor the 10-Week Industry-Grade Production Build\n"
        "Solo Developer + AI Pair  ·  ISRO SAC-style on-prem deployment\n\n"
        "Style: structured · table-heavy · flowchart-rich · beginner-to-advanced\n"
    )
    r.font.size = Pt(11); r.font.color.rgb = NAVY

    add_locked_banner(
        doc,
        "All Decisions Confirmed on 2026-05-16  -  Ready for Phase 3 (DB Design)"
    )

    add_page_break(doc)

    # ---------------- TOC ----------------
    add_heading(doc, "Table of Contents", level=1)
    toc = [
        "Section 1 — Executive Summary (One-Page Brief)",
        "Section 2 — Project at a Glance",
        "Section 3 — The 9-Module Architecture Map",
        "Section 4 — Role Model (5 Roles) + 3-Layer RBAC",
        "Section 5 — Final Consolidated Permission Matrix",
        "Section 6 — Authentication & Session Model",
        "Section 7 — Critical State Machines",
        "Section 8 — Business Rules Catalogue",
        "Section 9 — Functional Requirements (MVP)",
        "Section 10 — Non-Functional Requirements",
        "Section 11 — Database Strategy & Approach",
        "Section 12 — Locked Tech Stack (Frontend + Backend + Additions)",
        "Section 13 — Architectural Decision Register (D1–D11, C1–C5)",
        "Section 14 — Code Architecture Patterns "
        "(JS+JSDoc+Zod, Repository, Feature-FE)",
        "Section 15 — System Architecture & Request Trace",
        "Section 16 — Express Middleware Pipeline",
        "Section 17 — Production Deployment Topology (Nginx + PM2)",
        "Section 18 — Operational Strategies "
        "(PDF, Logging, Testing, CSRF, Security)",
        "Section 19 — Project Folder Structure (FINAL)",
        "Section 20 — Coding Conventions & API Standards",
        "Section 21 — MVP Scope vs Phase 2 Backlog",
        "Section 22 — 10-Week War Plan",
        "Section 23 — Real-World End-to-End Walk-Through",
        "Section 24 — Constraints, Quick Recap, What's Next",
    ]
    for line in toc:
        add_para(doc, "•  " + line, size=11)

    add_page_break(doc)

    # =====================================================================
    # SECTION 1 — EXECUTIVE SUMMARY
    # =====================================================================
    add_section_banner(doc, "Section 1 — Executive Summary (One-Page Brief)")

    add_callout(
        doc,
        "CMCMIS is a defence-grade Computerized Maintenance & Calibration "
        "Management Information System for a high-precision engineering org "
        "(ISRO SAC-style). It tracks every test/measurement and functional "
        "instrument through its full lifecycle — register, calibrate, "
        "maintain, repair, retire — under a strict 3-layer RBAC system."
    )

    add_heading(doc, "1.1 What CMCMIS Does", level=2)
    add_bullets(doc, [
        "Knows every instrument: serial, model, division, last/next "
        "calibration date, status.",
        "Routes work: Normal Users raise Job Requests; Lab In-charge "
        "approves; Lab Engineer executes via Job Cards.",
        "Records evidence: tasks done, readings taken, before/after, "
        "environmental conditions, signatures.",
        "Generates PDFs on demand: job cards, calibration certificates "
        "(no storage — pure print).",
        "Logs everything: audit trail on every state-changing action.",
        "Enforces order: state machines for Job Request + Equipment "
        "prevent illegal transitions.",
    ])

    add_heading(doc, "1.2 Who Uses CMCMIS", level=2)
    add_table(doc,
        ["Role", "Primary Job"],
        [
            ["Super Admin", "Assigns roles, owns master data, ultimate authority"],
            ["Lab In-charge", "Approves requests, verifies equipment, closes job cards"],
            ["Lab Engineer", "Executes calibration/repair, fills job cards"],
            ["Normal User", "Raises job requests, registers new equipment"],
            ["View-Only User", "Reads everything; cannot write"],
        ],
    )

    add_heading(doc, "1.3 How It's Built (Locked Stack)", level=2)
    add_table(doc,
        ["Layer", "Technology"],
        [
            ["Frontend", "React 18 + Vite + Tailwind v3 + TanStack Query + Zustand + Axios + RHF + Zod"],
            ["Backend", "Node + Express 4 + mysql2/promise + Zod + JWT + bcryptjs + Pino + pdfkit + PM2"],
            ["Database", "MySQL 8.x (existing ~64 tables + new) via phpMyAdmin"],
            ["Production", "Nginx (TLS + SPA + /api proxy) → PM2 cluster → MySQL on-prem"],
            ["Type system", "JavaScript + JSDoc + Zod (NO TypeScript)"],
            ["Data access", "Raw SQL + Repository pattern (NO ORM)"],
        ],
    )

    add_heading(doc, "1.4 Timeline & Outcome", level=2)
    add_para(doc,
        "10 weeks for MVP (Auth+RBAC, Equipment, Job Requests, Job Cards, "
        "Dashboard, Inquiry, PDF, audit). Phase 2 post-internship "
        "(Schedule, Procurement, Reports, Admin master CRUD, Notifications). "
        "Deployed on-prem to org private infrastructure. Demo-ready first; "
        "full go-live after stakeholder sign-off.")

    add_page_break(doc)

    # =====================================================================
    # SECTION 2 — PROJECT AT A GLANCE
    # =====================================================================
    add_section_banner(doc, "Section 2 — Project at a Glance")

    add_heading(doc, "2.1 One-Sentence Definition", level=2)
    add_callout(doc,
        "\"CMCMIS is the system that knows where every instrument in our labs "
        "is, what state it's in, when it needs calibration, who's working on "
        "it, and what was done to it — for as long as the instrument exists.\""
    )

    add_heading(doc, "2.2 The Domain in 5 Words", level=2)
    add_mono_block(doc, "REGISTER  -->  CALIBRATE  -->  MAINTAIN  -->  REPAIR  -->  RETIRE")

    add_heading(doc, "2.3 The Strategic Picture", level=2)
    add_table(doc,
        ["Layer", "What Lives Here", "Why It Matters"],
        [
            ["People", "5 roles, hundreds of users", "Wrong access = data corruption / compliance breach"],
            ["Process", "2 strict state machines", "Wrong transitions = audit failure"],
            ["Assets", "Thousands of instruments + full history", "Lose history = lose calibration provenance"],
            ["Money", "POs, spares, vendor contracts (Phase 2)", "Wrong PO chain = procurement audit issues"],
            ["Knowledge", "Reports, certificates, signatures", "This IS the institutional memory"],
        ],
    )

    add_heading(doc, "2.4 Engagement Context", level=2)
    add_table(doc,
        ["Aspect", "Decision"],
        [
            ["Organization context", "ISRO SAC-like government / defence / space-grade"],
            ["Project status", "Real organizational production system (NOT a prototype)"],
            ["Owner", "Solo Software-Developer intern (with AI pair)"],
            ["Timeline", "10 weeks for MVP; full go-live after stakeholder approval"],
            ["Deployment", "Internal on-prem / private infra (NOT public cloud)"],
            ["Existing DB", "~64 tables already exist — reviewed in Phase 3"],
            ["Stack", "Locked — see Section 12"],
            ["Decisions", "All locked — see Section 13"],
        ],
    )

    add_page_break(doc)

    # =====================================================================
    # SECTION 3 — MODULE ARCHITECTURE
    # =====================================================================
    add_section_banner(doc, "Section 3 — The 9-Module Architecture Map")

    add_heading(doc, "3.1 Module Map (MVP + Phase 2)", level=2)
    add_mono_block(doc,
        "                  +---------------------------------+\n"
        "                  |   DASHBOARD (read-only)         |\n"
        "                  |   Reads -> ALL modules          |\n"
        "                  +----------------+----------------+\n"
        "                                   |\n"
        "        +--------------------------+--------------------------+\n"
        "        v                          v                          v\n"
        "  +-------------+         +------------------+        +----------------+\n"
        "  | JOB         |--creates| JOB CARDS        |<-touch-| EQUIPMENT      |\n"
        "  | REQUESTS    |-------->|  (work execution)|        |  (asset master)|\n"
        "  | (intake)    |         |                  |        |                |\n"
        "  +-----+-------+         +-----+------------+        +--------+-------+\n"
        "        |                       | updates                      | schedules\n"
        "        |                       v                              v\n"
        "        |             +------------------+             +------------------+\n"
        "        |             | DOCUMENTS        |             | SCHEDULE  (P2)   |\n"
        "        |             | (PDF generation) |             | (PM + Cal cal)   |\n"
        "        |             +------------------+             +------------------+\n"
        "        v\n"
        "  +--------------+     +----------------+\n"
        "  | PROCUREMENT  |<--->| VENDORS  (P2)  |\n"
        "  |  (P2)        |     | (master data)  |\n"
        "  +--------------+     +----------------+\n"
        "\n"
        "  +-----------------+   +-----------------+   +------------------+\n"
        "  | INQUIRY         |   | REPORTS  (P2)   |   | ADMIN  (P2)      |\n"
        "  | (search hub)    |   | (analytics)     |   | (master + RBAC)  |\n"
        "  +-----------------+   +-----------------+   +------------------+\n"
        "\n"
        "  Legend:  (P2) = Phase 2 / post-internship.  All others MVP."
    )

    add_heading(doc, "3.2 Module Responsibility Matrix", level=2)
    add_table(doc,
        ["#", "Module", "Owns (CRUD)", "Reads From", "Writes To", "Primary Actors", "MVP?"],
        [
            ["1", "Dashboard", "Nothing", "All modules", "Nothing", "All roles", "YES"],
            ["2", "Job Requests", "job_requests", "Equipment, Employees", "Job Cards on approval", "Normal -> Lab In-charge", "YES"],
            ["3", "Job Cards", "job_cards, tasks, observations", "Job Requests, Equipment", "Documents, Equipment", "Lab Engineer, In-charge", "YES"],
            ["4", "Equipment", "equipment, specs, cal_history", "Job Cards, Procurement", "Schedule", "All except View-Only", "YES"],
            ["5", "Schedule", "schedules", "Equipment", "Auto-Job Requests", "In-charge, Engineer", "Phase 2"],
            ["6", "Procurement", "purchase_orders, spare_parts", "Vendors, Equipment", "Equipment stock", "In-charge", "Phase 2"],
            ["7", "Vendors", "vendors", "—", "Procurement", "Super Admin", "Phase 2"],
            ["8", "Inquiry", "Nothing (read-only)", "Everything", "—", "All roles", "YES"],
            ["9", "Reports", "Nothing (read-only)", "All modules", "Export logs", "In-charge, Super Admin", "Phase 2"],
            ["10", "Admin", "Master data + Users + Roles", "—", "All master tables", "Super Admin only", "Phase 2 (RBAC live in MVP)"],
        ],
    )

    add_heading(doc, "3.3 The Golden Rule of Data Flow", level=2)
    add_mono_block(doc,
        "  EQUIPMENT (source of truth)\n"
        "       |\n"
        "       v\n"
        "  JOB REQUEST (references equipment_id)\n"
        "       |\n"
        "       v\n"
        "  JOB CARD (references job_request_id + equipment_id)\n"
        "       |\n"
        "       v\n"
        "  DOCUMENT (references job_card_id, generated on demand)\n"
        "       |\n"
        "       v\n"
        "  EQUIPMENT.last_calibration_date  <-- the ONLY back-write\n"
        "                                     (only via state machine)"
    )
    add_para(doc, "Rule: data flows downstream; only one well-defined hook "
                  "back-writes calibration metadata to Equipment.", italic=True)

    add_page_break(doc)

    # =====================================================================
    # SECTION 4 — ROLES + RBAC
    # =====================================================================
    add_section_banner(doc, "Section 4 — Role Model (5 Roles) + 3-Layer RBAC")

    add_locked_banner(doc, "C1: Five roles final. Admin role DELETED. Super Admin is the only admin tier.")

    add_heading(doc, "4.1 The 5-Role Hierarchy", level=2)
    add_mono_block(doc,
        "                +------------------------+\n"
        "                |  SUPER ADMIN           | <-- role-assigner,\n"
        "                |  (only admin tier)     |     master data owner,\n"
        "                +-----------+------------+     ultimate authority\n"
        "                            |\n"
        "                +-----------v------------+\n"
        "                |  LAB IN-CHARGE         | <-- approves jobs,\n"
        "                |                        |     verifies/closes\n"
        "                +-----------+------------+\n"
        "                            |\n"
        "                +-----------v------------+\n"
        "                |  LAB ENGINEER          | <-- executes work,\n"
        "                |                        |     fills job cards\n"
        "                +-----------+------------+\n"
        "                            |\n"
        "             +--------------+--------------+\n"
        "             v                             v\n"
        "  +--------------------+        +---------------------+\n"
        "  |  NORMAL USER       |        |  VIEW-ONLY USER     |\n"
        "  | raises requests +  |        | read-only --        |\n"
        "  | registers equip!   |        | no writes EVER      |\n"
        "  +--------------------+        +---------------------+"
    )

    add_heading(doc, "4.2 The 3-Layer RBAC Model", level=2)
    add_mono_block(doc,
        "   USER  --->  ROLE  --->  PERMISSION  --->  RESOURCE+ACTION\n"
        "    |          |              |                     |\n"
        "    |          |              |                     +-- \"equipment:create\"\n"
        "    |          |              +-- Granular toggle stored in DB\n"
        "    |          +-- Role bundle (e.g. \"Lab Engineer\")\n"
        "    +-- Identity (employee_id)\n"
    )

    add_table(doc,
        ["Layer", "Real-world Example", "Stored In"],
        [
            ["User", "Employee #EMP1234 (Deep)", "users"],
            ["Role", "Lab Engineer", "roles + user_roles"],
            ["Permission", "job_card:update, equipment:read", "permissions + role_permissions"],
            ["Resource:Action", "Checked at API middleware", "Runtime check"],
        ],
    )

    add_heading(doc, "4.3 Why 3 Layers (Not 2)", level=2)
    add_bullets(doc, [
        "2-layer (User -> Role) requires a redeploy every time business changes 'who can do what'.",
        "3-layer (User -> Role -> Permission) lets Super Admin toggle a checkbox and behaviour changes instantly.",
        "BR-RBAC-03: code NEVER checks by role name. Always by permission. This is the contract.",
    ])

    add_heading(doc, "4.4 The Equipment Registration Override", level=2)
    add_callout(doc,
        "All roles EXCEPT View-Only can register a new equipment. "
        "Even a Normal User can. Verification (PENDING_VERIFICATION -> ACTIVE) "
        "still requires Lab In-charge or Super Admin (per C5 + D10).",
        fill="FFF4E5"
    )

    add_page_break(doc)

    # =====================================================================
    # SECTION 5 — FINAL PERMISSION MATRIX
    # =====================================================================
    add_section_banner(doc, "Section 5 — Final Consolidated Permission Matrix")
    add_para(doc, "Replaces every prior permission matrix. Single source of truth.",
             italic=True, color=GREY)

    role_cols = ["Resource:Action", "Normal", "View-Only", "Lab Engr", "Lab In-charge", "Super Admin"]

    add_heading(doc, "5.1 Authentication & Identity", level=2)
    add_table(doc, role_cols, [
        ["auth:login", "YES", "YES", "YES", "YES", "YES"],
        ["auth:logout", "YES", "YES", "YES", "YES", "YES"],
        ["auth:refresh-token", "YES", "YES", "YES", "YES", "YES"],
        ["me:read (own profile)", "YES", "YES", "YES", "YES", "YES"],
    ])

    add_heading(doc, "5.2 User & Role Management (Super Admin only)", level=2)
    add_table(doc, role_cols, [
        ["user:read-list", "NO", "NO", "NO", "NO", "YES"],
        ["user:role-assign", "NO", "NO", "NO", "NO", "YES"],
        ["user:activate / deactivate", "NO", "NO", "NO", "NO", "YES"],
    ])

    add_heading(doc, "5.3 Equipment (create open to ALL except View-Only)", level=2)
    add_table(doc, role_cols, [
        ["equipment:read-list", "YES", "YES", "YES", "YES", "YES"],
        ["equipment:read-detail", "YES", "YES", "YES", "YES", "YES"],
        ["equipment:create", "YES", "NO", "YES", "YES", "YES"],
        ["equipment:update", "NO", "NO", "YES", "YES", "YES"],
        ["equipment:verify (PENDING -> ACTIVE)", "NO", "NO", "NO", "YES", "YES"],
        ["equipment:delete (hard)", "NO", "NO", "NO", "NO", "YES"],
        ["equipment:condemn (status flip)", "NO", "NO", "NO", "YES", "YES"],
    ])

    add_heading(doc, "5.4 Job Requests", level=2)
    add_table(doc, role_cols, [
        ["job_request:create", "YES", "NO", "YES", "YES", "YES"],
        ["job_request:read-own", "YES", "YES", "YES", "YES", "YES"],
        ["job_request:read-all", "NO", "YES", "YES", "YES", "YES"],
        ["job_request:approve", "NO", "NO", "NO", "YES", "YES"],
        ["job_request:reject", "NO", "NO", "NO", "YES", "YES"],
        ["job_request:assign-engineer", "NO", "NO", "NO", "YES", "YES"],
    ])

    add_heading(doc, "5.5 Job Cards (Super Admin + Lab In-charge full)", level=2)
    add_table(doc, role_cols, [
        ["job_card:read-list", "NO", "YES", "YES", "YES", "YES"],
        ["job_card:read-detail", "NO", "YES", "YES", "YES", "YES"],
        ["job_card:start-work", "NO", "NO", "YES (own)", "YES", "YES"],
        ["job_card:update-tasks", "NO", "NO", "YES (own)", "YES", "YES"],
        ["job_card:complete", "NO", "NO", "YES (own)", "YES", "YES"],
        ["job_card:verify-close", "NO", "NO", "NO", "YES", "YES"],
        ["job_card:reopen", "NO", "NO", "NO", "YES", "YES"],
        ["job_card:generate-pdf", "NO", "YES", "YES", "YES", "YES"],
    ])

    add_heading(doc, "5.6 Dashboard & Inquiry", level=2)
    add_table(doc, role_cols, [
        ["dashboard:view", "YES", "YES", "YES", "YES", "YES"],
        ["inquiry:search-vendors", "YES", "YES", "YES", "YES", "YES"],
        ["inquiry:search-products", "YES", "YES", "YES", "YES", "YES"],
        ["inquiry:search-job-cards", "NO", "YES", "YES", "YES", "YES"],
        ["inquiry:search-instruments", "YES", "YES", "YES", "YES", "YES"],
    ])

    add_heading(doc, "5.7 Master Data (Super Admin only — Phase 2)", level=2)
    add_table(doc, role_cols, [
        ["master:employees:manage", "NO", "NO", "NO", "NO", "YES"],
        ["master:vendors:manage", "NO", "NO", "NO", "NO", "YES"],
        ["master:equipment-types:manage", "NO", "NO", "NO", "NO", "YES"],
        ["master:divisions:manage", "NO", "NO", "NO", "NO", "YES"],
        ["master:lookup-values:manage", "NO", "NO", "NO", "NO", "YES"],
    ])

    add_heading(doc, "5.8 Audit & Logs", level=2)
    add_table(doc, role_cols, [
        ["audit_log:read", "NO", "NO", "NO", "NO", "YES"],
        ["export:trigger (PDF / future Excel)", "NO", "NO", "YES", "YES", "YES"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 6 — AUTH FLOW
    # =====================================================================
    add_section_banner(doc, "Section 6 — Authentication & Session Model")

    add_heading(doc, "6.1 Login Flow (v1: employee_id + password; SSO-ready)", level=2)
    add_mono_block(doc,
        "  User submits credentials  OR  (future) SSO identity assertion\n"
        "                       |\n"
        "                       v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 1: Validate employee_id                 |\n"
        "  |  - v1: bcryptjs hash check vs DB             |\n"
        "  |  - future: SSO signature verify              |\n"
        "  +----------------------------------------------+\n"
        "                       | valid\n"
        "                       v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 2: Lookup user in `users` table         |\n"
        "  |  - exists? is_active? is_locked?             |\n"
        "  +----------------------------------------------+\n"
        "                       |\n"
        "          +------------+------------+\n"
        "          v                         v\n"
        "    NOT FOUND                   FOUND\n"
        "          |                         |\n"
        "          v                         v\n"
        "  +----------------+   +---------------------------+\n"
        "  | Auto-provision |   | Fetch user_roles JOIN     |\n"
        "  | as Normal User |   |     roles -> permissions  |\n"
        "  +-------+--------+   +-------------+-------------+\n"
        "          |                          |\n"
        "          +------------+-------------+\n"
        "                       v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 3: Generate ACCESS token (JWT)          |\n"
        "  |   - sub: employee_id                         |\n"
        "  |   - role + permissions[]                     |\n"
        "  |   - exp: 15 min  +  jti (for revoke list)    |\n"
        "  +-----------------+----------------------------+\n"
        "                    v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 4: Issue REFRESH cookie                 |\n"
        "  |   - httpOnly + Secure + SameSite=Lax         |\n"
        "  |   - exp: 7 days (hashed in DB)               |\n"
        "  +-----------------+----------------------------+\n"
        "                    v\n"
        "       +---------------------------+\n"
        "       | Authenticated session     |\n"
        "       | -> write login_audit row  |\n"
        "       +---------------------------+"
    )

    add_heading(doc, "6.2 Session Envelopes (Three Concentric Timers)", level=2)
    add_mono_block(doc,
        "  +========================================================+\n"
        "  |  REFRESH TOKEN MAX  (7 days, absolute)                 |\n"
        "  |  +--------------------------------------------------+  |\n"
        "  |  |  60-MIN INACTIVITY WINDOW (sliding)              |  |\n"
        "  |  |  +--------------------------------------------+  |  |\n"
        "  |  |  | 15-MIN JWT ACCESS TOKEN (short-lived)      |  |  |\n"
        "  |  |  | refreshed silently on each authed API call |  |  |\n"
        "  |  |  +--------------------------------------------+  |  |\n"
        "  |  |  resets on every authenticated request           |  |\n"
        "  |  +--------------------------------------------------+  |\n"
        "  |  hard expiry; user MUST re-login                       |\n"
        "  +========================================================+"
    )

    add_heading(doc, "6.3 First Super Admin Bootstrap (D11)", level=2)
    add_locked_banner(doc, "D11: Seed at least 2 Super Admin employee IDs via "
                           "SUPER_ADMIN_EMPLOYEE_IDS env var")
    add_bullets(doc, [
        "Seed via DB migration on first deploy (NOT manual SQL).",
        "Env var: SUPER_ADMIN_EMPLOYEE_IDS (comma-separated; ≥ 2 required).",
        "Migration: INSERT users + INSERT user_roles (SUPER_ADMIN) + INSERT audit_log 'BOOTSTRAP' rows.",
        "First login forces password reset (password_must_change = TRUE).",
        "IDs will be provided by user before seed migration runs.",
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 7 — STATE MACHINES
    # =====================================================================
    add_section_banner(doc, "Section 7 — Critical State Machines")

    add_heading(doc, "7.1 Job Request Lifecycle", level=2)
    add_mono_block(doc,
        "  +---------+   submit    +--------------+   assign   +--------------+\n"
        "  | DRAFT   |------------>| SUBMITTED    |----------->| ASSIGNED     |\n"
        "  +---------+             +------+-------+            +------+-------+\n"
        "                                 | reject                    | start\n"
        "                                 v                           v\n"
        "                          +--------------+           +--------------+\n"
        "                          | REJECTED     |           | IN-PROGRESS  |\n"
        "                          +--------------+           +------+-------+\n"
        "                                                            | complete\n"
        "                                                            v\n"
        "                                                     +--------------+\n"
        "                                                     | COMPLETED    |\n"
        "                                                     +------+-------+\n"
        "                                                            | verify\n"
        "                                                            v\n"
        "                                                     +----------------+\n"
        "                                                     | VERIFIED/CLOSED|\n"
        "                                                     +------+---------+\n"
        "                                                            | reopen\n"
        "                                                            v\n"
        "                                                     +--------------+\n"
        "                                                     | REOPENED     |\n"
        "                                                     +--------------+"
    )

    add_table(doc, ["From -> To", "Allowed By", "Side Effects"], [
        ["DRAFT -> SUBMITTED", "Owner", "Notify Lab In-charge"],
        ["SUBMITTED -> ASSIGNED", "Lab In-charge", "Auto-create Job Card"],
        ["SUBMITTED -> REJECTED", "Lab In-charge", "Mandatory reason; lock further edits"],
        ["ASSIGNED -> IN-PROGRESS", "Assigned Engineer", "Set start_time"],
        ["IN-PROGRESS -> COMPLETED", "Assigned Engineer", "BR-JC-07: observations required"],
        ["COMPLETED -> VERIFIED/CLOSED", "Lab In-charge", "Cert PDF available; update last_cal_date"],
        ["VERIFIED -> REOPENED", "Lab In-charge / Super Admin", "Audit log mandatory reason"],
    ])

    add_heading(doc, "7.2 Equipment Lifecycle (with PENDING_VERIFICATION per D10)", level=2)
    add_locked_banner(doc, "D10: All new equipment defaults to PENDING_VERIFICATION; "
                           "Lab In-charge / Super Admin flip to ACTIVE")
    add_mono_block(doc,
        "  [PENDING_VERIFICATION] --verify by InC/SA--> [ACTIVE]\n"
        "                                                 |\n"
        "                                                 +---> [UNDER_CALIBRATION] --> [ACTIVE]\n"
        "                                                 |             |\n"
        "                                                 |             +--> [OUT_OF_TOLERANCE] --> [UNDER_REPAIR] --> [ACTIVE]\n"
        "                                                 |\n"
        "                                                 +---> [UNDER_REPAIR] --> [ACTIVE]\n"
        "                                                 |\n"
        "                                                 +---> [QUARANTINED] --> [CONDEMNED/RETIRED]  (terminal)"
    )

    add_heading(doc, "7.3 Implementation Pattern (Single Choke-Point)", level=2)
    add_mono_block(doc,
        "  // backend/src/stateMachines/jobRequest.fsm.js\n"
        "  // ONE function per entity. Single source of truth.\n"
        "  export function transitionJobRequest(currentState, action, actor) {\n"
        "    const allowed = ALLOWED_TRANSITIONS[currentState]?.[action];\n"
        "    if (!allowed) throw new IllegalTransitionError(currentState, action);\n"
        "    if (!actor.hasPermission(allowed.permission))\n"
        "        throw new ForbiddenError(allowed.permission);\n"
        "    return { newState: allowed.to, sideEffects: allowed.effects };\n"
        "  }\n"
        "  // Service layer calls transition() inside a DB transaction,\n"
        "  // then writes the *_status_history row + audit_log row."
    )
    add_para(doc,
        "Effect: no developer writes raw UPDATE status= anywhere. Every "
        "transition is logged automatically. Audit trail is free.",
        italic=True)

    add_page_break(doc)

    # =====================================================================
    # SECTION 8 — BUSINESS RULES (compressed reference)
    # =====================================================================
    add_section_banner(doc, "Section 8 — Business Rules Catalogue")
    add_para(doc, "37 rules across 9 categories. Contract between business and code.",
             italic=True, color=GREY)

    add_heading(doc, "8.1 BR-AUTH — Authentication & Identity", level=2)
    add_table(doc, ["ID", "Rule"], [
        ["BR-AUTH-01", "Login is by employee_id only (not email/username)."],
        ["BR-AUTH-02", "User must exist in org's employee directory — no self-registration."],
        ["BR-AUTH-03", "Authenticated user with no role → defaults to Normal User."],
        ["BR-AUTH-04", "Sessions expire after 60 minutes of inactivity. Refresh token valid 7 days."],
        ["BR-AUTH-05", "First Super Admin seeded via DB migration using SUPER_ADMIN_EMPLOYEE_IDS env var. Seed ≥ 2."],
        ["BR-AUTH-06", "All login attempts (success + failure) are logged."],
        ["BR-AUTH-07", "Deactivated user (status = INACTIVE) cannot log in but history preserved."],
    ])

    add_heading(doc, "8.2 BR-RBAC — Authorization", level=2)
    add_table(doc, ["ID", "Rule"], [
        ["BR-RBAC-01", "Only Super Admin can assign/change a user's role."],
        ["BR-RBAC-02", "A user has exactly ONE primary role at a time (no multi-role in v1)."],
        ["BR-RBAC-03", "Permissions derived: User -> Role -> Permissions. NEVER check by role name."],
        ["BR-RBAC-04", "Sidebar items and routes filtered by permissions, not role names."],
        ["BR-RBAC-05", "Every API endpoint enforces permission check at controller layer."],
        ["BR-RBAC-06", "View-Only users can read but cannot trigger any write operation."],
        ["BR-RBAC-07", "Role changes take effect on next login OR token refresh."],
    ])

    add_heading(doc, "8.3 BR-EQP — Equipment", level=2)
    add_table(doc, ["ID", "Rule"], [
        ["BR-EQP-01", "Every equipment has a unique serial number system-wide."],
        ["BR-EQP-02", "Equipment must belong to T&ME or F&PE category."],
        ["BR-EQP-03", "F&PE may have NO calibration frequency (optional)."],
        ["BR-EQP-04", "T&ME MUST have calibration frequency (months)."],
        ["BR-EQP-05", "next_cal_due_date = last_cal_date + calibration_frequency_months."],
        ["BR-EQP-06", "Status transitions follow the equipment state machine."],
        ["BR-EQP-07", "Hard DELETE = Super Admin only. CONDEMN = Lab In-charge or Super Admin."],
        ["BR-EQP-08", "Search is case-insensitive across serial, model, manufacturer, type."],
        ["BR-EQP-09", "Every registration captures registered_by + verified_by + timestamps."],
        ["BR-EQP-10", "New equipment defaults to PENDING_VERIFICATION; verify by Lab In-charge / Super Admin -> ACTIVE."],
    ])

    add_heading(doc, "8.4 BR-JR — Job Requests", level=2)
    add_table(doc, ["ID", "Rule"], [
        ["BR-JR-01", "Must reference existing equipment (or trigger registration)."],
        ["BR-JR-02", "Job type in {Calibration, Repair, Registration}."],
        ["BR-JR-03", "User can save as DRAFT before submitting."],
        ["BR-JR-04", "Once SUBMITTED, only Lab In-charge can change state."],
        ["BR-JR-05", "Submitted request must be assigned to a Lab Engineer before becoming a Job Card."],
        ["BR-JR-06", "'submitted by' auto-filled from current user — not overridable."],
        ["BR-JR-07", "High-priority repairs appear at top of Lab In-charge queue."],
        ["BR-JR-08", "Rejection requires mandatory reason (free text + reason code)."],
    ])

    add_heading(doc, "8.5 BR-JC — Job Cards", level=2)
    add_table(doc, ["ID", "Rule"], [
        ["BR-JC-01", "Auto-created when a request is approved + assigned."],
        ["BR-JC-02", "Lifecycle: ASSIGNED -> IN-PROGRESS -> COMPLETED -> VERIFIED/CLOSED."],
        ["BR-JC-03", "Only the assigned engineer can mark IN-PROGRESS or COMPLETED."],
        ["BR-JC-04", "Only Lab In-charge can verify/close."],
        ["BR-JC-05", "Closed job can only be REOPENED by Lab In-charge with reason."],
        ["BR-JC-06", "Tasks are configurable per job type."],
        ["BR-JC-07", "Calibration cards require before-reading + after-reading + environment before COMPLETED."],
        ["BR-JC-08", "Job card history is append-only — immutable state transition log."],
    ])

    add_heading(doc, "8.6 BR-PDF, BR-AUD, BR-VIS, BR-MASTER", level=2)
    add_table(doc, ["ID", "Rule"], [
        ["BR-PDF-01", "PDFs generated on demand from current DB state. Nothing stored."],
        ["BR-PDF-02", "Job card PDF: header, equipment info, tasks, observations, signatures, date."],
        ["BR-PDF-03", "Calibration cert PDF: equipment, standards, readings, environment, uncertainty, valid-until."],
        ["BR-PDF-04", "All PDFs include generation timestamp + footer + record ID."],
        ["BR-AUD-01", "Every write on critical tables is logged in audit_logs."],
        ["BR-AUD-02", "Audit log captures who, what, when, before/after, IP, user-agent."],
        ["BR-AUD-03", "Exports logged with user + record IDs accessed."],
        ["BR-AUD-04", "Sensitive fields filtered from API responses based on permission."],
        ["BR-AUD-05", "HTTPS in prod. JWT in Authorization header. Refresh in httpOnly cookie."],
        ["BR-VIS-01", "Normal User sees only their own job requests."],
        ["BR-VIS-02", "Lab Engineer sees all jobs assigned to them + the queue."],
        ["BR-VIS-03", "Lab In-charge and above see all jobs and equipment."],
        ["BR-VIS-04", "View-Only sees all data, performs no actions."],
        ["BR-MASTER-01", "All master-data CRUD is Super Admin only (Phase 2 UI)."],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 9 — FRs
    # =====================================================================
    add_section_banner(doc, "Section 9 — Functional Requirements (MVP)")
    add_para(doc, "45 FRs across 6 MVP modules. 44 MUST + 1 SHOULD.",
             italic=True, color=GREY)

    add_heading(doc, "9.1 FR-A — Auth & RBAC", level=2)
    add_table(doc, ["FR ID", "Description", "Priority"], [
        ["FR-A-01", "Login screen accepts employee_id + password", "MUST"],
        ["FR-A-02", "Successful login returns JWT + sets refresh cookie", "MUST"],
        ["FR-A-03", "Logout clears tokens and invalidates session", "MUST"],
        ["FR-A-04", "Protected routes redirect to login when unauthenticated", "MUST"],
        ["FR-A-05", "Sidebar items render based on user's permission set", "MUST"],
        ["FR-A-06", "Super Admin can view list of all users + their roles", "MUST"],
        ["FR-A-07", "Super Admin can assign/change a user's role", "MUST"],
        ["FR-A-08", "Super Admin can activate/deactivate a user", "MUST"],
        ["FR-A-09", "Backend exposes GET /me returning current user + permissions", "MUST"],
        ["FR-A-10", "Token refresh endpoint silently extends session", "MUST"],
    ])

    add_heading(doc, "9.2 FR-E — Equipment", level=2)
    add_table(doc, ["FR ID", "Description", "Priority"], [
        ["FR-E-01", "List equipment with pagination, search, filter (type, status)", "MUST"],
        ["FR-E-02", "View equipment detail page with specs, cal history, linked jobs", "MUST"],
        ["FR-E-03", "Register new equipment (all roles except View-Only)", "MUST"],
        ["FR-E-04", "Edit equipment details (specs, status)", "MUST"],
        ["FR-E-05", "Color-code calibration due dates (green/yellow/red)", "MUST"],
        ["FR-E-06", "Show equipment's calibration history timeline", "MUST"],
        ["FR-E-07", "Search equipment by serial, model, manufacturer", "MUST"],
        ["FR-E-08", "Soft-delete (status = CONDEMNED) for non-super-admins", "MUST"],
        ["FR-E-09", "Lab In-charge / Super Admin verify equipment (PENDING -> ACTIVE)", "MUST"],
    ])

    add_heading(doc, "9.3 FR-JR — Job Requests", level=2)
    add_table(doc, ["FR ID", "Description", "Priority"], [
        ["FR-JR-01", "List job requests with filters (type, status, mine vs all)", "MUST"],
        ["FR-JR-02", "Create job request form (multi-section, progressive disclosure)", "MUST"],
        ["FR-JR-03", "Save as draft", "MUST"],
        ["FR-JR-04", "Submit triggers state transition to SUBMITTED", "MUST"],
        ["FR-JR-05", "View request detail with full info + state history", "MUST"],
        ["FR-JR-06", "Lab In-charge can approve / reject with reason", "MUST"],
        ["FR-JR-07", "Approval triggers automatic Job Card creation", "MUST"],
        ["FR-JR-08", "Assignment dropdown shows lab engineers with current workload", "SHOULD"],
    ])

    add_heading(doc, "9.4 FR-JC — Job Cards", level=2)
    add_table(doc, ["FR ID", "Description", "Priority"], [
        ["FR-JC-01", "List job cards with filter (status, assigned-to-me, all)", "MUST"],
        ["FR-JC-02", "View job card with horizontal status stepper", "MUST"],
        ["FR-JC-03", "Task checklist (configurable per job type)", "MUST"],
        ["FR-JC-04", "Observations & readings log (free text + structured)", "MUST"],
        ["FR-JC-05", "Engineer can transition state (Start work, Mark complete)", "MUST"],
        ["FR-JC-06", "Lab In-charge can Verify/Close or Reopen", "MUST"],
        ["FR-JC-07", "Generate PDF of job card (on-demand download)", "MUST"],
        ["FR-JC-08", "State history visible as timeline", "MUST"],
    ])

    add_heading(doc, "9.5 FR-D — Dashboard", level=2)
    add_table(doc, ["FR ID", "Description", "Priority"], [
        ["FR-D-01", "Role-aware KPI widgets (different per role)", "MUST"],
        ["FR-D-02", "Calibration due alerts (next 30 days + overdue)", "MUST"],
        ["FR-D-03", "Engineer workload bar chart (for Lab In-charge)", "MUST"],
        ["FR-D-04", "Equipment status pie chart", "MUST"],
        ["FR-D-05", "Recently updated jobs table", "MUST"],
        ["FR-D-06", "Quick actions: Raise request, Add equipment (permission-gated)", "MUST"],
    ])

    add_heading(doc, "9.6 FR-I — Inquiry", level=2)
    add_table(doc, ["FR ID", "Description", "Priority"], [
        ["FR-I-01", "4-tab interface (Vendor, Product, Job Card, Instrument)", "MUST"],
        ["FR-I-02", "Real-time search with debounce", "MUST"],
        ["FR-I-03", "Result tables with drill-down to full record", "MUST"],
        ["FR-I-04", "Empty-state messages when no results", "MUST"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 10 — NFRs
    # =====================================================================
    add_section_banner(doc, "Section 10 — Non-Functional Requirements")
    add_table(doc, ["Category", "Requirement", "Target"], [
        ["Performance", "Initial page load (cold)", "< 3 sec on intranet"],
        ["Performance", "API response (95th percentile)", "< 500 ms"],
        ["Performance", "List pagination", "25 default, 100 max"],
        ["Security", "Password hashing", "bcryptjs ≥ 10 rounds"],
        ["Security", "JWT access lifetime", "15 min"],
        ["Security", "Refresh token lifetime", "7 days, httpOnly cookie, SameSite=Lax"],
        ["Security", "SQL injection defence", "Parameterized queries everywhere (mysql2 placeholders)"],
        ["Security", "XSS", "React auto-escape + helmet CSP header"],
        ["Security", "CSRF", "SameSite cookies + double-submit CSRF token on /auth/refresh"],
        ["Reliability", "Server uptime (post go-live)", "99% during business hours"],
        ["Reliability", "DB connection pool", "10–20 connections"],
        ["Reliability", "Process supervision", "PM2 cluster mode + auto-restart"],
        ["Usability", "Max clicks to any feature", "≤ 3"],
        ["Usability", "Keyboard shortcuts", "Power-user screens"],
        ["Usability", "Responsive viewport", "1280–1920 primary; graceful 768"],
        ["Audit", "Every state-changing op logged", "Yes"],
        ["Audit", "Audit log retention", "Indefinite (until specified)"],
        ["Maintainability", "Code style", "ESLint + Prettier via husky pre-commit"],
        ["Maintainability", "Folder structure (BE)", "routes / controllers / services / repositories"],
        ["Maintainability", "Folder structure (FE)", "feature-based (per D8)"],
        ["Maintainability", "API versioning", "/api/v1/... from day 1"],
        ["Compatibility", "Browsers", "Chrome, Edge, Firefox (latest 2 versions)"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 11 — DATABASE STRATEGY
    # =====================================================================
    add_section_banner(doc, "Section 11 — Database Strategy & Approach")

    add_callout(doc,
        "REALITY: ~64 tables already exist with records. CMCMIS does NOT "
        "start blank. Phase 3 (the next step) reviews each existing table, "
        "classifies, and designs new tables alongside.",
        fill="FFF4E5"
    )

    add_heading(doc, "11.1 The Phase 3 Inventory Approach (Next)", level=2)
    add_mono_block(doc,
        "  STEP 1  Inventory existing schema (user provides tomorrow):\n"
        "          - mysqldump --no-data  OR  SHOW CREATE TABLE for all 64\n"
        "          - user-provided notes: definitely-used / probably-dead / unsure\n"
        "  STEP 2  Classify each existing table:\n"
        "          KEEP    | EXTEND | REFACTOR | DEPRECATE | NEW\n"
        "  STEP 3  Decide migration strategy:\n"
        "          - Add new tables alongside existing (low risk)\n"
        "          - Write data migrations for refactored tables\n"
        "          - Never drop in MVP - only deprecate\n"
        "  STEP 4  Lock final ERD (~70 tables target)\n"
        "  STEP 5  Write seed scripts (roles, permissions, super admins,\n"
        "          system_types T&ME/F&PE, divisions, lookup data)"
    )

    add_heading(doc, "11.2 Approximate Group Map (~70 tables)", level=2)
    add_table(doc, ["Group", "~Tables", "Examples"], [
        ["Identity & RBAC", "~10", "users, employees, roles, permissions, user_roles, role_permissions, refresh_tokens, login_audit, password_reset_tokens, sso_identity_mappings"],
        ["Job Requests", "~7", "job_requests, types, status_history, attachments, comments, approvals, priorities"],
        ["Job Cards", "~9", "job_cards, tasks, task_templates, observations, status_history, assignees, signatures, time_logs, attachments"],
        ["Equipment", "~11", "equipment, types, specifications, cal_history, maint_history, status_history, documents, warranty, amc, manufacturers, categories"],
        ["Schedule (P2)", "~4", "schedules, recurrence_rules, notifications, assignments"],
        ["Procurement + Vendors (P2)", "~10", "purchase_orders, po_items, status_history, approvals, spare_parts, stock_movements, vendors, contacts, categories, documents"],
        ["Master / Lookups", "~10", "divisions, departments, locations, labs, system_types, units, currencies, config_settings, numbering_sequences, holiday_calendar"],
        ["Documents", "~4", "documents, versions, categories, signatures"],
        ["System / Observability", "~7", "notifications, templates, audit_log, error_log, background_jobs, email_dispatch_log, system_announcements"],
    ])

    add_heading(doc, "11.3 Conventions Locked (Per D2)", level=2)
    add_locked_banner(doc, "D2: Raw SQL + Repository pattern. NO ORM. mysql2/promise pool throughout.")
    add_table(doc, ["Convention", "Rule"], [
        ["Engine", "InnoDB always (transactions, FK enforcement)"],
        ["Charset", "utf8mb4 + utf8mb4_unicode_ci"],
        ["Table names", "snake_case, plural"],
        ["Primary keys", "BIGINT UNSIGNED AUTO_INCREMENT named id"],
        ["Foreign keys", "<ref_table_singular>_id"],
        ["Timestamps", "Every table has created_at, updated_at (NOT NULL, defaults)"],
        ["Soft deletes", "is_deleted BOOLEAN + deleted_at on transactional tables"],
        ["Status fields", "ENUM only for tiny fixed lists; otherwise FK to lookup table"],
        ["Audit fields", "created_by, updated_by (FK -> users.id)"],
        ["Money fields", "DECIMAL(18,4) — never FLOAT"],
        ["Query style", "Parameterized via mysql2 ? placeholders; NEVER concat"],
        ["Transactions", "Required for any multi-table write (state transition + audit log)"],
    ])

    add_para(doc,
        "Note: deep schema work — final ERD, table-by-table classification, "
        "SQL DDL — happens in Phase 3 starting tomorrow with user's existing "
        "schema export.",
        italic=True
    )

    add_page_break(doc)

    # =====================================================================
    # SECTION 12 — LOCKED TECH STACK
    # =====================================================================
    add_section_banner(doc, "Section 12 — Locked Tech Stack (FE + BE + Additions)")
    add_locked_banner(doc, "All stack items + 6 additions confirmed by user on 2026-05-16")

    add_heading(doc, "12.1 Frontend Stack (React 18 + Vite + Tailwind v3)", level=2)
    add_table(doc, ["#", "Concern", "Library", "Why"], [
        ["1", "Build", "vite", "Fast HMR, modern"],
        ["2", "Framework", "react@18", "Locked"],
        ["3", "Routing", "react-router-dom@v6", "SPA standard"],
        ["4", "State (local)", "useState / useReducer", "Built-in"],
        ["5", "State (global UI)", "zustand (D4)", "~1KB, no boilerplate"],
        ["6", "Server state", "@tanstack/react-query (D3)", "Cache, refetch, mutations"],
        ["7", "HTTP", "axios", "Interceptors for auth/refresh"],
        ["8", "Forms", "react-hook-form + zod", "Performant + schema shared with BE"],
        ["9", "UI primitives", "Custom on Tailwind", "Avoids shadcn/MUI lock-in"],
        ["10", "Icons", "lucide-react", "Tree-shakable"],
        ["11", "Charts", "recharts", "Dashboard widgets"],
        ["12", "Date", "dayjs", "Symmetric with backend"],
        ["13", "Toasts", "sonner", "Lightweight"],
        ["14", "Styling", "tailwindcss@v3", "Locked"],
        ["15", "Table", "@tanstack/react-table (ADDED)", "Headless tables for equipment/job lists"],
        ["16", "Form defaults", "@tailwindcss/forms (ADDED)", "Sane form-element defaults"],
        ["17", "Lint/Format", "eslint + prettier + husky + lint-staged", "Pre-commit enforced"],
        ["18", "Testing (FE)", "vitest (ADDED)", "Component + hook tests"],
    ])

    add_heading(doc, "12.2 Backend Stack (Node + Express 4)", level=2)
    add_table(doc, ["#", "Concern", "Library", "Ver", "Why"], [
        ["1", "Web framework", "express", "^4.x", "Stable, ubiquitous"],
        ["2", "DB driver", "mysql2/promise", "^3.x", "Fastest MySQL driver"],
        ["3", "Validation", "zod", "^3.x", "Single source of truth (FE+BE share)"],
        ["4", "Auth tokens", "jsonwebtoken", "^9.x", "Industry standard"],
        ["5", "Password hash", "bcryptjs", "^2.x", "Pure-JS, no native build pain (Windows)"],
        ["6", "Logger", "pino + pino-pretty (D6)", "^8.x", "Faster than Winston, structured JSON"],
        ["7", "Env loader", "dotenv", "^16.x", "Standard"],
        ["8", "Env validation", "envalid", "^8.x", "Fail at boot if env missing"],
        ["9", "Date", "dayjs", "^1.x", "Symmetric with frontend"],
        ["10", "PDF", "pdfkit (D5)", "^0.14.x", "Programmatic, no Chromium"],
        ["11", "Rate limit", "express-rate-limit", "^7.x", "In-memory, sufficient v1"],
        ["12", "Security", "helmet", "^7.x", "OWASP-friendly defaults"],
        ["13", "CORS", "cors", "^2.x", "Standard"],
        ["14", "Process mgr", "pm2", "latest", "Cluster mode, restarts, logs"],
        ["15", "Cookie", "cookie-parser (ADDED)", "^1.x", "Parse httpOnly refresh cookie"],
        ["16", "Compression", "compression (ADDED)", "^1.x", "gzip/brotli for API"],
        ["17", "CSRF", "double-submit token (ADDED)", "custom", "On /api/v1/auth/refresh"],
        ["18", "Lint/Format", "eslint + prettier + husky + lint-staged", "latest", "Pre-commit enforced"],
        ["19", "Testing (BE)", "vitest + supertest (ADDED)", "latest", "Unit + API tests"],
    ])

    add_heading(doc, "12.3 The Three Big Architectural Wins", level=2)
    add_bullets(doc, [
        "Schema symmetry — zod on BOTH frontend and backend. One definition, two enforcers.",
        "Date symmetry — dayjs on BOTH sides. No timezone/format drift on BR-EQP-05 math.",
        "On-prem friendly — pdfkit (no Chromium), pm2 (no systemd quirks), no Redis/SMTP/object-storage dependencies.",
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 13 — DECISION REGISTER
    # =====================================================================
    add_section_banner(doc, "Section 13 — Architectural Decision Register (D1–D11, C1–C5)")
    add_locked_banner(doc, "All decisions LOCKED by user on 2026-05-16. Do not re-open.")

    add_heading(doc, "13.1 Major Decisions (D-series)", level=2)
    add_table(doc, ["ID", "Decision", "Choice", "Lock"], [
        ["D1", "Type system", "JavaScript + JSDoc + Zod (NOT TypeScript)", "LOCKED"],
        ["D2", "Data access", "Raw SQL + Repository pattern (NOT ORM)", "LOCKED"],
        ["D3", "Server state (FE)", "TanStack Query", "LOCKED"],
        ["D4", "Global state (FE)", "Zustand (NOT Redux)", "LOCKED"],
        ["D5", "PDF generation", "pdfkit (NOT Puppeteer)", "LOCKED"],
        ["D6", "Logging", "Pino (NOT Winston)", "LOCKED"],
        ["D7", "BE layering", "routes -> controllers -> services -> repositories", "LOCKED"],
        ["D8", "FE folder org", "Feature-based (NOT layer-based)", "LOCKED"],
        ["D9", "Production proxy", "Nginx reverse proxy + TLS termination", "LOCKED"],
        ["D10", "New equipment", "Defaults to PENDING_VERIFICATION", "LOCKED"],
        ["D11", "Super Admin seed", "≥ 2 via SUPER_ADMIN_EMPLOYEE_IDS env var", "LOCKED"],
    ])

    add_heading(doc, "13.2 Confirmation Points (C-series)", level=2)
    add_table(doc, ["ID", "Confirmation", "Lock"], [
        ["C1", "5 roles final (Super Admin, Lab In-charge, Lab Engineer, Normal, View-Only)", "LOCKED"],
        ["C2", "Master Data Management = Phase 2; Super Admin only when built", "LOCKED"],
        ["C3", "Lookup data seeded Week 2; edited via phpMyAdmin during MVP", "LOCKED"],
        ["C4", "Bootstrap with ≥ 2 Super Admin employee IDs (user provides)", "LOCKED"],
        ["C5", "Equipment verify (PENDING -> ACTIVE) by Lab In-charge + Super Admin", "LOCKED"],
    ])

    add_heading(doc, "13.3 Pending Inputs from User (Tomorrow)", level=2)
    add_table(doc, ["#", "Input", "Use"], [
        ["1", "2 Super Admin employee IDs", "Goes into SUPER_ADMIN_EMPLOYEE_IDS env; seed migration"],
        ["2", "Existing DB schema export (mysqldump --no-data or SHOW CREATE TABLE x 64)", "Phase 3 schema inventory"],
        ["3", "Note on existing tables: used / dead / unsure", "Classification step in Phase 3"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 14 — CODE ARCHITECTURE PATTERNS
    # =====================================================================
    add_section_banner(doc, "Section 14 — Code Architecture Patterns")

    add_heading(doc, "14.1 JavaScript + JSDoc + Zod (D1)", level=2)
    add_para(doc, "Type system without TypeScript compilation overhead. "
                  "Zod schemas are the runtime contract; JSDoc + z.infer give "
                  "IDE-time hints in VSCode.")
    add_mono_block(doc,
        "  // backend/src/modules/equipment/equipment.schema.js\n"
        "  import { z } from 'zod';\n"
        "\n"
        "  export const EquipmentCreateSchema = z.object({\n"
        "    serial_no:        z.string().min(1).max(64),\n"
        "    model:            z.string().min(1).max(128),\n"
        "    manufacturer_id:  z.number().int().positive(),\n"
        "    system_type:      z.enum(['TME', 'FPE']),\n"
        "    cal_freq_months:  z.number().int().positive().nullable(),\n"
        "    division_id:      z.number().int().positive(),\n"
        "  });\n"
        "\n"
        "  /** @typedef {z.infer<typeof EquipmentCreateSchema>} EquipmentCreateInput */\n"
        "\n"
        "  // Use the type alias for IDE hints; no .ts compilation needed.\n"
        "  /** @param {EquipmentCreateInput} input */\n"
        "  export async function createEquipment(input) { ... }"
    )
    add_bullets(doc, [
        "Validation: every controller calls schema.parse(req.body) — fails fast on invalid input.",
        "Frontend: same schema file is imported by FE; react-hook-form uses zodResolver(schema).",
        "Result: one schema, two enforcers — perfect symmetry.",
    ])

    add_heading(doc, "14.2 Repository Pattern (D2 + D7)", level=2)
    add_para(doc, "Each layer has a single responsibility. No layer skips down to a lower layer.")
    add_mono_block(doc,
        "  HTTP request\n"
        "      |\n"
        "      v\n"
        "  +-------------------------------+\n"
        "  | ROUTE                         |  one-liner: router.post('/', ctrl.create)\n"
        "  +-------------------------------+\n"
        "      |\n"
        "      v\n"
        "  +-------------------------------+\n"
        "  | CONTROLLER                    |  HTTP layer:\n"
        "  |                               |   - parse req\n"
        "  |                               |   - validate with Zod\n"
        "  |                               |   - call service\n"
        "  |                               |   - format response + status\n"
        "  +-------------------------------+\n"
        "      |\n"
        "      v\n"
        "  +-------------------------------+\n"
        "  | SERVICE                       |  Business logic:\n"
        "  |                               |   - run state machine\n"
        "  |                               |   - orchestrate multiple repos\n"
        "  |                               |   - manage transactions\n"
        "  |                               |   - write audit_log\n"
        "  +-------------------------------+\n"
        "      |\n"
        "      v\n"
        "  +-------------------------------+\n"
        "  | REPOSITORY                    |  Data access ONLY:\n"
        "  |                               |   - parameterized SQL via mysql2\n"
        "  |                               |   - returns plain rows / objects\n"
        "  |                               |   - NO business rules here\n"
        "  +-------------------------------+\n"
        "      |\n"
        "      v\n"
        "    MySQL"
    )
    add_para(doc, "Why: easy to mock repositories in tests; easy to swap raw SQL "
                  "for batched SQL; business logic isolated from data access.",
             italic=True)

    add_heading(doc, "14.3 Feature-Based Frontend Folders (D8)", level=2)
    add_mono_block(doc,
        "  frontend/src/features/\n"
        "    equipment/\n"
        "      api/          (RQ hooks: useEquipmentList, useEquipmentCreate)\n"
        "      components/   (EquipmentTable, EquipmentForm, StatusBadge)\n"
        "      pages/        (EquipmentListPage, EquipmentDetailPage)\n"
        "      schemas/      (zod schemas — shared from backend)\n"
        "      hooks/        (useEquipmentStateTransition, ...)\n"
        "      types.js      (JSDoc typedefs via z.infer)\n"
        "      index.js      (public API of this feature)\n"
        "    jobRequests/\n"
        "    jobCards/\n"
        "    auth/\n"
        "    dashboard/\n"
        "    inquiry/"
    )
    add_para(doc, "Why over layer-based: every change to a feature lives in ONE folder. "
                  "Adding a Job Request field doesn't require touching 6 layer folders.",
             italic=True)

    add_page_break(doc)

    # =====================================================================
    # SECTION 15 — SYSTEM ARCHITECTURE & REQUEST TRACE
    # =====================================================================
    add_section_banner(doc, "Section 15 — System Architecture & Request Trace")

    add_heading(doc, "15.1 Runtime Architecture", level=2)
    add_mono_block(doc,
        "  +---------------------------------------------------------------+\n"
        "  | DESKTOP / LAPTOP BROWSER (NO MOBILE)                          |\n"
        "  | React 18 + Vite build + Tailwind                              |\n"
        "  |   Routes -> React Router                                      |\n"
        "  |   Server state -> TanStack Query  (D3)                        |\n"
        "  |   Global UI state -> Zustand      (D4)                        |\n"
        "  |   Forms -> React Hook Form + Zod                              |\n"
        "  |   HTTP -> axios (interceptor attaches JWT)                    |\n"
        "  +---------------------------+-----------------------------------+\n"
        "                              | HTTPS (JWT + httpOnly refresh cookie)\n"
        "                              v\n"
        "  +---------------------------------------------------------------+\n"
        "  | NGINX REVERSE PROXY  (D9, prod only)                          |\n"
        "  |  - TLS termination                                            |\n"
        "  |  - Serves static SPA (dist/)                                  |\n"
        "  |  - proxy_pass /api/*  ->  http://127.0.0.1:3000               |\n"
        "  |  - gzip / brotli                                              |\n"
        "  +---------------------------+-----------------------------------+\n"
        "                              v\n"
        "  +---------------------------------------------------------------+\n"
        "  | PM2 CLUSTER: NODE + EXPRESS API SERVER (D9)                   |\n"
        "  |  Middleware: helmet > cors > compression > json > cookies     |\n"
        "  |              > pino-http > rate-limit > authenticate          |\n"
        "  |              > authorize > rowLevelScope > validate(zod)      |\n"
        "  |              > controller                                     |\n"
        "  |  Service layer: state-machine transition() (single chokepoint)|\n"
        "  |                  pdfkit cert generator        (D5)            |\n"
        "  |                  audit-log helper                             |\n"
        "  |  Repository layer: mysql2/promise pool, parameterized SQL (D2)|\n"
        "  +---------------------------+-----------------------------------+\n"
        "                              v\n"
        "  +---------------------------------------------------------------+\n"
        "  | MySQL 8.x (existing on-prem; phpMyAdmin admin)                |\n"
        "  |   ~70 tables (existing 64 + new) · InnoDB · utf8mb4           |\n"
        "  |   FK + CHECK constraints · indexes for dashboard / inquiry    |\n"
        "  +---------------------------------------------------------------+\n"
        "\n"
        "  NOT IN MVP:  Redis  ·  SMTP  ·  S3/Object store  ·  SSO/AD"
    )

    add_heading(doc, "15.2 End-to-End Trace: Equipment Registration", level=2)
    add_mono_block(doc,
        "  Lab Engineer fills equipment register form\n"
        "    -> RHF + zod validates on submit (no network if invalid)\n"
        "    -> RQ mutation calls axios.post('/api/v1/equipment', body)\n"
        "    -> axios interceptor attaches Bearer JWT\n"
        "    ----------- NETWORK BOUNDARY ------------\n"
        "    -> nginx (TLS, static, proxy)  -> Node @ 3000\n"
        "    -> helmet > cors > compression > express.json > cookie-parser\n"
        "    -> pino logs request meta\n"
        "    -> express-rate-limit (auth-route bucket only)\n"
        "    -> authenticate (jsonwebtoken verify)\n"
        "    -> authorize ('equipment:create' permission check)\n"
        "    -> rowLevelScope (no-op for create)\n"
        "    -> validate(EquipmentCreateSchema) re-checks body\n"
        "    -> equipmentController.create(req, res)\n"
        "    -> equipmentService.create(input, actor):\n"
        "          - BEGIN TRANSACTION\n"
        "          - dayjs.add for next_cal_due_date (if T&ME)\n"
        "          - state machine: initial status = PENDING_VERIFICATION (D10)\n"
        "          - equipmentRepository.insert(payload)\n"
        "          - auditLogRepository.write({...})\n"
        "          - COMMIT\n"
        "    -> 201 Created + JSON body\n"
        "    -> pino logs response status + duration\n"
        "    ----------- NETWORK BOUNDARY ------------\n"
        "    -> RQ invalidates 'equipment-list' cache\n"
        "    -> sonner: 'Equipment registered. Pending verification.'\n"
        "    -> router navigates to detail page"
    )

    add_page_break(doc)

    # =====================================================================
    # SECTION 16 — MIDDLEWARE PIPELINE
    # =====================================================================
    add_section_banner(doc, "Section 16 — Express Middleware Pipeline")
    add_mono_block(doc,
        "  incoming request\n"
        "         |\n"
        "         v\n"
        "  +----------------------------------------------------+\n"
        "  | 1.  helmet()              <- security headers      |\n"
        "  | 2.  cors(orgAllowlist)    <- reject foreign origins|\n"
        "  | 3.  compression()         <- gzip / brotli         |\n"
        "  | 4.  express.json({limit}) <- parse JSON body       |\n"
        "  | 5.  cookie-parser()       <- parse refresh cookie  |\n"
        "  | 6.  pino-http()           <- structured log line   |\n"
        "  | 7.  rateLimit(login)      <- only on auth routes   |\n"
        "  | 8.  csrfCheck()           <- /auth/refresh only    |\n"
        "  | 9.  authenticate()        <- JWT verify            |\n"
        "  | 10. authorize(perm)       <- RBAC check            |\n"
        "  | 11. rowLevelScope()       <- BR-VIS WHERE clause   |\n"
        "  | 12. validate(zodSchema)   <- input validation      |\n"
        "  | 13. controller            <- business handler      |\n"
        "  | 14. errorHandler()        <- centralized errors    |\n"
        "  +----------------------------------------------------+\n"
        "         |\n"
        "         v\n"
        "  outgoing response"
    )
    add_para(doc, "Order matters. Helmet/cors/compression run for ALL requests. "
                  "CSRF check only fires on /auth/refresh (which uses cookies). "
                  "RBAC + scoping happen only after authenticate.", italic=True)

    add_page_break(doc)

    # =====================================================================
    # SECTION 17 — PRODUCTION DEPLOYMENT TOPOLOGY
    # =====================================================================
    add_section_banner(doc, "Section 17 — Production Deployment Topology")
    add_locked_banner(doc, "D9: Nginx reverse proxy in production")

    add_heading(doc, "17.1 The Topology", level=2)
    add_mono_block(doc,
        "                Org intranet (no public exposure)\n"
        "                          |\n"
        "                          | HTTPS :443\n"
        "                          v\n"
        "  +----------------------------------------------------------+\n"
        "  | NGINX  (on the same VM, port 80/443)                     |\n"
        "  |   - TLS termination  (org-signed cert)                   |\n"
        "  |   - Serves /            ->  /var/www/cmcmis-spa/dist     |\n"
        "  |   - Proxies /api/*      ->  http://127.0.0.1:3000        |\n"
        "  |   - gzip + brotli                                        |\n"
        "  |   - cache headers for static assets (hashed filenames)   |\n"
        "  |   - access logs                                          |\n"
        "  +----------------------+-----------------------------------+\n"
        "                         |\n"
        "                         v\n"
        "  +----------------------------------------------------------+\n"
        "  | PM2 CLUSTER  (node, port 3000 — cluster on N workers)    |\n"
        "  |   ecosystem.config.js:                                   |\n"
        "  |     - instances: 'max'  (one per CPU core)               |\n"
        "  |     - exec_mode: 'cluster'                               |\n"
        "  |     - autorestart: true                                  |\n"
        "  |     - max_memory_restart: '500M'                         |\n"
        "  |     - log rotation                                       |\n"
        "  +----------------------+-----------------------------------+\n"
        "                         |\n"
        "                         v\n"
        "  +----------------------------------------------------------+\n"
        "  | MySQL 8.x  (port 3306, on-prem)                          |\n"
        "  |   InnoDB · utf8mb4 · pooled connections (10-20)          |\n"
        "  +----------------------------------------------------------+"
    )

    add_heading(doc, "17.2 Nginx Responsibilities", level=2)
    add_bullets(doc, [
        "Terminate TLS with org-signed cert.",
        "Serve Vite static build (`dist/`) for any non-/api/* path.",
        "Reverse-proxy /api/* to Node:3000 (preserves Authorization header + cookies).",
        "Compress responses (gzip/brotli).",
        "Add immutable cache headers for hashed asset filenames.",
        "Optional second-layer rate limit (e.g., 100 req/min/IP) on top of express-rate-limit.",
        "Forward client IP via X-Forwarded-For (Express trust proxy = 1).",
    ])

    add_heading(doc, "17.3 PM2 Responsibilities", level=2)
    add_bullets(doc, [
        "Cluster mode across CPU cores for throughput.",
        "Auto-restart on crash; respawn within seconds.",
        "Memory-based restart guard (kill worker if RAM > 500MB).",
        "Centralized log rotation (pm2-logrotate).",
        "Zero-downtime reload on deploy (pm2 reload all).",
    ])

    add_heading(doc, "17.4 Deploy Checklist (Phase 8)", level=2)
    add_bullets(doc, [
        "Build FE: `vite build` → dist/.",
        "Copy dist/ to /var/www/cmcmis-spa/dist on the server.",
        "Install BE: `npm ci --omit=dev`.",
        "Validate env via envalid (boot fails if missing).",
        "Run seed migration (creates Super Admins from SUPER_ADMIN_EMPLOYEE_IDS).",
        "pm2 start ecosystem.config.js.",
        "Nginx reload.",
        "Smoke test: hit /api/v1/health, login as one of the seeded Super Admins, force password change.",
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 18 — OPERATIONAL STRATEGIES
    # =====================================================================
    add_section_banner(doc, "Section 18 — Operational Strategies")

    add_heading(doc, "18.1 PDF Generation Strategy (D5 — pdfkit)", level=2)
    add_bullets(doc, [
        "Templates as pure functions: (data) => Buffer.",
        "Two templates in MVP: jobCard.pdf.js + calibrationCert.pdf.js (per BR-PDF-02 / BR-PDF-03).",
        "Generated on GET /.../<id>/pdf — streamed back; nothing persisted (BR-PDF-01).",
        "Every PDF includes footer: 'Generated by CMCMIS · <timestamp> · Record #<id>' (BR-PDF-04).",
        "Per-export audit_log row written before the stream begins (BR-AUD-03).",
    ])

    add_heading(doc, "18.2 Logging Strategy (D6 — Pino)", level=2)
    add_bullets(doc, [
        "Pino emits structured JSON; logs ingested as files via pm2 log rotation.",
        "Levels: trace / debug / info / warn / error / fatal. Default INFO in prod, DEBUG in dev.",
        "Request log line: method, path, status, duration_ms, user_id, request_id.",
        "Business events (state transitions, login attempts) logged at INFO with `event` key.",
        "Errors logged with stack + request_id; never console.log anywhere.",
        "Sensitive fields (password, token, refresh) redacted via pino redact.",
    ])

    add_heading(doc, "18.3 Testing Strategy (vitest + supertest)", level=2)
    add_table(doc, ["Layer", "Tool", "What to Test"], [
        ["Unit (BE)", "vitest", "State-machine transition() functions; repository SQL builders; auth helpers"],
        ["API (BE)", "vitest + supertest", "Each endpoint × each of 5 roles; happy path + 401/403/422/409"],
        ["Unit (FE)", "vitest", "Permission gate component; form schemas; date formatters"],
        ["Component (FE)", "vitest + React Testing Library", "Critical UI: login, equipment form, job card stepper"],
        ["E2E", "Manual + scripted smoke", "Phase 9 hardening: full DMM-2034 walk-through across roles"],
    ])
    add_para(doc, "Coverage target for MVP: state machines + RBAC = 100%. Everything "
                  "else: best-effort, no minimum gate.", italic=True)

    add_heading(doc, "18.4 CSRF Strategy (Added Stack Item #3)", level=2)
    add_para(doc, "Standard JWT in Authorization header is CSRF-safe (no cookie). "
                  "The ONLY endpoint at risk is /api/v1/auth/refresh — it reads the "
                  "httpOnly refresh cookie. Strategy:")
    add_bullets(doc, [
        "On login, in addition to the refresh cookie, set a non-httpOnly cookie 'csrf_token' with a random value.",
        "Frontend reads csrf_token and echoes it in X-CSRF-Token header on /auth/refresh.",
        "Server middleware compares cookie value vs header value (double-submit). Mismatch = 403.",
        "Same-site=Lax on refresh cookie blocks most CSRF anyway; this is defence in depth.",
    ])

    add_heading(doc, "18.5 Security Strategy (Cross-Cutting)", level=2)
    add_table(doc, ["Concern", "Mitigation"], [
        ["SQL injection", "Parameterized queries only (mysql2 ? placeholders). No string concat. Repository layer enforces."],
        ["XSS", "React auto-escape + helmet CSP header. No dangerouslySetInnerHTML in MVP."],
        ["CSRF", "Double-submit token on /auth/refresh (see 18.4)."],
        ["Brute-force login", "express-rate-limit on /auth/login + is_locked column after N failures."],
        ["Token theft", "15-min access JWT + httpOnly Secure SameSite refresh cookie + hashed in DB."],
        ["Privilege escalation", "Every endpoint requires explicit permission. Default-deny."],
        ["Data leakage", "BR-AUD-04: sensitive fields stripped from API responses by permission."],
        ["Audit gap", "BR-AUD-01: every write to critical tables logged in audit_logs."],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 19 — FOLDER STRUCTURE
    # =====================================================================
    add_section_banner(doc, "Section 19 — Project Folder Structure (FINAL)")

    add_mono_block(doc,
        "  cmcmis-simplified/\n"
        "  |\n"
        "  +-- backend/                                    (BE: D2 + D7)\n"
        "  |   +-- src/\n"
        "  |   |   +-- config/        (db.js, env.js, logger.js, jwt.js)\n"
        "  |   |   +-- middleware/    (authenticate, authorize, rowLevelScope,\n"
        "  |   |   |                   validate, csrfCheck, errorHandler)\n"
        "  |   |   +-- modules/\n"
        "  |   |   |   +-- auth/\n"
        "  |   |   |   |   +-- auth.routes.js\n"
        "  |   |   |   |   +-- auth.controller.js\n"
        "  |   |   |   |   +-- auth.service.js\n"
        "  |   |   |   |   +-- auth.repository.js\n"
        "  |   |   |   |   +-- auth.schema.js     (zod)\n"
        "  |   |   |   +-- users/\n"
        "  |   |   |   +-- equipment/\n"
        "  |   |   |   +-- jobRequests/\n"
        "  |   |   |   +-- jobCards/\n"
        "  |   |   |   +-- dashboard/\n"
        "  |   |   |   +-- inquiry/\n"
        "  |   |   |   +-- audit/\n"
        "  |   |   +-- stateMachines/ (jobRequest.fsm.js, equipment.fsm.js)\n"
        "  |   |   +-- pdf/           (jobCard.pdf.js, calibrationCert.pdf.js)\n"
        "  |   |   +-- utils/\n"
        "  |   |   +-- server.js\n"
        "  |   +-- db/\n"
        "  |   |   +-- migrations/    (numbered .sql files)\n"
        "  |   |   +-- seeds/         (roles, permissions, super admins, lookups)\n"
        "  |   |   +-- schema.sql\n"
        "  |   +-- tests/             (vitest + supertest)\n"
        "  |   +-- .env.example\n"
        "  |   +-- ecosystem.config.js   (pm2)\n"
        "  |   +-- package.json\n"
        "  |\n"
        "  +-- frontend/                                   (FE: D8 feature-based)\n"
        "  |   +-- src/\n"
        "  |   |   +-- components/    (shared: Button, Input, Table, Modal)\n"
        "  |   |   +-- features/\n"
        "  |   |   |   +-- auth/         (api/, components/, pages/, schemas/)\n"
        "  |   |   |   +-- equipment/\n"
        "  |   |   |   +-- jobRequests/\n"
        "  |   |   |   +-- jobCards/\n"
        "  |   |   |   +-- dashboard/\n"
        "  |   |   |   +-- inquiry/\n"
        "  |   |   +-- layouts/       (AppLayout, AuthLayout)\n"
        "  |   |   +-- lib/           (axios client, authContext,\n"
        "  |   |   |                   permissions, formatters)\n"
        "  |   |   +-- stores/        (zustand stores)\n"
        "  |   |   +-- router.jsx\n"
        "  |   |   +-- main.jsx\n"
        "  |   +-- public/\n"
        "  |   +-- tailwind.config.js\n"
        "  |   +-- vite.config.js\n"
        "  |   +-- package.json\n"
        "  |\n"
        "  +-- docs/\n"
        "  |   +-- FINAL-DESC-CMCMIS.docx   (this document)\n"
        "  |   +-- ERD.md                   (after Phase 3)\n"
        "  |   +-- rbac-matrix.md\n"
        "  |   +-- state-machines.md\n"
        "  |   +-- business-rules.md\n"
        "  +-- .gitignore\n"
        "  +-- .editorconfig\n"
        "  +-- README.md\n"
        "  +-- package.json   (monorepo root w/ workspaces)"
    )

    add_page_break(doc)

    # =====================================================================
    # SECTION 20 — CODING CONVENTIONS
    # =====================================================================
    add_section_banner(doc, "Section 20 — Coding Conventions & API Standards")
    add_table(doc, ["Concern", "Convention"], [
        ["API base", "/api/v1/..."],
        ["HTTP verbs", "REST: GET / POST / PATCH / DELETE"],
        ["Response shape (success)", "{ data, meta }"],
        ["Response shape (error)", "{ error: { code, message, details } }"],
        ["Status codes used", "200 / 201 / 204 / 400 / 401 / 403 / 404 / 409 / 422 / 429 / 500"],
        ["Pagination", "?page=1&limit=25 (max 100)"],
        ["Filtering", "?status=ACTIVE&type=TME (whitelisted per endpoint)"],
        ["Sorting", "?sort=-created_at,name (- prefix = DESC)"],
        ["Naming (DB)", "snake_case, plural tables, id PKs, *_id FKs"],
        ["Naming (JS)", "camelCase vars/functions, PascalCase components/classes"],
        ["Naming (files)", "kebab-case.js for utilities; PascalCase.jsx for React components"],
        ["Errors", "Throw typed errors -> centralized handler -> standard response"],
        ["Logs", "Structured JSON via pino; never console.log"],
        ["Commits", "Conventional Commits (feat / fix / chore / docs / refactor)"],
        ["Branches", "feature/*, fix/*, chore/* from main"],
        ["Type system", "JS + JSDoc + Zod (NO TypeScript) — D1"],
        ["Backend layering", "routes -> controllers -> services -> repositories — D7"],
        ["Frontend layering", "Feature folders, each self-contained — D8"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 21 — MVP SCOPE
    # =====================================================================
    add_section_banner(doc, "Section 21 — MVP Scope vs Phase 2 Backlog")

    add_heading(doc, "21.1 MVP (10 weeks) — Signed Off", level=2)
    add_bullets(doc, [
        "Auth + RBAC (login, session, role loading, permissions, protected routes, SSO-ready)",
        "Equipment master + register + verify flow (PENDING_VERIFICATION -> ACTIVE)",
        "Job Requests (form, listing, lifecycle, approve/reject)",
        "Job Cards (tasks, observations, verify/close, reopen)",
        "Dashboard (role-aware KPIs, alerts, widgets)",
        "Inquiry (search hub — vendors, products, job cards, instruments)",
        "PDF generation (download/generate only — no storage)",
        "Audit logs (basic, all state changes)",
        "Responsive UI (desktop + laptop; 1280–1920 primary)",
        "Row-level visibility (per BR-VIS rules)",
        "Optimized SQL with pagination + connection pool",
    ])

    add_heading(doc, "21.2 Phase 2 — Post-Internship Backlog", level=2)
    add_bullets(doc, [
        "Schedule module (PM + Calibration calendar)",
        "Procurement (POs + spares)",
        "Reports (analytics + exports)",
        "Admin master-data CRUD UI (Super Admin only)",
        "Notifications (in-app feed)",
    ])

    add_heading(doc, "21.3 Explicitly NOT in MVP", level=2)
    add_table(doc, ["Item", "Decision"], [
        ["SSO / Active Directory integration", "Architecture SSO-ready, integration deferred"],
        ["Email / SMTP", "Out of scope until specified"],
        ["Redis / caching layer", "MySQL + pagination + pool is enough"],
        ["File storage", "PDFs generated + downloaded only, never stored"],
        ["Backup infrastructure", "Out of our scope"],
        ["Mobile / tablet UI", "Desktop + laptop only (responsive within range)"],
        ["Barcode / QR", "Not in scope"],
        ["NABL / ISO 17025 / AS9100 specifics", "Deferred — user will instruct"],
        ["Calibration cert template format", "Deferred — user will instruct"],
        ["Notification channels (email/in-app/SMS/push)", "Deferred — user will instruct"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 22 — 10-WEEK WAR PLAN
    # =====================================================================
    add_section_banner(doc, "Section 22 — 10-Week War Plan (Phased Timeline)")
    add_table(doc, ["Week", "Phase", "Deliverables"], [
        ["1", "Phase 0 — Foundation", "Repo scaffolding · ESLint/Prettier · Tailwind base · MySQL pool · 'hello world' protected route end-to-end"],
        ["2", "Phase 3 — DB Design", "Inventory ~64 existing tables · Classify · Draft new tables (RBAC + Job + Equipment first) · Seed scripts · Lock ERD"],
        ["3", "Block 1 — Auth + RBAC", "employee_id + password login · JWT + refresh · cookie-parser + CSRF · RBAC middleware · Sidebar visibility · Row-level scoping helper · Audit-log helper"],
        ["4–5", "Block 2a — Equipment Master", "Equipment list/detail/register · State machine (incl. PENDING_VERIFICATION) · Cal history · Specs · Verify by In-charge/SA"],
        ["5–6", "Block 2b — Job Requests", "Multi-section form (T&ME/F&PE; cal/repair/registration) · Draft/submit · Approve/reject · State machine"],
        ["6–7", "Block 2c — Job Cards + PDF", "Auto-create on approval · Task checklist + observations · Verify/close -> Equipment status update · PDF cert generation"],
        ["8", "Block 3 — Dashboard + Inquiry", "Dashboard widgets (KPIs, due alerts, workload) · Inquiry 4-tab search"],
        ["9", "Hardening", "Audit coverage check · RBAC e2e test (all 5 roles) · Illegal-transition tests · SQL tuning · Responsive QA · CSRF test"],
        ["10", "Demo Prep + Deploy", "Demo data seeded · Nginx + PM2 cluster set up · On-prem deploy · Stakeholder demo · Bug-fix buffer"],
    ])

    add_heading(doc, "22.1 Buffer Strategy (Slack Plan)", level=2)
    add_table(doc, ["Risk", "Buffer Plan"], [
        ["Existing 64-table DB messier than expected", "Phase 3 stretches into Week 3; Auth gets compressed"],
        ["Job Card observations more complex than planned", "Cut Inquiry tabs to 2 (job cards + equipment) for demo"],
        ["PDF formatting takes longer", "Ship plain-template PDFs; polish post-internship"],
        ["RBAC edge cases discovered", "Core — never cut. Cut Dashboard widgets instead."],
        ["Nginx setup on-prem unfamiliar", "Use docker-compose fallback for demo; raw nginx after"],
    ])

    add_page_break(doc)

    # =====================================================================
    # SECTION 23 — WALK-THROUGH
    # =====================================================================
    add_section_banner(doc, "Section 23 — Real-World End-to-End Walk-Through")
    add_para(doc, "Tracing Digital Multimeter DMM-2034 through CMCMIS to make every decision concrete.",
             italic=True, color=GREY)

    add_table(doc, ["Step", "Actor", "Role", "Module", "Action"], [
        ["1", "Engineer Deep", "Lab Engineer", "Dashboard", "Sees DMM-2034 calibration due alert"],
        ["2", "Deep", "Lab Engineer", "Job Requests", "Raises calibration request"],
        ["3", "Lab In-charge", "Lab In-charge", "Job Requests", "Approves -> Job Card auto-created"],
        ["4", "Deep (auto-assigned)", "Lab Engineer", "Job Cards", "Completes tasks, enters before/after readings + environment, marks COMPLETE"],
        ["5", "Lab In-charge", "Lab In-charge", "Job Cards", "Verifies + closes -> triggers PDF cert generation (pdfkit)"],
        ["6", "Anyone (except View-Only)", "Any of 4 roles", "Equipment", "Could register new DMM-2035 tomorrow (PENDING_VERIFICATION)"],
        ["7", "Lab In-charge", "Lab In-charge", "Equipment", "Verifies DMM-2035 -> ACTIVE"],
        ["8", "Super Admin", "Super Admin", "Admin (seed/CLI in MVP)", "Promotes Normal User -> Lab Engineer"],
        ["9", "View-Only auditor", "View-Only", "Inquiry", "Searches DMM-2034 history; reads everything; writes nothing"],
        ["10", "Super Admin", "Super Admin", "Audit Log", "Reviews every state transition"],
    ])

    add_para(doc,
        "One instrument, one calibration cycle: ~20 DB writes across ~12 tables, "
        "1 PDF cert, ~9 audit rows. This is why ~70 tables is the MINIMUM to "
        "track this lawfully.",
        italic=True
    )

    add_page_break(doc)

    # =====================================================================
    # SECTION 24 — RECAP & WHAT'S NEXT
    # =====================================================================
    add_section_banner(doc, "Section 24 — Constraints, Quick Recap, What's Next")

    add_heading(doc, "24.1 Locked Constraints", level=2)
    add_table(doc, ["#", "Constraint", "Decision"], [
        ["1", "Organization context", "ISRO SAC-like defence/space-grade"],
        ["2", "Existing DB", "~64 tables; review in Phase 3; never drop in MVP"],
        ["3", "SSO in v1", "NO — employee_id login from existing DB; SSO-ready for future"],
        ["4", "Deployment target", "Internal on-prem / org private infra"],
        ["5", "File storage", "NO — PDFs generated + downloaded only"],
        ["6", "Email / SMTP", "NO in v1"],
        ["7", "Redis / cache", "NO — MySQL + pagination + pool"],
        ["8", "Backup infrastructure", "Out of scope (user handles separately)"],
        ["9", "Mobile / tablet UI", "NO — desktop/laptop only; fully responsive within range"],
        ["10", "Barcode / QR", "NO"],
        ["11", "PDF generation", "YES — pdfkit, server-side, download only"],
        ["12", "Sensitive data", "Defence-grade; row-level visibility + secure sessions + limited exports"],
    ])

    add_heading(doc, "24.2 Print-and-Pin Quick Recap", level=2)
    add_table(doc, ["Topic", "Locked Decision"], [
        ["Roles", "5 only — Super Admin, Lab In-charge, Lab Engineer, Normal, View-Only"],
        ["equipment:create", "Open to ALL roles except View-Only"],
        ["Equipment lifecycle", "Starts at PENDING_VERIFICATION -> ACTIVE (D10)"],
        ["Per-user role count", "Exactly ONE primary role (BR-RBAC-02)"],
        ["Super Admin seed", "≥ 2 via SUPER_ADMIN_EMPLOYEE_IDS env var (D11)"],
        ["Session model", "15-min JWT · 60-min idle · 7-day refresh (concentric)"],
        ["Type system", "JS + JSDoc + Zod (D1)"],
        ["Data access", "Raw SQL + Repository (D2 + D7)"],
        ["FE folders", "Feature-based (D8)"],
        ["FE state", "TanStack Query (D3) + Zustand (D4)"],
        ["BE logging", "Pino (D6)"],
        ["PDF", "pdfkit (D5) — generate + download, never store"],
        ["Production proxy", "Nginx + PM2 cluster + MySQL (D9)"],
        ["Audit retention", "Indefinite (until user specifies)"],
        ["API base", "/api/v1/... from day 1"],
        ["Timeline", "10 weeks · solo dev + AI pair · MVP first · full go-live after sign-off"],
        ["DB target", "~70 tables (existing 64 reviewed + new added)"],
        ["Bottom line", "Industry-grade defence context. Demo-ready MVP in 10 weeks. Josh: MAXED."],
    ])

    add_heading(doc, "24.3 What's Next (Tomorrow Morning)", level=2)
    add_callout(doc,
        "User will deliver:  (1) 2 Super Admin employee IDs.  "
        "(2) Existing DB schema export (mysqldump --no-data or SHOW CREATE TABLE x 64).  "
        "(3) Notes on which existing tables are used / dead / unsure.\n"
        "\n"
        "Next phase opens immediately: Phase 3 — DB Design. "
        "Existing schema review + final ERD lock + new tables drafted.",
        fill="E8F5E9"
    )

    add_locked_banner(
        doc,
        "END OF FINAL-DESC-CMCMIS  -  All decisions locked  -  "
        "Energy level: maxed  -  Discipline level: locked  -  Let's ship it."
    )

    return doc


if __name__ == "__main__":
    out = build()
    out_path = r"e:\SOFTWAREs By DS\cmcmis-simplified\Documents\FINAL-DESC-CMCMIS.docx"
    out.save(out_path)
    print(f"Saved: {out_path}")
