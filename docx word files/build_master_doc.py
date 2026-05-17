"""
CMCMIS Master Plan -> Word document generator.
Builds a comprehensive .docx capturing everything decided in the planning phase
so far, formatted in the user's preferred style (structured, table-heavy,
flowchart-rich, beginner-to-advanced friendly).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
#                  STYLING HELPERS
# ============================================================

NAVY = RGBColor(0x0B, 0x2A, 0x5B)
ACCENT = RGBColor(0x1E, 0x6F, 0xA8)
GREEN = RGBColor(0x1B, 0x7F, 0x3A)
RED = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x55, 0x55, 0x55)
TABLE_HEADER_FILL = "1E6FA8"
SECTION_FILL = "0B2A5B"


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size is None:
        size = {1: 22, 2: 17, 3: 14, 4: 12}.get(level, 12)
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_callout(doc, text, fill="EAF3FB"):
    """Single-cell coloured callout for emphasized takeaways."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.bold = True
    return tbl


def add_mono_block(doc, text):
    """Code/ASCII-diagram block in monospace, light grey background."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F4F4F4")
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Force Consolas in East-Asian glyph fallback too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return tbl


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, TABLE_HEADER_FILL)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    # Rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = tbl.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if val == "YES" or val == "Y":
                run.font.color.rgb = GREEN
                run.bold = True
            elif val == "NO" or val == "N":
                run.font.color.rgb = RED
                run.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)


def add_page_break(doc):
    doc.add_page_break()


def add_section_banner(doc, text):
    """Big section banner — coloured single-cell row."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, SECTION_FILL)
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


# ============================================================
#                  DOCUMENT BUILD
# ============================================================


def build_document():
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------------- COVER ----------------
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("\n\nCMCMIS")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Computerized Maintenance & Calibration\nManagement Information System")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\nMaster Planning Document — v3\n")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Project Definition · Module Architecture · Roles & RBAC\n"
        "State Machines · Business Rules · FR/NFR\n"
        "Tech Stack · System Architecture · Build Roadmap"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "\n\nPrepared collaboratively for the\nIndustry-Grade Production Build (10-Week Sprint)\n\n"
        "Style: structured · table-heavy · flowchart-rich · beginner-to-advanced\n"
    )
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = GREY

    add_page_break(doc)

    # ---------------- TABLE OF CONTENTS ----------------
    add_heading(doc, "Table of Contents", level=1)
    toc = [
        "Section 1 — Project at a Glance",
        "Section 2 — The 9-Module Architecture Map",
        "Section 3 — Role Model (5 Roles, Single-Tier Admin)",
        "Section 4 — Authorization & RBAC (3-Layer Model)",
        "Section 5 — Final Consolidated Permission Matrix",
        "Section 6 — Authentication Flow",
        "Section 7 — Critical State Machines",
        "Section 8 — Business Rules Catalogue",
        "Section 9 — Functional Requirements (MVP)",
        "Section 10 — Non-Functional Requirements",
        "Section 11 — Database Strategy & 70-Table Map",
        "Section 12 — Tech Stack (Frontend + Backend)",
        "Section 13 — System Architecture & Request Trace",
        "Section 14 — Express Middleware Pipeline",
        "Section 15 — Project Folder Structure",
        "Section 16 — Coding Conventions & API Standards",
        "Section 17 — MVP Scope vs Phase 2 Backlog",
        "Section 18 — 10-Week War Plan (Phased Timeline)",
        "Section 19 — Real-World End-to-End Walk-Through",
        "Section 20 — Constraints, Open Items, Quick Recap",
    ]
    for line in toc:
        add_para(doc, "•  " + line, size=11)

    add_page_break(doc)

    # ============================================================
    # SECTION 1 — PROJECT AT A GLANCE
    # ============================================================
    add_section_banner(doc, "Section 1 — Project at a Glance")

    add_heading(doc, "1.1 The One-Sentence Version (Simple)", level=2)
    add_callout(
        doc,
        "“CMCMIS is the system that knows where every instrument in our labs is, "
        "what state it’s in, when it needs calibration, who’s working on it, and "
        "what was done to it — for as long as the instrument exists.”",
    )

    add_heading(doc, "1.2 The Domain in 5 Words", level=2)
    add_mono_block(doc, "REGISTER  →  CALIBRATE  →  MAINTAIN  →  REPAIR  →  RETIRE")
    add_para(
        doc,
        "Every record in CMCMIS is one event on one instrument’s journey through these 5 phases.",
        italic=True,
    )

    add_heading(doc, "1.3 The Strategic Picture (Deep)", level=2)
    add_table(
        doc,
        ["Layer", "What Lives Here", "Why It Matters"],
        [
            ["People", "5 roles, hundreds of users", "Wrong access = data corruption / compliance breach"],
            ["Process", "2 strict state machines", "Wrong transitions = audit failures (ISRO-grade orgs cannot tolerate this)"],
            ["Assets", "Thousands of instruments + their full history", "Lose history = lose calibration provenance = unsafe results"],
            ["Money", "POs, spares, vendor contracts (Phase 2)", "Wrong PO chains = procurement audit issues"],
            ["Knowledge", "Reports, certificates, signatures", "This IS the institutional memory"],
        ],
    )

    add_heading(doc, "1.4 Engagement Context", level=2)
    add_table(
        doc,
        ["Aspect", "Decision"],
        [
            ["Organization context", "ISRO SAC-like government / technical / defence-grade environment"],
            ["Project status", "Real organizational production system (NOT a prototype)"],
            ["Owner", "Solo Software-Developer intern (with AI pair)"],
            ["Timeline", "10 weeks for MVP; full go-live after testing + stakeholder approval"],
            ["Deployment", "Internal on-prem / private infra (NOT public cloud)"],
            ["Existing system", "~64 tables already exist with records — reviewed in Phase 1"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 2 — MODULE ARCHITECTURE
    # ============================================================
    add_section_banner(doc, "Section 2 — The 9-Module Architecture Map")

    add_heading(doc, "2.1 The Big Picture", level=2)
    add_mono_block(
        doc,
        "                  +---------------------------------+\n"
        "                  |   DASHBOARD (read-only)         |\n"
        "                  |   Reads -> ALL modules          |\n"
        "                  +----------------+----------------+\n"
        "                                   |\n"
        "        +--------------------------+--------------------------+\n"
        "        v                          v                          v\n"
        "  +-------------+         +------------------+        +----------------+\n"
        "  | JOB         |--creates| JOB CARDS        |<-touch-| EQUIPMENT      |\n"
        "  | REQUESTS    |-------->|  (work execution)|        |  (asset master)|\n"
        "  | (intake)    |         |                  |        |                |\n"
        "  +-----+-------+         +-----+------------+        +--------+-------+\n"
        "        |                       | updates                      | schedules\n"
        "        |                       v                              v\n"
        "        |             +------------------+             +------------------+\n"
        "        |             | DOCUMENTS        |             | SCHEDULE  (P2)   |\n"
        "        |             | (PDF generation) |             | (PM + Cal cal)   |\n"
        "        |             +------------------+             +------------------+\n"
        "        v\n"
        "  +--------------+     +----------------+\n"
        "  | PROCUREMENT  |<--->| VENDORS  (P2)  |\n"
        "  |  (P2)        |     | (master data)  |\n"
        "  +--------------+     +----------------+\n"
        "\n"
        "  +-----------------+   +-----------------+   +------------------+\n"
        "  | INQUIRY         |   | REPORTS  (P2)   |   | ADMIN  (P2)      |\n"
        "  | (search hub)    |   | (analytics)     |   | (master + RBAC)  |\n"
        "  +-----------------+   +-----------------+   +------------------+\n"
        "\n"
        "  (P2) = Phase 2 / post-internship. NOT in MVP.",
    )

    add_heading(doc, "2.2 Module Responsibility Matrix", level=2)
    add_table(
        doc,
        ["#", "Module", "Owns (CRUD)", "Reads From", "Writes To", "Primary Actors", "MVP?"],
        [
            ["1", "Dashboard", "Nothing", "All modules", "Nothing", "All roles", "YES"],
            ["2", "Job Requests", "job_requests", "Equipment, Employees", "Job Cards (on approval)", "Normal User → Lab In-charge", "YES"],
            ["3", "Job Cards", "job_cards, tasks, observations", "Job Requests, Equipment", "Documents, Schedule, Equipment status", "Lab Engineer, Lab In-charge", "YES"],
            ["4", "Equipment", "equipment, specs, cal_history", "Job Cards, Procurement", "Schedule", "All except View-Only", "YES"],
            ["5", "Schedule", "schedules (PM + Cal)", "Equipment", "Notifications, auto-Job Requests", "Lab In-charge, Engineer", "Phase 2"],
            ["6", "Procurement", "purchase_orders, spare_parts", "Vendors, Equipment", "Equipment stock", "Lab In-charge", "Phase 2"],
            ["7", "Vendors", "vendors", "—", "Procurement", "Super Admin", "Phase 2"],
            ["8", "Inquiry", "Nothing (read-only)", "Everything", "—", "All roles", "YES"],
            ["9", "Reports", "Nothing (read-only)", "All modules", "Export logs", "Lab In-charge, Super Admin", "Phase 2"],
            ["10", "Admin", "Master data + Users + Roles", "—", "All master tables", "Super Admin only", "Phase 2 (RBAC actions live in MVP)"],
        ],
    )

    add_heading(doc, "2.3 The Golden Rule of Data Flow", level=2)
    add_mono_block(
        doc,
        "  EQUIPMENT (source of truth)\n"
        "       |\n"
        "       v\n"
        "  JOB REQUEST (references equipment_id)\n"
        "       |\n"
        "       v\n"
        "  JOB CARD (references job_request_id + equipment_id)\n"
        "       |\n"
        "       v\n"
        "  DOCUMENT (references job_card_id, generated on demand)\n"
        "       |\n"
        "       v\n"
        "  EQUIPMENT.last_calibration_date  <-- the ONLY back-write\n"
        "                                     (and it goes through the state machine)",
    )
    add_para(
        doc,
        "Why: If Job Cards could freely modify Equipment, audit trail breaks. ONE well-defined hook back-writes calibration metadata.",
        italic=True,
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 3 — ROLES
    # ============================================================
    add_section_banner(doc, "Section 3 — Role Model (5 Roles, Single-Tier Admin)")

    add_callout(
        doc,
        "MAJOR CHANGE: The earlier 'Admin' role has been DELETED. Super Admin is "
        "now the single admin tier. The system has exactly 5 roles.",
        fill="FFF4E5",
    )

    add_heading(doc, "3.1 Role Hierarchy", level=2)
    add_mono_block(
        doc,
        "                +------------------------+\n"
        "                |  SUPER ADMIN           | <-- role-assigner,\n"
        "                |  (only admin tier)     |     master data owner,\n"
        "                +-----------+------------+     ultimate authority\n"
        "                            |\n"
        "                +-----------v------------+\n"
        "                |  LAB IN-CHARGE         | <-- approves jobs,\n"
        "                |                        |     verifies/closes\n"
        "                +-----------+------------+\n"
        "                            |\n"
        "                +-----------v------------+\n"
        "                |  LAB ENGINEER          | <-- executes work,\n"
        "                |                        |     fills job cards\n"
        "                +-----------+------------+\n"
        "                            |\n"
        "             +--------------+--------------+\n"
        "             v                             v\n"
        "  +--------------------+        +---------------------+\n"
        "  |  NORMAL USER       |        |  VIEW-ONLY USER      |\n"
        "  | raises requests +  |        | read-only — no       |\n"
        "  | registers equip!   |        | writes EVER          |\n"
        "  +--------------------+        +---------------------+",
    )

    add_heading(doc, "3.2 Equipment Registration — Cross-Role Capability", level=2)
    add_callout(
        doc,
        "All roles EXCEPT View-Only can register a new equipment. "
        "Even a Normal User can register. The verification (PENDING → ACTIVE) "
        "still requires Lab In-charge or Super Admin.",
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 4 — RBAC
    # ============================================================
    add_section_banner(doc, "Section 4 — Authorization & RBAC (3-Layer Model)")

    add_heading(doc, "4.1 Why 3 Layers, Not 2", level=2)
    add_mono_block(
        doc,
        "   USER  --->  ROLE  --->  PERMISSION  --->  RESOURCE+ACTION\n"
        "    |          |              |                     |\n"
        "    |          |              |                     +-- \"equipment:create\"\n"
        "    |          |              +-- Granular toggle stored in DB\n"
        "    |          +-- Role bundle (e.g. \"Lab Engineer\")\n"
        "    +-- Identity (employee_id)\n",
    )

    add_table(
        doc,
        ["Layer", "Real-world Example", "Stored In"],
        [
            ["User", "Employee #EMP1234 (Deep)", "users"],
            ["Role", "Lab Engineer", "roles + user_roles"],
            ["Permission", "job_card:update, equipment:read", "permissions + role_permissions"],
            ["Resource:Action", "Checked at API middleware", "Runtime check"],
        ],
    )

    add_heading(doc, "4.2 Why This Matters", level=2)
    add_bullets(
        doc,
        [
            "2-layer (User→Role) = you must redeploy every time business changes “who can do what”.",
            "3-layer (User→Role→Permission) = Super Admin toggles a checkbox; behavior changes instantly. No redeploy.",
            "BR-RBAC-03: code NEVER checks by role name. Always by permission. This is the contract.",
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 5 — FINAL PERMISSION MATRIX
    # ============================================================
    add_section_banner(doc, "Section 5 — Final Consolidated Permission Matrix")
    add_para(
        doc,
        "This replaces every prior permission matrix in this conversation. One source of truth.",
        italic=True,
        color=GREY,
    )

    role_cols = ["Resource:Action", "Normal", "View-Only", "Lab Engr", "Lab In-charge", "Super Admin"]

    add_heading(doc, "5.1 Authentication & Identity", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["auth:login", "YES", "YES", "YES", "YES", "YES"],
            ["auth:logout", "YES", "YES", "YES", "YES", "YES"],
            ["auth:refresh-token", "YES", "YES", "YES", "YES", "YES"],
            ["me:read (own profile)", "YES", "YES", "YES", "YES", "YES"],
        ],
    )

    add_heading(doc, "5.2 User & Role Management (Super Admin only)", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["user:read-list", "NO", "NO", "NO", "NO", "YES"],
            ["user:role-assign", "NO", "NO", "NO", "NO", "YES"],
            ["user:activate / user:deactivate", "NO", "NO", "NO", "NO", "YES"],
        ],
    )

    add_heading(doc, "5.3 Equipment (equipment:create open to ALL except View-Only)", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["equipment:read-list", "YES", "YES", "YES", "YES", "YES"],
            ["equipment:read-detail", "YES", "YES", "YES", "YES", "YES"],
            ["equipment:create  (NEW BEHAVIOUR)", "YES", "NO", "YES", "YES", "YES"],
            ["equipment:update", "NO", "NO", "YES", "YES", "YES"],
            ["equipment:verify (PENDING -> ACTIVE)", "NO", "NO", "NO", "YES", "YES"],
            ["equipment:delete (hard)", "NO", "NO", "NO", "NO", "YES"],
            ["equipment:condemn (status flip)", "NO", "NO", "NO", "YES", "YES"],
        ],
    )

    add_heading(doc, "5.4 Job Requests", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["job_request:create", "YES", "NO", "YES", "YES", "YES"],
            ["job_request:read-own", "YES", "YES", "YES", "YES", "YES"],
            ["job_request:read-all", "NO", "YES", "YES", "YES", "YES"],
            ["job_request:approve", "NO", "NO", "NO", "YES", "YES"],
            ["job_request:reject", "NO", "NO", "NO", "YES", "YES"],
            ["job_request:assign-engineer", "NO", "NO", "NO", "YES", "YES"],
        ],
    )

    add_heading(doc, "5.5 Job Cards (Super Admin + Lab In-charge full rights)", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["job_card:read-list", "NO", "YES", "YES", "YES", "YES"],
            ["job_card:read-detail", "NO", "YES", "YES", "YES", "YES"],
            ["job_card:start-work", "NO", "NO", "YES (own)", "YES", "YES"],
            ["job_card:update-tasks", "NO", "NO", "YES (own)", "YES", "YES"],
            ["job_card:complete", "NO", "NO", "YES (own)", "YES", "YES"],
            ["job_card:verify-close", "NO", "NO", "NO", "YES", "YES"],
            ["job_card:reopen", "NO", "NO", "NO", "YES", "YES"],
            ["job_card:generate-pdf", "NO", "YES", "YES", "YES", "YES"],
        ],
    )

    add_heading(doc, "5.6 Dashboard & Inquiry", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["dashboard:view", "YES", "YES", "YES", "YES", "YES"],
            ["inquiry:search-vendors", "YES", "YES", "YES", "YES", "YES"],
            ["inquiry:search-products", "YES", "YES", "YES", "YES", "YES"],
            ["inquiry:search-job-cards", "NO", "YES", "YES", "YES", "YES"],
            ["inquiry:search-instruments", "YES", "YES", "YES", "YES", "YES"],
        ],
    )

    add_heading(doc, "5.7 Master Data (Super Admin only — Phase 2 build)", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["master:employees:manage", "NO", "NO", "NO", "NO", "YES"],
            ["master:vendors:manage", "NO", "NO", "NO", "NO", "YES"],
            ["master:equipment-types:manage", "NO", "NO", "NO", "NO", "YES"],
            ["master:divisions:manage", "NO", "NO", "NO", "NO", "YES"],
            ["master:lookup-values:manage", "NO", "NO", "NO", "NO", "YES"],
        ],
    )

    add_heading(doc, "5.8 Audit & Logs", level=2)
    add_table(
        doc,
        role_cols,
        [
            ["audit_log:read", "NO", "NO", "NO", "NO", "YES"],
            ["export:trigger (PDF / future Excel)", "NO", "NO", "YES", "YES", "YES"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 6 — AUTH FLOW
    # ============================================================
    add_section_banner(doc, "Section 6 — Authentication Flow")

    add_heading(doc, "6.1 Login Flow (v1: employee_id + password; SSO-ready)", level=2)
    add_mono_block(
        doc,
        "  User submits credentials  OR  SSO returns identity assertion (future)\n"
        "                       |\n"
        "                       v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 1: Validate employee_id                 |\n"
        "  |  - v1: bcryptjs hash check vs DB             |\n"
        "  |  - future: SSO signature verify              |\n"
        "  +----------------------------------------------+\n"
        "                       | valid\n"
        "                       v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 2: Lookup user in `users` table         |\n"
        "  |  - exists?                                   |\n"
        "  |  - is_active = TRUE?                         |\n"
        "  |  - is_locked = FALSE?                        |\n"
        "  +----------------------------------------------+\n"
        "                       |\n"
        "          +------------+------------+\n"
        "          v                         v\n"
        "    NOT FOUND                   FOUND\n"
        "          |                         |\n"
        "          v                         v\n"
        "  +----------------+   +---------------------------+\n"
        "  | Auto-provision |   | Fetch user_roles JOIN     |\n"
        "  | as Normal User |   |     roles -> permissions  |\n"
        "  +-------+--------+   +-------------+-------------+\n"
        "          |                          |\n"
        "          +------------+-------------+\n"
        "                       v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 3: Generate ACCESS token (JWT)          |\n"
        "  |   - sub: employee_id                         |\n"
        "  |   - role + permissions[]                     |\n"
        "  |   - exp: 15 min                              |\n"
        "  |   - jti: unique id (for revoke list)         |\n"
        "  +-----------------+----------------------------+\n"
        "                    |\n"
        "                    v\n"
        "  +----------------------------------------------+\n"
        "  | STEP 4: Issue REFRESH cookie                 |\n"
        "  |   - httpOnly + Secure + SameSite=Lax         |\n"
        "  |   - exp: 7 days (hashed in DB)               |\n"
        "  +-----------------+----------------------------+\n"
        "                    |\n"
        "                    v\n"
        "       +---------------------------+\n"
        "       | Authenticated session     |\n"
        "       | -> write login_audit row  |\n"
        "       +---------------------------+",
    )

    add_heading(doc, "6.2 Session Envelopes (Three Concentric Timers)", level=2)
    add_mono_block(
        doc,
        "  +========================================================+\n"
        "  |  REFRESH TOKEN MAX  (7 days, absolute)                 |\n"
        "  |  +--------------------------------------------------+  |\n"
        "  |  |  60-MIN INACTIVITY WINDOW (sliding)              |  |\n"
        "  |  |  +--------------------------------------------+  |  |\n"
        "  |  |  | 15-MIN JWT ACCESS TOKEN (short-lived)      |  |  |\n"
        "  |  |  | refreshed silently on each authed API call |  |  |\n"
        "  |  |  +--------------------------------------------+  |  |\n"
        "  |  |  resets on every authenticated request           |  |\n"
        "  |  +--------------------------------------------------+  |\n"
        "  |  hard expiry; user MUST re-login                       |\n"
        "  +========================================================+",
    )

    add_heading(doc, "6.3 First Super Admin Bootstrap", level=2)
    add_bullets(
        doc,
        [
            "Seeded via DB migration on first deploy (NOT manual SQL).",
            "Env var: SUPER_ADMIN_EMPLOYEE_IDS (comma-separated, at least 2).",
            "Migration inserts users + user_roles (SUPER_ADMIN) + audit_log 'BOOTSTRAP' rows.",
            "First login forces password reset (password_must_change = TRUE).",
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 7 — STATE MACHINES
    # ============================================================
    add_section_banner(doc, "Section 7 — Critical State Machines")

    add_heading(doc, "7.1 Job Request Lifecycle", level=2)
    add_mono_block(
        doc,
        "  +---------+   submit    +--------------+   assign   +--------------+\n"
        "  | DRAFT   |------------>| SUBMITTED    |----------->| ASSIGNED     |\n"
        "  +---------+             +------+-------+            +------+-------+\n"
        "                                 | reject                    | start\n"
        "                                 v                           v\n"
        "                          +--------------+           +--------------+\n"
        "                          | REJECTED     |           | IN-PROGRESS  |\n"
        "                          +--------------+           +------+-------+\n"
        "                                                            | complete\n"
        "                                                            v\n"
        "                                                     +--------------+\n"
        "                                                     | COMPLETED    |\n"
        "                                                     +------+-------+\n"
        "                                                            | verify\n"
        "                                                            v\n"
        "                                                     +----------------+\n"
        "                                                     | VERIFIED/CLOSED|\n"
        "                                                     +------+---------+\n"
        "                                                            | reopen\n"
        "                                                            v\n"
        "                                                     +--------------+\n"
        "                                                     | REOPENED     |\n"
        "                                                     +--------------+",
    )

    add_table(
        doc,
        ["From → To", "Allowed By", "Side Effects"],
        [
            ["DRAFT → SUBMITTED", "Owner", "Notify Lab In-charge"],
            ["SUBMITTED → ASSIGNED", "Lab In-charge", "Auto-create Job Card"],
            ["SUBMITTED → REJECTED", "Lab In-charge", "Mandatory reason; lock further edits"],
            ["ASSIGNED → IN-PROGRESS", "Assigned Engineer", "Set start_time"],
            ["IN-PROGRESS → COMPLETED", "Assigned Engineer", "Required: observations exist (BR-JC-07)"],
            ["COMPLETED → VERIFIED/CLOSED", "Lab In-charge", "Generate cert; update Equipment.last_cal_date"],
            ["VERIFIED → REOPENED", "Lab In-charge / Super Admin", "Audit log mandatory reason"],
        ],
    )

    add_heading(doc, "7.2 Equipment Lifecycle (with NEW PENDING_VERIFICATION state)", level=2)
    add_callout(
        doc,
        "NEW: All equipment registrations now START at PENDING_VERIFICATION. "
        "Only Lab In-charge or Super Admin can flip to ACTIVE (BR-EQP-10).",
        fill="E8F5E9",
    )
    add_mono_block(
        doc,
        "  [PENDING_VERIFICATION] --verify by InC/SA--> [ACTIVE]\n"
        "                                                 |\n"
        "                                                 +---> [UNDER_CALIBRATION] --> [ACTIVE]\n"
        "                                                 |             |\n"
        "                                                 |             +--> [OUT_OF_TOLERANCE] --> [UNDER_REPAIR] --> [ACTIVE]\n"
        "                                                 |\n"
        "                                                 +---> [UNDER_REPAIR] --> [ACTIVE]\n"
        "                                                 |\n"
        "                                                 +---> [QUARANTINED] --> [CONDEMNED/RETIRED]   (terminal)",
    )

    add_heading(doc, "7.3 Implementation Pattern (Single Choke-point)", level=2)
    add_mono_block(
        doc,
        "  // ONE function per entity. Single source of truth.\n"
        "  function transitionJobRequest(currentState, action, actor) {\n"
        "    const allowed = ALLOWED_TRANSITIONS[currentState]?.[action];\n"
        "    if (!allowed) throw new Error('Illegal transition');\n"
        "    if (!actor.hasPermission(allowed.permission)) throw new Error('Forbidden');\n"
        "    return { newState: allowed.to, sideEffects: allowed.effects };\n"
        "  }",
    )
    add_para(
        doc,
        "Why this matters: no developer writes raw UPDATE statements for status anywhere. "
        "Every transition writes to *_status_history automatically. Audit trail is free.",
        italic=True,
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 8 — BUSINESS RULES
    # ============================================================
    add_section_banner(doc, "Section 8 — Business Rules Catalogue")
    add_para(
        doc,
        "These are the contract between the business and the code. Every API endpoint enforces them. "
        "36+ rules across 9 categories.",
        italic=True,
        color=GREY,
    )

    add_heading(doc, "8.1 BR-AUTH — Authentication & Identity", level=2)
    add_table(
        doc,
        ["ID", "Rule"],
        [
            ["BR-AUTH-01", "Login is by employee_id only (not email/username)."],
            ["BR-AUTH-02", "User must exist in org's employee directory — no self-registration."],
            ["BR-AUTH-03", "Authenticated user with no role → defaults to Normal User."],
            ["BR-AUTH-04", "Sessions expire after 60 minutes of inactivity. Refresh token valid 7 days."],
            ["BR-AUTH-05", "First Super Admin seeded via DB migration using SUPER_ADMIN_EMPLOYEE_IDS env var. Seed at least 2."],
            ["BR-AUTH-06", "All login attempts (success + failure) are logged."],
            ["BR-AUTH-07", "Deactivated user (status = INACTIVE) cannot log in but history is preserved."],
        ],
    )

    add_heading(doc, "8.2 BR-RBAC — Authorization", level=2)
    add_table(
        doc,
        ["ID", "Rule"],
        [
            ["BR-RBAC-01", "Only Super Admin can assign/change a user's role."],
            ["BR-RBAC-02", "A user has exactly ONE primary role at a time (no multi-role in v1)."],
            ["BR-RBAC-03", "Permissions derived: User → Role → Permissions. NEVER check by role name."],
            ["BR-RBAC-04", "Sidebar items and routes filtered by permissions, not role names."],
            ["BR-RBAC-05", "Every API endpoint enforces permission check at controller layer."],
            ["BR-RBAC-06", "View-Only users can read but cannot trigger any write operation."],
            ["BR-RBAC-07", "Role changes take effect on next login OR token refresh."],
        ],
    )

    add_heading(doc, "8.3 BR-EQP — Equipment", level=2)
    add_table(
        doc,
        ["ID", "Rule"],
        [
            ["BR-EQP-01", "Every equipment has a unique serial number system-wide."],
            ["BR-EQP-02", "Equipment must belong to T&ME or F&PE category."],
            ["BR-EQP-03", "F&PE may have NO calibration frequency (optional)."],
            ["BR-EQP-04", "T&ME MUST have calibration frequency (months)."],
            ["BR-EQP-05", "next_cal_due_date = last_cal_date + calibration_frequency_months."],
            ["BR-EQP-06", "Status transitions follow the equipment state machine."],
            ["BR-EQP-07", "Hard DELETE = Super Admin only. CONDEMN = Lab In-charge OR Super Admin."],
            ["BR-EQP-08", "Search is case-insensitive across serial, model, manufacturer, type."],
            ["BR-EQP-09", "NEW. Every registration captures registered_by + verified_by + timestamps."],
            ["BR-EQP-10", "NEW. New equipment defaults to PENDING_VERIFICATION; verify by Lab In-charge / Super Admin → ACTIVE."],
        ],
    )

    add_heading(doc, "8.4 BR-JR — Job Requests", level=2)
    add_table(
        doc,
        ["ID", "Rule"],
        [
            ["BR-JR-01", "Must reference existing equipment (or trigger registration)."],
            ["BR-JR-02", "Job type ∈ {Calibration, Repair, Registration}."],
            ["BR-JR-03", "User can save as DRAFT before submitting."],
            ["BR-JR-04", "Once SUBMITTED, only Lab In-charge can change state."],
            ["BR-JR-05", "Submitted request must be assigned to a Lab Engineer before becoming a Job Card."],
            ["BR-JR-06", "'submitted by' auto-filled from current user — not overridable."],
            ["BR-JR-07", "High-priority repairs appear at top of Lab In-charge queue."],
            ["BR-JR-08", "Rejection requires mandatory reason (free text + reason code)."],
        ],
    )

    add_heading(doc, "8.5 BR-JC — Job Cards", level=2)
    add_table(
        doc,
        ["ID", "Rule"],
        [
            ["BR-JC-01", "Auto-created when a request is approved + assigned."],
            ["BR-JC-02", "Lifecycle: ASSIGNED → IN-PROGRESS → COMPLETED → VERIFIED/CLOSED."],
            ["BR-JC-03", "Only the assigned engineer can mark IN-PROGRESS or COMPLETED."],
            ["BR-JC-04", "Only Lab In-charge can verify/close."],
            ["BR-JC-05", "Closed job can only be REOPENED by Lab In-charge with reason."],
            ["BR-JC-06", "Tasks are configurable per job type."],
            ["BR-JC-07", "Calibration cards require before-reading + after-reading + environment before COMPLETED."],
            ["BR-JC-08", "Job card history is append-only — immutable state transition log."],
        ],
    )

    add_heading(doc, "8.6 BR-PDF, BR-AUD, BR-VIS, BR-MASTER", level=2)
    add_table(
        doc,
        ["ID", "Rule"],
        [
            ["BR-PDF-01", "PDFs generated on demand from current DB state. Nothing stored."],
            ["BR-PDF-02", "Job card PDF: header, equipment info, tasks, observations, signatures, date."],
            ["BR-PDF-03", "Calibration cert PDF: equipment, standards, readings, environment, uncertainty, valid-until."],
            ["BR-PDF-04", "All PDFs include generation timestamp + footer + record ID."],
            ["BR-AUD-01", "Every write on critical tables is logged in audit_logs."],
            ["BR-AUD-02", "Audit log captures who, what, when, before/after, IP, user-agent."],
            ["BR-AUD-03", "Exports logged with user + record IDs accessed."],
            ["BR-AUD-04", "Sensitive fields filtered from API responses based on permission."],
            ["BR-AUD-05", "HTTPS in prod. JWT in Authorization header. Refresh in httpOnly cookie."],
            ["BR-VIS-01", "Normal User sees only their own job requests."],
            ["BR-VIS-02", "Lab Engineer sees all jobs assigned to them + the queue."],
            ["BR-VIS-03", "Lab In-charge and above see all jobs and equipment."],
            ["BR-VIS-04", "View-Only sees all data, performs no actions."],
            ["BR-MASTER-01", "NEW. All master-data CRUD is Super Admin only."],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 9 — FRs
    # ============================================================
    add_section_banner(doc, "Section 9 — Functional Requirements (MVP)")
    add_para(doc, "45 FRs across 6 MVP modules. 44 MUST + 1 SHOULD.", italic=True, color=GREY)

    add_heading(doc, "9.1 FR-A — Auth & RBAC Module", level=2)
    add_table(
        doc,
        ["FR ID", "Description", "Priority"],
        [
            ["FR-A-01", "Login screen accepts employee_id + password", "MUST"],
            ["FR-A-02", "Successful login returns JWT + sets refresh cookie", "MUST"],
            ["FR-A-03", "Logout clears tokens and invalidates session", "MUST"],
            ["FR-A-04", "Protected routes redirect to login when unauthenticated", "MUST"],
            ["FR-A-05", "Sidebar items render based on user's permission set", "MUST"],
            ["FR-A-06", "Super Admin can view list of all users + their roles", "MUST"],
            ["FR-A-07", "Super Admin can assign/change a user's role", "MUST"],
            ["FR-A-08", "Super Admin can activate/deactivate a user", "MUST"],
            ["FR-A-09", "Backend exposes GET /me returning current user + permissions", "MUST"],
            ["FR-A-10", "Token refresh endpoint silently extends session", "MUST"],
        ],
    )

    add_heading(doc, "9.2 FR-E — Equipment Module", level=2)
    add_table(
        doc,
        ["FR ID", "Description", "Priority"],
        [
            ["FR-E-01", "List equipment with pagination, search, filter (type, status)", "MUST"],
            ["FR-E-02", "View equipment detail page with specs, cal history, linked jobs", "MUST"],
            ["FR-E-03", "Register new equipment (available to all roles except View-Only)", "MUST"],
            ["FR-E-04", "Edit equipment details (specs, status)", "MUST"],
            ["FR-E-05", "Color-code calibration due dates (green/yellow/red)", "MUST"],
            ["FR-E-06", "Show equipment's calibration history timeline", "MUST"],
            ["FR-E-07", "Search equipment by serial, model, manufacturer", "MUST"],
            ["FR-E-08", "Soft-delete (status = CONDEMNED) for non-super-admins", "MUST"],
            ["FR-E-09", "NEW. Lab In-charge / Super Admin verify equipment (PENDING → ACTIVE)", "MUST"],
        ],
    )

    add_heading(doc, "9.3 FR-JR — Job Request Module", level=2)
    add_table(
        doc,
        ["FR ID", "Description", "Priority"],
        [
            ["FR-JR-01", "List job requests with filters (type, status, my requests vs all)", "MUST"],
            ["FR-JR-02", "Create job request form (multi-section, progressive disclosure)", "MUST"],
            ["FR-JR-03", "Save as draft", "MUST"],
            ["FR-JR-04", "Submit triggers state transition to SUBMITTED", "MUST"],
            ["FR-JR-05", "View request detail with full info + state history", "MUST"],
            ["FR-JR-06", "Lab In-charge can approve / reject with reason", "MUST"],
            ["FR-JR-07", "Approval triggers automatic Job Card creation", "MUST"],
            ["FR-JR-08", "Assignment dropdown shows lab engineers with current workload", "SHOULD"],
        ],
    )

    add_heading(doc, "9.4 FR-JC — Job Card Module", level=2)
    add_table(
        doc,
        ["FR ID", "Description", "Priority"],
        [
            ["FR-JC-01", "List job cards with filter (status, assigned-to-me, all)", "MUST"],
            ["FR-JC-02", "View job card with horizontal status stepper", "MUST"],
            ["FR-JC-03", "Task checklist (configurable per job type)", "MUST"],
            ["FR-JC-04", "Observations & readings log (free text + structured)", "MUST"],
            ["FR-JC-05", "Engineer can transition state (Start work, Mark complete)", "MUST"],
            ["FR-JC-06", "Lab In-charge can Verify/Close or Reopen", "MUST"],
            ["FR-JC-07", "Generate PDF of job card (on-demand download)", "MUST"],
            ["FR-JC-08", "State history visible as timeline", "MUST"],
        ],
    )

    add_heading(doc, "9.5 FR-D — Dashboard", level=2)
    add_table(
        doc,
        ["FR ID", "Description", "Priority"],
        [
            ["FR-D-01", "Role-aware KPI widgets (different per role)", "MUST"],
            ["FR-D-02", "Calibration due alerts (next 30 days + overdue)", "MUST"],
            ["FR-D-03", "Engineer workload bar chart (for Lab In-charge)", "MUST"],
            ["FR-D-04", "Equipment status pie chart", "MUST"],
            ["FR-D-05", "Recently updated jobs table", "MUST"],
            ["FR-D-06", "Quick actions: Raise request, Add equipment (permission-gated)", "MUST"],
        ],
    )

    add_heading(doc, "9.6 FR-I — Inquiry", level=2)
    add_table(
        doc,
        ["FR ID", "Description", "Priority"],
        [
            ["FR-I-01", "4-tab interface (Vendor, Product, Job Card, Instrument)", "MUST"],
            ["FR-I-02", "Real-time search with debounce", "MUST"],
            ["FR-I-03", "Result tables with drill-down link to full record", "MUST"],
            ["FR-I-04", "Empty-state messages when no results", "MUST"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 10 — NFRs
    # ============================================================
    add_section_banner(doc, "Section 10 — Non-Functional Requirements")

    add_table(
        doc,
        ["Category", "Requirement", "Target"],
        [
            ["Performance", "Initial page load (cold)", "< 3 sec on intranet"],
            ["Performance", "API response (95th percentile)", "< 500 ms"],
            ["Performance", "List pagination", "25 default, 100 max"],
            ["Security", "Password hashing", "bcrypt ≥ 10 rounds"],
            ["Security", "JWT access lifetime", "15 min"],
            ["Security", "Refresh token lifetime", "7 days, httpOnly cookie, SameSite=Lax"],
            ["Security", "SQL injection defence", "Parameterized queries everywhere"],
            ["Security", "XSS", "React auto-escape + CSP header"],
            ["Security", "CSRF", "SameSite cookies + CSRF token on /auth/refresh"],
            ["Reliability", "Server uptime (post go-live)", "99% during business hours"],
            ["Reliability", "DB connection pool", "10–20 connections"],
            ["Usability", "Max clicks to any feature", "≤ 3"],
            ["Usability", "Keyboard shortcuts", "On power-user screens"],
            ["Usability", "Responsive viewport", "1280–1920 primary; graceful to 768"],
            ["Audit", "Every state-changing op logged", "Yes"],
            ["Audit", "Audit log retention", "Indefinite (until specified)"],
            ["Maintainability", "Code style", "ESLint + Prettier (enforced via husky)"],
            ["Maintainability", "Folder structure", "routes / controllers / services / models / utils"],
            ["Maintainability", "API versioning", "/api/v1/... from day 1"],
            ["Compatibility", "Browsers", "Chrome, Edge, Firefox (latest 2 versions)"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 11 — DB STRATEGY
    # ============================================================
    add_section_banner(doc, "Section 11 — Database Strategy & 70-Table Map")

    add_callout(
        doc,
        "REALITY CHECK: ~64 tables already exist with records. CMCMIS does NOT "
        "start on a blank slate — we refactor/extend existing schema in Phase 1.",
        fill="FFF4E5",
    )

    add_heading(doc, "11.1 Phase 1 Classification Approach", level=2)
    add_mono_block(
        doc,
        "  STEP 1: Inventory existing schema\n"
        "          (export current DB structure via phpMyAdmin)\n"
        "  STEP 2: Classify each existing table\n"
        "          - KEEP AS-IS         (use without changes)\n"
        "          - EXTEND             (add columns, keep data)\n"
        "          - REFACTOR           (restructure, migrate data)\n"
        "          - DEPRECATE          (mark unused, remove post-MVP)\n"
        "          - NEW                (CMCMIS additions)\n"
        "  STEP 3: Decide migration strategy\n"
        "          - Add new tables alongside existing (low risk)\n"
        "          - Write data migration for refactored tables\n"
        "          - Never drop in MVP — only deprecate\n"
        "  STEP 4: Lock final ERD (~70 tables)\n"
        "  STEP 5: Write seed scripts for lookup data",
    )

    add_heading(doc, "11.2 Logical Group Map (Targeting ~70 Tables)", level=2)
    add_table(
        doc,
        ["Group", "Approx Tables", "Key Tables"],
        [
            ["Identity & RBAC", "~10", "users, employees, roles, permissions, user_roles, role_permissions, refresh_tokens, login_audit, password_reset_tokens, sso_identity_mappings"],
            ["Job Requests", "~7", "job_requests, job_request_types, status_history, attachments, comments, approvals, priorities"],
            ["Job Cards", "~9", "job_cards, tasks, task_templates, observations, status_history, assignees, signatures, time_logs, attachments"],
            ["Equipment", "~11", "equipment, types, specifications, calibration_history, maintenance_history, status_history, documents, warranty, amc, manufacturers, categories"],
            ["Schedule (P2)", "~4", "schedules, recurrence_rules, notifications, assignments"],
            ["Procurement + Vendors (P2)", "~10", "purchase_orders, po_items, status_history, approvals, spare_parts, stock_movements, vendors, contacts, categories, documents"],
            ["Master / Lookups", "~10", "divisions, departments, locations, labs, system_types, units, currencies, config_settings, numbering_sequences, holiday_calendar"],
            ["Documents", "~4", "documents, versions, categories, signatures"],
            ["System / Observability", "~7", "notifications, templates, audit_log, error_log, background_jobs, email_dispatch_log, system_announcements"],
        ],
    )

    add_heading(doc, "11.3 Naming & Design Conventions (Locked)", level=2)
    add_table(
        doc,
        ["Convention", "Rule"],
        [
            ["Table names", "snake_case, plural"],
            ["Primary keys", "BIGINT UNSIGNED AUTO_INCREMENT named id"],
            ["Foreign keys", "<ref_table_singular>_id"],
            ["Timestamps", "Every table has created_at, updated_at"],
            ["Soft deletes", "is_deleted BOOLEAN + deleted_at (no hard deletes for transactional data)"],
            ["Status fields", "ENUM only for tiny fixed lists; otherwise FK to lookup table"],
            ["Audit fields", "created_by, updated_by (FK → users.id)"],
            ["Money fields", "DECIMAL(18,4) — never FLOAT"],
            ["Charset", "utf8mb4 + utf8mb4_unicode_ci"],
            ["Engine", "InnoDB always (transactions, FK enforcement)"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 12 — TECH STACK
    # ============================================================
    add_section_banner(doc, "Section 12 — Tech Stack (Frontend + Backend)")

    add_heading(doc, "12.1 Frontend Stack (React 18 + Vite + Tailwind v3)", level=2)
    add_table(
        doc,
        ["#", "Concern", "Library", "Why"],
        [
            ["1", "Build", "vite", "Fast HMR, modern"],
            ["2", "Framework", "react@18", "Locked"],
            ["3", "Routing", "react-router-dom v6", "SPA standard"],
            ["4", "State (local)", "useState / useReducer", "Built-in"],
            ["5", "State (global UI)", "zustand", "~1KB, no boilerplate"],
            ["6", "Server state", "@tanstack/react-query", "Cache/refetch/mutations"],
            ["7", "HTTP", "axios", "Interceptors for auth/refresh"],
            ["8", "Forms", "react-hook-form + zod", "Performant + schema shared with BE"],
            ["9", "UI primitives", "Custom on Tailwind", "Avoids shadcn/MUI lock-in"],
            ["10", "Icons", "lucide-react", "Tree-shakable"],
            ["11", "Charts", "recharts", "Dashboard widgets"],
            ["12", "Date", "dayjs", "Symmetric with backend"],
            ["13", "Toasts", "sonner", "Lightweight"],
            ["14", "Styling", "tailwindcss v3", "Locked"],
            ["15", "Lint/Format", "eslint + prettier + husky + lint-staged", "Pre-commit"],
        ],
    )

    add_heading(doc, "12.2 Backend Stack (Node + Express 4)", level=2)
    add_table(
        doc,
        ["#", "Concern", "Library", "Ver", "Why"],
        [
            ["1", "Web framework", "express", "^4.x", "Stable, ubiquitous"],
            ["2", "DB driver", "mysql2/promise", "^3.x", "Fastest MySQL driver"],
            ["3", "Validation", "zod", "^3.x", "Single source of truth (FE+BE share)"],
            ["4", "Auth tokens", "jsonwebtoken", "^9.x", "Industry standard"],
            ["5", "Password hash", "bcryptjs", "^2.x", "Pure-JS, no native build pain"],
            ["6", "Logger", "pino + pino-pretty", "^8.x", "Faster than Winston, JSON"],
            ["7", "Env loader", "dotenv", "^16.x", "Standard"],
            ["8", "Env validation", "envalid", "^8.x", "Fail at boot if env missing"],
            ["9", "Date", "dayjs", "^1.x", "Symmetric with frontend"],
            ["10", "PDF", "pdfkit", "^0.14.x", "Programmatic, no Chromium"],
            ["11", "Rate limit", "express-rate-limit", "^7.x", "In-memory, sufficient v1"],
            ["12", "Security", "helmet", "^7.x", "OWASP-friendly defaults"],
            ["13", "CORS", "cors", "^2.x", "Standard"],
            ["14", "Process mgr", "pm2", "latest", "Cluster mode, restarts, logs"],
            ["15", "Lint/Format", "eslint + prettier + husky + lint-staged", "latest", "Pre-commit"],
        ],
    )

    add_heading(doc, "12.3 The Three Big Architectural Wins", level=2)
    add_bullets(
        doc,
        [
            "Schema symmetry — zod on BOTH frontend and backend → one definition, two enforcers.",
            "Date symmetry — dayjs on BOTH sides → no timezone/format drift on BR-EQP-05 math.",
            "On-prem friendly — pdfkit (no Chromium), pm2 (no systemd quirks), no Redis/SMTP/object-storage dependencies.",
        ],
    )

    add_heading(doc, "12.4 Pending Stack Additions Awaiting Your OK", level=2)
    add_table(
        doc,
        ["#", "Addition", "Why"],
        [
            ["1", "cookie-parser (BE)", "Needed to read httpOnly refresh cookie"],
            ["2", "compression (BE)", "gzip/brotli — helps NFR p95 < 500ms"],
            ["3", "CSRF token on /auth/refresh", "NFR mandates; custom double-submit token recommended"],
            ["4", "@tanstack/react-table (FE)", "Headless table for equipment/job lists"],
            ["5", "@tailwindcss/forms (FE)", "Sane form defaults"],
            ["6", "vitest + supertest", "Unit + API testing for auth + state machines"],
            ["7", "swagger-ui-express + zod-to-openapi", "Auto API docs (optional)"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 13 — SYSTEM ARCHITECTURE
    # ============================================================
    add_section_banner(doc, "Section 13 — System Architecture & Request Trace")

    add_heading(doc, "13.1 Runtime Architecture", level=2)
    add_mono_block(
        doc,
        "  +---------------------------------------------------------------+\n"
        "  | DESKTOP / LAPTOP BROWSER (NO MOBILE)                          |\n"
        "  | React 18 + Vite + Tailwind                                    |\n"
        "  |   Routes -> React Router                                      |\n"
        "  |   Server state -> TanStack Query                              |\n"
        "  |   Forms -> React Hook Form + zod                              |\n"
        "  |   HTTP -> axios (interceptor attaches JWT)                    |\n"
        "  +---------------------------+-----------------------------------+\n"
        "                              | HTTPS (JWT + httpOnly refresh)\n"
        "                              v\n"
        "  +---------------------------------------------------------------+\n"
        "  | NODE + EXPRESS API SERVER (on-prem)                           |\n"
        "  |  Middleware: helmet > cors > compression > json > cookies     |\n"
        "  |              > pino > rate-limit > auth > authorize > scope   |\n"
        "  |              > validate(zod) > controller                     |\n"
        "  |  Service layer: state-machine transition()                    |\n"
        "  |                  pdfkit cert generator                        |\n"
        "  |                  audit-log helper                             |\n"
        "  |  Data access: mysql2/promise pool, parameterized SQL          |\n"
        "  +---------------------------+-----------------------------------+\n"
        "                              v\n"
        "  +---------------------------------------------------------------+\n"
        "  | MySQL 8.x (existing on-prem; phpMyAdmin admin)                |\n"
        "  |   ~70 tables (existing 64 + new) · InnoDB · utf8mb4           |\n"
        "  |   FK + CHECK constraints · indexes per dashboard/inquiry      |\n"
        "  +---------------------------------------------------------------+\n"
        "\n"
        "  NOT IN MVP:  Redis  ·  SMTP  ·  S3/Object store  ·  SSO/AD",
    )

    add_heading(doc, "13.2 End-to-End Request Trace (Equipment Register)", level=2)
    add_mono_block(
        doc,
        "  Lab Engineer fills equipment register form\n"
        "    -> react-hook-form + zod validates on submit (no network if invalid)\n"
        "    -> @tanstack/react-query mutation calls axios.post('/api/v1/equipment')\n"
        "    -> axios interceptor attaches Bearer JWT\n"
        "    ----------------- NETWORK BOUNDARY -----------------\n"
        "    -> helmet (headers) > cors > compression > express.json > cookie-parser\n"
        "    -> pino logs request meta\n"
        "    -> express-rate-limit (auth routes only)\n"
        "    -> authenticate (jsonwebtoken verify)\n"
        "    -> authorize ('equipment:create' permission check)\n"
        "    -> rowLevelScope (no-op for create)\n"
        "    -> validate(zod) re-checks body\n"
        "    -> equipmentController.create\n"
        "    -> equipmentService.create:\n"
        "          - dayjs.add for next_cal_due_date\n"
        "          - state machine: status = PENDING_VERIFICATION\n"
        "          - audit_log row\n"
        "    -> mysql2 transaction: INSERT equipment + INSERT audit_log\n"
        "    -> 201 Created + JSON body\n"
        "    -> pino logs response status + duration\n"
        "    ----------------- NETWORK BOUNDARY -----------------\n"
        "    -> react-query invalidates 'equipment-list' cache\n"
        "    -> sonner: 'Equipment registered. Pending verification.'\n"
        "    -> router navigates to detail page",
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 14 — MIDDLEWARE PIPELINE
    # ============================================================
    add_section_banner(doc, "Section 14 — Express Middleware Pipeline")
    add_mono_block(
        doc,
        "  incoming request\n"
        "         |\n"
        "         v\n"
        "  +----------------------------------------------------+\n"
        "  | 1.  helmet()              <- security headers      |\n"
        "  | 2.  cors(orgAllowlist)    <- reject foreign origins|\n"
        "  | 3.  compression()         <- gzip responses        |\n"
        "  | 4.  express.json({limit}) <- parse JSON body       |\n"
        "  | 5.  cookie-parser()       <- parse refresh cookie  |\n"
        "  | 6.  pino-http()           <- log request meta      |\n"
        "  | 7.  rateLimit(login)      <- only on auth routes   |\n"
        "  | 8.  authenticate()        <- JWT verify            |\n"
        "  | 9.  authorize(perm)       <- RBAC check            |\n"
        "  | 10. rowLevelScope()       <- BR-VIS WHERE clause   |\n"
        "  | 11. validate(zodSchema)   <- input validation      |\n"
        "  | 12. controller            <- business handler      |\n"
        "  | 13. errorHandler()        <- centralized errors    |\n"
        "  +----------------------------------------------------+\n"
        "         |\n"
        "         v\n"
        "  outgoing response",
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 15 — FOLDER STRUCTURE
    # ============================================================
    add_section_banner(doc, "Section 15 — Project Folder Structure")

    add_mono_block(
        doc,
        "  cmcmis-simplified/\n"
        "  |\n"
        "  +-- backend/\n"
        "  |   +-- src/\n"
        "  |   |   +-- config/        (db.js, env.js, logger.js, jwt.js)\n"
        "  |   |   +-- middleware/    (authenticate, authorize, rowLevelScope,\n"
        "  |   |   |                   validate, errorHandler)\n"
        "  |   |   +-- modules/\n"
        "  |   |   |   +-- auth/      (routes, controller, service, schema)\n"
        "  |   |   |   +-- users/\n"
        "  |   |   |   +-- equipment/\n"
        "  |   |   |   +-- jobRequests/\n"
        "  |   |   |   +-- jobCards/\n"
        "  |   |   |   +-- dashboard/\n"
        "  |   |   |   +-- inquiry/\n"
        "  |   |   |   +-- audit/\n"
        "  |   |   +-- stateMachines/ (jobRequest.fsm.js, equipment.fsm.js)\n"
        "  |   |   +-- pdf/           (jobCard.pdf.js, calibrationCert.pdf.js)\n"
        "  |   |   +-- utils/\n"
        "  |   |   +-- server.js\n"
        "  |   +-- db/\n"
        "  |   |   +-- migrations/    (numbered .sql files)\n"
        "  |   |   +-- seeds/         (roles, permissions, super admins)\n"
        "  |   |   +-- schema.sql\n"
        "  |   +-- tests/\n"
        "  |   +-- .env.example\n"
        "  |   +-- ecosystem.config.js (pm2)\n"
        "  |   +-- package.json\n"
        "  |\n"
        "  +-- frontend/\n"
        "  |   +-- src/\n"
        "  |   |   +-- api/           (axios client + per-module hooks)\n"
        "  |   |   +-- components/    (Button, Input, Table, Modal, FormField)\n"
        "  |   |   +-- features/      (auth, equipment, jobRequests, jobCards,\n"
        "  |   |   |                   dashboard, inquiry)\n"
        "  |   |   +-- layouts/       (AppLayout, AuthLayout)\n"
        "  |   |   +-- lib/           (authContext, permissions, formatters)\n"
        "  |   |   +-- pages/         (thin route components)\n"
        "  |   |   +-- schemas/       (zod schemas — shared with BE)\n"
        "  |   |   +-- stores/        (zustand)\n"
        "  |   |   +-- router.jsx\n"
        "  |   |   +-- main.jsx\n"
        "  |   +-- public/\n"
        "  |   +-- tailwind.config.js\n"
        "  |   +-- vite.config.js\n"
        "  |   +-- package.json\n"
        "  |\n"
        "  +-- docs/\n"
        "  |   +-- ERD.md\n"
        "  |   +-- rbac-matrix.md\n"
        "  |   +-- state-machines.md\n"
        "  |   +-- business-rules.md\n"
        "  +-- .gitignore\n"
        "  +-- .editorconfig\n"
        "  +-- README.md\n"
        "  +-- package.json   (monorepo root)",
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 16 — CODING CONVENTIONS
    # ============================================================
    add_section_banner(doc, "Section 16 — Coding Conventions & API Standards")
    add_table(
        doc,
        ["Concern", "Convention"],
        [
            ["API base", "/api/v1/..."],
            ["HTTP verbs", "REST: GET / POST / PATCH / DELETE"],
            ["Response shape (success)", "{ data, meta }"],
            ["Response shape (error)", "{ error: { code, message, details } }"],
            ["Status codes used", "200 / 201 / 204 / 400 / 401 / 403 / 404 / 409 / 422 / 429 / 500"],
            ["Pagination", "?page=1&limit=25 (max 100)"],
            ["Filtering", "?status=ACTIVE&type=TME (whitelisted per endpoint)"],
            ["Sorting", "?sort=-created_at,name (- prefix = DESC)"],
            ["Naming (DB)", "snake_case, plural tables, id PKs, *_id FKs"],
            ["Naming (JS)", "camelCase vars/functions, PascalCase components/classes"],
            ["Naming (files)", "kebab-case.js or PascalCase.jsx for React components"],
            ["Errors", "Throw typed errors → centralized handler → standard response"],
            ["Logs", "Structured JSON via pino; never console.log"],
            ["Commits", "Conventional Commits (feat / fix / chore / ...)"],
            ["Branches", "feature/*, fix/*, chore/* from main"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 17 — MVP SCOPE
    # ============================================================
    add_section_banner(doc, "Section 17 — MVP Scope vs Phase 2 Backlog")

    add_heading(doc, "17.1 MVP (10 weeks) — Signed Off", level=2)
    add_bullets(
        doc,
        [
            "Auth + RBAC (login, session, role loading, permissions, protected routes, SSO-ready)",
            "Equipment master + register + verify flow",
            "Job Requests (form, listing, lifecycle)",
            "Job Cards (tasks, observations, sign-off)",
            "Dashboard (role-aware KPIs, alerts, widgets)",
            "Inquiry (search hub — vendors, products, job cards, instruments)",
            "PDF generation (download/generate only — no storage)",
            "Audit logs (basic, all state changes)",
            "Responsive UI (desktop + laptop; 1280–1920 primary)",
            "Row-level visibility (per BR-VIS rules)",
            "Optimized SQL with pagination + connection pool",
        ],
    )

    add_heading(doc, "17.2 Phase 2 — Post-Internship Backlog", level=2)
    add_bullets(
        doc,
        [
            "Schedule module (PM + Calibration calendar)",
            "Procurement (POs + spares)",
            "Reports (analytics + exports)",
            "Admin master-data CRUD UI",
            "Notifications (in-app feed)",
        ],
    )

    add_heading(doc, "17.3 Explicitly NOT in MVP", level=2)
    add_table(
        doc,
        ["Item", "Decision"],
        [
            ["SSO / Active Directory integration", "Architecture SSO-ready, integration deferred"],
            ["Email / SMTP", "Out of scope until specified"],
            ["Redis / caching layer", "MySQL + pagination + pool is enough"],
            ["File storage", "PDFs generated + downloaded only, never stored"],
            ["Backup infrastructure", "Out of our scope"],
            ["Mobile / tablet UI", "Desktop + laptop only (responsive within range)"],
            ["Barcode / QR", "Not in scope"],
            ["NABL / ISO 17025 / AS9100 specifics", "Deferred — user will instruct"],
            ["Calibration cert template format", "Deferred — user will instruct"],
            ["Notification channels (email/in-app/SMS/push)", "Deferred — user will instruct"],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 18 — 10-WEEK WAR PLAN
    # ============================================================
    add_section_banner(doc, "Section 18 — 10-Week War Plan (Phased Timeline)")

    add_table(
        doc,
        ["Week", "Phase", "Deliverables"],
        [
            ["1", "Phase 0 — Foundation", "Repo scaffolding · ESLint/Prettier · Tailwind base layout · MySQL pool · 'hello world' protected route end-to-end"],
            ["2", "Phase 1 — DB Design", "Inventory ~64 existing tables · Classify · Draft new tables (RBAC + Job + Equipment first) · Seed scripts · Lock ERD"],
            ["3", "Block 1 — Auth + RBAC", "employee_id + password login · JWT + refresh · RBAC middleware (role + granular perm) · Sidebar visibility · Row-level scoping helper · Audit-log helper"],
            ["4–5", "Block 2a — Equipment Master", "Equipment list/detail/register · State machine (incl. PENDING_VERIFICATION) · Cal history · Specs"],
            ["5–6", "Block 2b — Job Requests", "Multi-section form (T&ME/F&PE; cal/repair/registration) · Draft/submit · Approve/reject · State machine"],
            ["6–7", "Block 2c — Job Cards + PDF", "Auto-create on approval · Task checklist + observations · Verify/close → Equipment status update · PDF cert generation"],
            ["8", "Block 3 — Dashboard + Inquiry", "Dashboard widgets (KPIs, due alerts, workload) · Inquiry 4-tab search"],
            ["9", "Hardening", "Audit coverage check · RBAC e2e test (all 5 roles) · Illegal-transition tests · SQL tuning · Responsive QA"],
            ["10", "Demo Prep + Deploy", "Demo data seeded · On-prem deploy · Stakeholder demo · Bug-fix buffer"],
        ],
    )

    add_heading(doc, "18.1 Slack / Buffer Strategy", level=2)
    add_table(
        doc,
        ["Risk", "Buffer Plan"],
        [
            ["Existing 64-table DB messier than expected", "Phase 1 stretches into Week 3; Auth gets compressed"],
            ["Job Card observations more complex than planned", "Cut Inquiry tabs to 2 (job cards + equipment) for demo"],
            ["PDF formatting takes longer", "Ship plain-template PDFs; polish post-internship"],
            ["RBAC edge cases discovered", "Core — never cut. Cut Dashboard widgets instead."],
        ],
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 19 — REAL-WORLD WALK-THROUGH
    # ============================================================
    add_section_banner(doc, "Section 19 — Real-World End-to-End Walk-Through")
    add_para(
        doc,
        "Tracing a single instrument (Digital Multimeter DMM-2034) through CMCMIS — to make all the architecture concrete.",
        italic=True,
        color=GREY,
    )

    add_table(
        doc,
        ["Step", "Actor", "Role", "Module", "Action"],
        [
            ["1", "Engineer Deep", "Lab Engineer", "Dashboard", "Sees DMM-2034 calibration due alert"],
            ["2", "Deep", "Lab Engineer", "Job Requests", "Raises calibration request"],
            ["3", "Lab In-charge", "Lab In-charge", "Job Requests", "Approves → Job Card auto-created"],
            ["4", "Deep (auto-assigned)", "Lab Engineer", "Job Cards", "Completes tasks, enters readings, marks COMPLETE"],
            ["5", "Lab In-charge", "Lab In-charge", "Job Cards", "Verifies + closes → triggers PDF cert generation"],
            ["6", "Anyone (except View-Only)", "Any of 4 roles", "Equipment", "Could register new DMM-2035 tomorrow (PENDING_VERIFICATION)"],
            ["7", "Lab In-charge", "Lab In-charge", "Equipment", "Verifies DMM-2035 → ACTIVE"],
            ["8", "Super Admin", "Super Admin", "Admin (seed/CLI in MVP)", "Promotes Normal User → Lab Engineer"],
            ["9", "View-Only auditor", "View-Only", "Inquiry", "Searches DMM-2034 history; reads everything; edits nothing"],
            ["10", "Super Admin", "Super Admin", "Audit Log", "Reviews every state transition"],
        ],
    )
    add_para(
        doc,
        "One instrument, one calibration cycle ≈ 20 DB writes across ~12 tables, 1 PDF cert, ~9 audit rows. "
        "This is why the ~70-table schema is the MINIMUM to track this lawfully.",
        italic=True,
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 20 — RECAP
    # ============================================================
    add_section_banner(doc, "Section 20 — Constraints, Open Items, Quick Recap")

    add_heading(doc, "20.1 Locked Constraints", level=2)
    add_table(
        doc,
        ["#", "Constraint", "Decision"],
        [
            ["1", "Organization context", "ISRO SAC-like defence/space-grade"],
            ["2", "Existing DB", "~64 tables; review in Phase 1; never drop in MVP"],
            ["3", "SSO in v1", "NO — employee_id login from existing DB; SSO-ready for future"],
            ["4", "Deployment target", "Internal on-prem / org private infra"],
            ["5", "File storage", "NO — PDFs generated + downloaded only"],
            ["6", "Email / SMTP", "NO in v1"],
            ["7", "Redis / cache", "NO — MySQL + pagination + pool"],
            ["8", "Backup infrastructure", "Out of scope (user handles separately)"],
            ["9", "Mobile / tablet UI", "NO — desktop/laptop only; fully responsive within range"],
            ["10", "Barcode / QR", "NO"],
            ["11", "PDF generation", "YES — pdfkit, server-side, download only"],
            ["12", "Sensitive data", "Defence-grade; row-level visibility + secure sessions + limited exports"],
        ],
    )

    add_heading(doc, "20.2 Open Items Awaiting Your Call", level=2)
    add_bullets(
        doc,
        [
            "Lock the 7 small stack additions (cookie-parser, compression, CSRF, table lib, tailwind-forms, testing, swagger).",
            "Confirm proposed folder structure (Section 15).",
            "Confirm coding conventions (Section 16).",
            "Hand over phpMyAdmin export of existing ~64 tables to start Phase 1 DB inventory.",
        ],
    )

    add_heading(doc, "20.3 Big Quick Recap (Print-and-Pin)", level=2)
    add_table(
        doc,
        ["Topic", "Locked Decision"],
        [
            ["Roles", "5 only — Super Admin, Lab In-charge, Lab Engineer, Normal, View-Only (Admin role DELETED)"],
            ["equipment:create", "Open to ALL roles except View-Only"],
            ["Equipment lifecycle", "Starts at PENDING_VERIFICATION → ACTIVE (BR-EQP-10)"],
            ["Per-user role count", "Exactly ONE primary role (BR-RBAC-02)"],
            ["Super Admin seed", "≥ 2 via SUPER_ADMIN_EMPLOYEE_IDS env var"],
            ["Session model", "15-min JWT · 60-min idle · 7-day refresh (concentric envelopes)"],
            ["Tech stack", "React 18+Vite+TW+Zustand+RQ+Axios+RHF+Zod / Express 4+mysql2+Zod+JWT+bcryptjs+Pino+PDFKit+Helmet+PM2"],
            ["Symmetry wins", "zod schemas + dayjs date math shared across FE/BE"],
            ["Audit retention", "Indefinite (until user specifies)"],
            ["API base", "/api/v1/... from day 1"],
            ["Timeline", "10 weeks (solo dev + AI pair), MVP first, full go-live after stakeholder sign-off"],
            ["DB target", "~70 tables (existing 64 reviewed + new added)"],
            ["Bottom line", "Industry-grade, defence-grade context, demo-ready MVP in 10 weeks. Josh: MAXED."],
        ],
    )

    add_para(doc, "", size=11)
    add_callout(
        doc,
        "END OF DOCUMENT — v3 master plan locked. "
        "Next step: confirm small additions + hand over existing DB export → enter Phase 1.",
        fill="E8F5E9",
    )

    return doc


if __name__ == "__main__":
    out_doc = build_document()
    out_path = r"e:\SOFTWAREs By DS\cmcmis-simplified\Documents\CMCMIS_Master_Plan_v3.docx"
    out_doc.save(out_path)
    print(f"Saved: {out_path}")
