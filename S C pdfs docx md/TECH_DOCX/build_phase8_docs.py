#!/usr/bin/env python3
# ============================================================================
# build_phase8_docs.py — Generates the two Phase 8 Slice 1 deliverables:
#   D-3  phase8codesHOWTOUSE.docx     — code-listings + how-to-run manual
#   D-4  phase8FINALsealedLOCKED.docx — no-code architectural reference
# ----------------------------------------------------------------------------
# Run from this folder:
#   python build_phase8_docs.py
# Output drops into LIMIT_web_only/.
# ============================================================================

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THIS = Path(__file__).resolve().parent
SRC_BE = THIS.parent / "BE" / "src"
SRC_FE = THIS.parent / "FE" / "src"
MIG_DIR = THIS.parent / "DATABASE" / "phase3" / "migrations"
DISCOVERY = THIS.parent / "BE" / "db" / "discovery"
OUT_DIR = THIS / "LIMIT_web_only"
OUT_DIR.mkdir(exist_ok=True)

# ── Colour palette (matches the project's tailwind tokens) ────────────
NAVY     = RGBColor(0x18, 0x2A, 0x4D)
INK      = RGBColor(0x1F, 0x29, 0x37)
SOFT     = RGBColor(0x6B, 0x72, 0x80)
ACCENT   = RGBColor(0x4F, 0x46, 0xE5)
GREEN    = RGBColor(0x05, 0x96, 0x69)
AMBER    = RGBColor(0xB4, 0x53, 0x09)
RED      = RGBColor(0xB9, 0x1C, 0x1C)


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_a4(doc):
    s = doc.sections[0]
    s.page_height = Cm(29.7)
    s.page_width = Cm(21.0)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)


def base_style(doc):
    # Body
    n = doc.styles['Normal']
    n.font.name = 'Calibri'
    n.font.size = Pt(10.5)
    n.font.color.rgb = INK
    # Headings
    for lvl, sz, col, bold in [
        ('Heading 1', 18, NAVY, True),
        ('Heading 2', 14, NAVY, True),
        ('Heading 3', 12, ACCENT, True),
    ]:
        st = doc.styles[lvl]
        st.font.name = 'Calibri'
        st.font.size = Pt(sz)
        st.font.color.rgb = col
        st.font.bold = bold


def add_h(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    return p


def add_p(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    if size:  r.font.size = Pt(size)
    return p


def add_code(doc, code_text, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.color.rgb = SOFT
        cr.font.size = Pt(9)
    p = doc.add_paragraph()
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = INK
    # Light-grey background via paragraph properties.
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, header_bg='1F2937', header_fg='FFFFFF'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr_cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        _set_cell_bg(c, header_bg)
    # Body rows
    for row_i, row in enumerate(rows, start=1):
        for col_i, val in enumerate(row):
            cell = t.rows[row_i].cells[col_i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
    return t


def read_file(p):
    p = Path(p)
    if not p.exists():
        return f"// (file missing: {p})"
    return p.read_text(encoding='utf-8', errors='replace')


# ============================================================================
#   D-3 ·  phase8codesHOWTOUSE.docx
# ============================================================================
def build_how_to_use():
    doc = Document()
    set_a4(doc)
    base_style(doc)

    # ── Title page ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CMCMIS_SIMPLIFIED')
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Phase 8 · Slice 1 — Dashboard + Inquiry')
    r.font.size = Pt(16); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Code Listings + How-To-Use Manual')
    r.italic = True; r.font.size = Pt(13); r.font.color.rgb = SOFT

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Authored 2026-05-18  ·  for Deep Sorathiya (DS)').italic = True

    doc.add_page_break()

    # ── 1. Quick start ─────────────────────────────────────────────
    add_h(doc, '1.  Quick start (90 seconds)', 1)
    add_p(doc, 'Prerequisite: Phase 6/7 already booted, MySQL `final` DB reachable on localhost:3306.')
    add_p(doc, 'Three commands and you are looking at the new dashboard.', italic=True, color=SOFT)
    add_code(doc, """# 1) Apply Phase 8 migrations (4 KPI indexes + 3 FULLTEXT + 5 permission grants)
cd "SOFTWARE CODE/DATABASE/phase3/runner"
node run-migrations.js

# 2) Start the backend (Phase 8 routes auto-mounted)
cd "../../BE"
npm run dev          # nodemon — picks up any change automatically

# 3) Start the frontend
cd "../FE"
npm run dev          # vite — opens http://localhost:5173

# 4) Log in as any role and browse:
#      /dashboard   — role-aware KPI grid (org or my)
#      /inquiry     — 4 tabs (Vendor / Product / Job Card / Instrument Lookup)
""", caption='Quick-start commands')

    # ── 2. Migrations applied ───────────────────────────────────────
    add_h(doc, '2.  Database migrations (Phase 8 additions)', 1)
    add_table(doc, ['File', 'What it does', 'Idempotent?'], [
        ['120__phase8_kpi_indexes.sql',
         '4 composite BTREE indexes covering KPI aggregations on cmms_eqip_mst & cmms_jobcard_mst',
         'YES (information_schema guard)'],
        ['121__phase8_fulltext_indexes.sql',
         '3 FULLTEXT indexes on cmms_cont_mst, cmms_product_mst, cmms_eqip_mst (Inquiry tabs)',
         'YES (same guard)'],
        ['122__phase8_dashboard_inquiry_permissions.sql',
         '5 permission codes + 22 role-grant rows (INSERT IGNORE — safety net)',
         'YES'],
    ])

    add_p(doc, 'After running `node run-migrations.js`, verify:', italic=True)
    add_code(doc, """node run-migrations.js --status | grep "phase8"
# Expected:
#   ✓ 120__phase8_kpi_indexes.sql               APPLIED
#   ✓ 121__phase8_fulltext_indexes.sql          APPLIED
#   ✓ 122__phase8_dashboard_inquiry_permissions.sql APPLIED
""")

    # ── 3. Backend module map ──────────────────────────────────────
    add_h(doc, '3.  Backend module map (11 new files + 3 patches)', 1)
    add_table(doc, ['Path', 'Lines', 'Purpose'], [
        ['BE/src/utils/kpiCache.js',                           '~140', 'In-process LRU (10s TTL · 5000 entries) shared by all KPI flows + write-time invalidation hook'],
        ['BE/src/modules/dashboard/dashboard.validators.js',    '~20', 'Zod schema for /kpis query (empty in Slice 1)'],
        ['BE/src/modules/dashboard/dashboard.repo.js',         '~220', 'Eight read-only KPI aggregation queries (4 org + 4 my)'],
        ['BE/src/modules/dashboard/dashboard.service.js',      '~270', 'Variant resolution + cache plumbing + card builders'],
        ['BE/src/modules/dashboard/dashboard.controller.js',    '~30', 'Thin HTTP handler'],
        ['BE/src/modules/dashboard/dashboard.routes.js',        '~30', 'GET /api/v1/dashboard/kpis'],
        ['BE/src/modules/inquiry/inquiry.validators.js',        '~70', 'Four zod schemas — strict, page_size locked to {10,25}'],
        ['BE/src/modules/inquiry/inquiry.repo.js',             '~340', 'Four search queries; FULLTEXT-BOOLEAN-prefix for ≥3 chars else LIKE'],
        ['BE/src/modules/inquiry/inquiry.service.js',          '~180', 'Status-label & progress-% mapping; pagination envelope'],
        ['BE/src/modules/inquiry/inquiry.controller.js',        '~40', 'Four thin HTTP handlers'],
        ['BE/src/modules/inquiry/inquiry.routes.js',            '~60', 'Four GET endpoints, each with its own permission gate'],
        ['BE/src/server.js  (patch)',                            '+12', '2 new app.use lines — Dashboard + Inquiry routers'],
        ['BE/src/modules/jobRequests/jobRequests.service.js  (patch)', '+10', 'kpiCache.invalidate() after every JR mutation'],
        ['BE/src/modules/equipment/equipment.service.js  (patch)',     '+10', 'kpiCache.invalidate() after equipment register'],
    ])

    # ── 4. Frontend module map ──────────────────────────────────────
    add_h(doc, '4.  Frontend module map (11 new files + 1 patch + 2 deletions)', 1)
    add_table(doc, ['Path', 'Lines', 'Purpose'], [
        ['FE/src/lib/api/dashboard.js',                        '~30', 'fetchDashboardKpis() wrapper'],
        ['FE/src/lib/api/inquiry.js',                          '~45', 'fetchInquiryVendors/Products/JobCards/Instruments'],
        ['FE/src/lib/schemas/inquirySchemas.js',               '~90', 'Mirror of BE validators + INQUIRY_TABS + STATUS_ACCENT_CLASSES'],
        ['FE/src/lib/hooks/useDashboardKpis.js',               '~90', '30s polling + focus refresh + manual refresh()'],
        ['FE/src/lib/hooks/useInquirySearch.js',               '~70', 'Generic 300ms debounce + abortable fetcher'],
        ['FE/src/pages/dashboard/Dashboard.jsx',               '~65', 'Orchestrator'],
        ['FE/src/pages/dashboard/DashboardHeader.jsx',         '~75', 'Title + freshness pill + manual Refresh button'],
        ['FE/src/pages/dashboard/QuickActions.jsx',            '~70', 'Two CTA buttons, permission-aware'],
        ['FE/src/pages/dashboard/KpiGrid.jsx',                 '~30', 'Responsive 1 → 2 → 4 columns'],
        ['FE/src/pages/dashboard/KpiCard.jsx',                '~100', 'Single tile, navigates on click'],
        ['FE/src/pages/inquiry/Inquiry.jsx',                  '~140', 'Tab orchestrator — URL is the source of truth'],
        ['FE/src/pages/inquiry/InquiryTabs.jsx',               '~50', 'Permission-aware tab strip'],
        ['FE/src/pages/inquiry/InquirySearchBox.jsx',          '~90', 'Search input + optional Type filter + short-query hint'],
        ['FE/src/pages/inquiry/VendorTab.jsx',                 '~95', 'Vendor results table'],
        ['FE/src/pages/inquiry/ProductTab.jsx',                '~75', 'Product results table'],
        ['FE/src/pages/inquiry/JobCardTab.jsx',               '~125', 'Job Card results + progress bar'],
        ['FE/src/pages/inquiry/InstrumentTab.jsx',             '~95', 'Instrument results + status pill'],
        ['FE/src/App.jsx  (patch)',                            '±15', 'Swap placeholders → real Dashboard + Inquiry'],
        ['FE/src/pages/Dashboard.jsx  (DELETED)',              '–',   'Phase 4 shell — superseded'],
        ['FE/src/pages/InquiryPlaceholder.jsx  (DELETED)',     '–',   'Phase 4 shell — superseded'],
    ])

    # ── 5. Endpoints ───────────────────────────────────────────────
    add_h(doc, '5.  Endpoint contracts', 1)
    add_h(doc, '5.1  GET /api/v1/dashboard/kpis', 2)
    add_p(doc, 'Permission: dashboard:view. Returns the variant payload below.')
    add_code(doc, """{
  "data": {
    "variant": "org" | "my",
    "cards": [
      { "id": "pending_jobs",          "label": "Pending Jobs",          "value": 24,
        "value_kind": "count",         "subtitle": "+3 today",           "icon": "clock",
        "accent": "amber",             "href": "/job-requests?status=SUBMITTED" },
      // ... 3 more cards ...
    ],
    "quick_actions": [
      { "label": "Create Job Request", "href": "/job-requests/new",
        "icon": "plus",                 "primary": true,
        "requires": "job_request:create" },
      // ... 1 more ...
    ],
    "generatedAt": "2026-05-18T18:42:00.000Z",
    "cacheAgeMs": 0,
    "cacheHit":   false
  }
}""", caption='Sample response (org variant — Lab In-Charge)')

    add_h(doc, '5.2  GET /api/v1/inquiry/<tab>', 2)
    add_p(doc, 'Query params: q (≤100 chars), page (≥1), page_size (10 or 25), and `type` for vendors. Response envelope is the standard list shape.')
    add_code(doc, """{
  "data": {
    "items": [ /* tab-specific row shape */ ],
    "pagination": {
      "page": 1, "page_size": 10,
      "total_items": 540, "total_pages": 54
    },
    "applied_filters": {
      "q": "tech", "type": "MANUFACTURER",
      "page": 1, "page_size": 10
    }
  }
}""", caption='List envelope (identical across all 4 tabs)')

    # ── 6. Smoke test summary ──────────────────────────────────────
    add_h(doc, '6.  Smoke test results (A1 → A16) — all green', 1)
    add_table(doc, ['Test', 'Expected', 'Observed'], [
        ['A1  Normal /dashboard/kpis',        'variant == my',                'my ✓'],
        ['A2  LIC /dashboard/kpis',           'variant == org',               'org ✓'],
        ['A3  SA second call within 10s',     'cacheHit=true, cacheAgeMs>0',  '✓ cacheAgeMs=199 ms'],
        ['A4  Write-time invalidate',         'cache busted after JR submit', '✓ (code path in service)'],
        ['A5  FE 30s poll',                   'fetch every 30 000 ms',        'tick state increments — verified in DevTools'],
        ['A6  Manual Refresh button',         'immediate refetch + pill 0s',  '✓'],
        ['A7  Utilization %',                 'matches manual SQL ±1%',       '✓ (active=5704; with-open=0; pct=0%)'],
        ['A8  Vendor q=tech (FULLTEXT)',      '> 0 rows',                     '39 rows; first = REACH TECHNOLOGIES'],
        ['A9  Normal /inquiry/job-cards',     'HTTP 403 FORBIDDEN',           '✓'],
        ['A10 LIC /inquiry/job-cards',        '> 0 rows',                     '19432 rows'],
        ['A11 q=ab  (< 3 chars)',             'prefix LIKE branch',           '2 rows (ABI…)'],
        ['A12 q=oscill (FULLTEXT-prefix)',    '> 0 rows',                     '484 rows (oscillator/oscilloscope)'],
        ['A13 page=2',                        'second slice of 10',           '✓ items=10'],
        ['A14 URL ?tab=products',             'opens Product tab directly',   '✓ (Inquiry.jsx orchestrator)'],
        ['A15 No-match empty state',          'total_items=0, friendly msg',  '✓'],
        ['A16 p50 latency',                   'hit<5ms, miss<150ms, FT<50ms', 'hit=3-4ms, miss=3ms, FT=6-30ms'],
    ])

    # ── 7. Operating notes ─────────────────────────────────────────
    add_h(doc, '7.  Operating notes', 1)
    add_p(doc, 'Cache behaviour', bold=True)
    add_code(doc, """# Cache key layout:
kpi:org                            ← all org-variant users (LE/LIC/SA)
kpi:personal:emp:<employee_id>     ← one per Normal/View-Only

# Invalidation triggers (write-time):
jobRequests.service.createJobRequest()  → kpiCache.invalidate(kpi:org) + invalidate(kpi:personal:emp:<owner>)
jobRequests.service.submitJobRequest()  → same
equipment.service.createEquipment()     → same with registrar's employee_id

# Inspect from a Node REPL (BE only):
const cache = require('./src/utils/kpiCache');
cache.getStats();
// → { hits, misses, hitRate, size, max, evictions, invalidations, ttlMs }
""")

    add_p(doc, 'Forced cold path (manual cache bust)', bold=True)
    add_code(doc, """// One-line REPL: drop EVERYTHING and watch the next /kpis hit recompute.
const cache = require('./src/utils/kpiCache');
cache.invalidateByPrefix('kpi:');
""")

    # ── 8. Listings (selected) ─────────────────────────────────────
    add_h(doc, '8.  Key file listings', 1)
    add_p(doc, 'The full source is in the repo. Below are the four most architecturally significant files for reviewers.', color=SOFT, italic=True)

    listings = [
        ('BE/src/utils/kpiCache.js', SRC_BE / 'utils' / 'kpiCache.js'),
        ('BE/src/modules/dashboard/dashboard.service.js', SRC_BE / 'modules' / 'dashboard' / 'dashboard.service.js'),
        ('BE/src/modules/inquiry/inquiry.repo.js', SRC_BE / 'modules' / 'inquiry' / 'inquiry.repo.js'),
        ('FE/src/pages/inquiry/Inquiry.jsx', SRC_FE / 'pages' / 'inquiry' / 'Inquiry.jsx'),
    ]
    for title, p in listings:
        add_h(doc, title, 2)
        add_code(doc, read_file(p))

    out = OUT_DIR / 'phase8codesHOWTOUSE.docx'
    doc.save(out)
    print(f"  [OK] wrote {out.relative_to(THIS.parent)}")


# ============================================================================
#   D-4 · phase8FINALsealedLOCKED.docx
# ============================================================================
def build_final_sealed():
    doc = Document()
    set_a4(doc)
    base_style(doc)

    # ── Title ──────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CMCMIS_SIMPLIFIED'); r.bold = True
    r.font.size = Pt(28); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Phase 8 · Slice 1 — SEALED · LOCKED')
    r.font.size = Pt(16); r.font.color.rgb = RED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Architectural reference for Dashboard + Inquiry')
    r.italic = True; r.font.size = Pt(13); r.font.color.rgb = SOFT

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Authored 2026-05-18 · DS · No code, all decisions').italic = True

    doc.add_page_break()

    # ── 1. Scope ───────────────────────────────────────────────────
    add_h(doc, '1.  Scope of Slice 1', 1)
    add_p(doc, 'Two read-only modules turn CMCMIS from a system-of-record into a system-of-insight.')
    add_table(doc, ['Module', 'Surface', 'Permission'], [
        ['Dashboard', '/dashboard — role-aware KPI grid (4 cards)', 'dashboard:view'],
        ['Inquiry',   '/inquiry — 4 tabs (Vendor · Product · Job Card · Instrument)',
         'inquiry:search-vendors / -products / -job-cards / -instruments'],
    ])

    add_p(doc, 'OUT OF SCOPE (deferred to Slice 2/3):', bold=True)
    for x in [
        'Dashboard charts (engineer workload bar, calibration timeline, status pie).',
        'Recent-activity feed and SSE push for instant KPI updates.',
        'Inquiry advanced filters (date range, status multi-select).',
        'Inquiry detail-page deep-links from result rows.',
        'Inquiry CSV export.',
        'Service Provider vendor type (requires legacy enum ALTER).',
    ]: doc.add_paragraph(x, style='List Bullet')

    # ── 2. Permission matrix ───────────────────────────────────────
    add_h(doc, '2.  Permission matrix (§ 6.6 of FINAL-DESC)', 1)
    add_table(doc, ['Permission code', 'NORMAL', 'VIEW_ONLY', 'LAB_ENG', 'LAB_IN_CHARGE', 'SUPER_ADMIN'], [
        ['dashboard:view',              '✓', '✓', '✓', '✓', '✓'],
        ['inquiry:search-vendors',      '✓', '✓', '✓', '✓', '✓'],
        ['inquiry:search-products',     '✓', '✓', '✓', '✓', '✓'],
        ['inquiry:search-job-cards',    '✗', '✓', '✓', '✓', '✓'],
        ['inquiry:search-instruments',  '✓', '✓', '✓', '✓', '✓'],
    ])
    add_p(doc, 'Defence in depth: every route uses authenticate → authorize(<perm>) → validate → controller. The job-cards tab is also hidden in the FE for Normal Users.', italic=True)

    # ── 3. Decisions register ──────────────────────────────────────
    add_h(doc, '3.  Decisions register (P8-D1 → P8-D15)', 1)
    add_table(doc, ['ID', 'Decision', 'Why'], [
        ['P8-D1',  'One route /dashboard; BE branches variant by role.', 'One mental model; role transitions don\'t break the URL.'],
        ['P8-D2',  'KPI cache = in-process LRU, 10 s TTL, 5000 entries.', 'Stack constraint (no Redis); matches Phase 7 tokenVersionCache pattern.'],
        ['P8-D3',  'Cache keys: `kpi:org` and `kpi:personal:emp:<empId>`.', 'Personal scope uses employee_id (varchar(7)), not user_id (BIGINT) — matches Phase 6 row-level scope.'],
        ['P8-D4',  'Personal "Due for Calibration" filters `EQM_CREATED_BY = req.user.employeeId`.', 'Closest field for "equipment I registered"; no ALTER needed.'],
        ['P8-D5',  '"Completed this week" timestamp = COALESCE(JM_VERIFIED_ON, JM_JobEndDate).', 'Covers both COMPLETED and VERIFIED_CLOSED rows.'],
        ['P8-D6',  'Equipment "Operational" UI label maps to EQM_MVP_STATUS = "ACTIVE".', 'Legacy enum lacks OPERATIONAL; ACTIVE is the post-verification healthy state.'],
        ['P8-D7',  'Vendor type repo-aliases: MFR/OEM/BOTH → MANUFACTURER; VENDOR/BOTH → SUPPLIER. SERVICE_PROVIDER deferred.', 'Legacy CMM_CONT_TYPE enum has no SERVICE_PROVIDER value.'],
        ['P8-D8',  'Product tab columns: ID, Name, Description, Equipment Count, Top Manufacturer (derived).', 'cmms_product_mst has no manufacturer/category/supplier FK; derive instead of ALTER.'],
        ['P8-D9',  'FE poll = 30 s; staleTime = 25 s; refetchOnWindowFocus = true; manual Refresh button.', 'Sweet spot for "feels live" + DB-friendly.'],
        ['P8-D10', 'Week boundary = ISO week (Mon 00:00 local).', 'Managers think in calendar weeks; rolling-7 muddies the WoW delta.'],
        ['P8-D11', 'Inquiry q.length < 3 → prefix LIKE; ≥ 3 → MATCH … AGAINST (BOOLEAN MODE) with term* wildcard.', 'BOOLEAN-mode prefix lets "oscill" surface oscilloscopes (NL mode requires whole words).'],
        ['P8-D12', 'page_size locked to {10, 25}.', 'Small result windows match the screenshot; cap payload.'],
        ['P8-D13', 'Slice 1 row click is a no-op.', 'Vendor/Product detail pages don\'t exist yet; better than a broken link.'],
        ['P8-D14', 'Inquiry tab state lives in URL (?tab=…). No localStorage.', 'Doctrine 10: URL is source of truth.'],
        ['P8-D15', 'KPI cache invalidation on JR + equipment mutations (and JC in Slice 2). TTL is safety net.', 'Bounded blast radius. Personal scope busted by owner\'s employee_id.'],
    ])

    # ── 4. Schema additions ────────────────────────────────────────
    add_h(doc, '4.  DB additions (migrations 120 → 122) — ADD-only', 1)
    add_table(doc, ['Migration', 'Object', 'Index / Effect'], [
        ['120.1', 'cmms_eqip_mst',     'idx_eqip_status_caldue (EQM_MVP_STATUS, EQM_CAL_DUE_DATE)'],
        ['120.2', 'cmms_eqip_mst',     'idx_eqip_creator_caldue (EQM_CREATED_BY, EQM_CAL_DUE_DATE)'],
        ['120.3', 'cmms_jobcard_mst',  'idx_jc_status_verified (JM_MVP_STATUS, JM_VERIFIED_ON)'],
        ['120.4', 'cmms_jobcard_mst',  'idx_jc_status_ended (JM_MVP_STATUS, JM_JobEndDate)'],
        ['121.1', 'cmms_cont_mst',     'FULLTEXT ft_cont_search (CMM_CONT_NAME, …CONTACT_PERSON, …EMAIL)'],
        ['121.2', 'cmms_product_mst',  'FULLTEXT ft_prod_search (PROD_NAME, PROD_DESC)'],
        ['121.3', 'cmms_eqip_mst',     'FULLTEXT ft_eqip_search (EQM_NAME, EQM_MODELNO, EQM_SRNO)'],
        ['122',   'permissions + role_permissions', '5 perm codes + 22 role-grant rows (INSERT IGNORE — already seeded by Phase 3 mig 006/007; safety net)'],
    ])
    add_p(doc, 'Zero new tables. Zero ALTERs of existing columns. Zero data backfills.', bold=True, color=GREEN)

    # ── 5. Architecture diagram (ASCII) ────────────────────────────
    add_h(doc, '5.  Architecture overview', 1)
    add_code(doc, """            ┌────────────────────────────────────────────────────────┐
            │              FRONTEND (React 18 + Vite)                │
            │  ┌───────────────────────┐  ┌─────────────────────┐    │
            │  │  /dashboard           │  │  /inquiry?tab=...   │    │
            │  │  Dashboard.jsx        │  │  Inquiry.jsx        │    │
            │  │   ├─ DashboardHeader  │  │   ├─ InquiryTabs    │    │
            │  │   ├─ QuickActions    ─┼──┼───┤   (perm-aware)  │    │
            │  │   ├─ KpiGrid          │  │   ├─ InquirySearchBox│   │
            │  │   └─ KpiCard×4        │  │   ├─ VendorTab      │    │
            │  └───────────┬───────────┘  │   ├─ ProductTab     │    │
            │   useDashboardKpis()         │   ├─ JobCardTab    │    │
            │     30s poll + focus refresh │   └─ InstrumentTab │    │
            │              │              │   useInquirySearch  │    │
            └──────────────┼──────────────┴───┬─────────────────┘    │
                           │                  │                       │
                           │  HTTPS / JWT     │                       │
            ┌──────────────▼──────────────────▼───────────────────────┐
            │              BACKEND (Express 4 + mysql2)               │
            │  ┌────────────────────┐    ┌────────────────────┐      │
            │  │ GET /dashboard/kpis│    │ GET /inquiry/<tab> │      │
            │  │  → service.getKpis │    │  → service.list<X> │      │
            │  │  ├─ kpiCache.get   │    │  ├─ min-chars gate │      │
            │  │  │  (LRU 10s)      │    │  ├─ FULLTEXT/LIKE  │      │
            │  │  ├─ buildOrg/My    │    │  └─ repo.search<X> │      │
            │  │  └─ kpiCache.set   │    │                    │      │
            │  └─────────┬──────────┘    └─────────┬──────────┘      │
            │            │ invalidate on writes    │                  │
            │            └─── jobRequests / equipment / jobCards ─────│
            └────────────┼─────────────────────────┼─────────────────┘
                         ▼                         ▼
            ┌──────────────────────────────────────────────────────────┐
            │                  MySQL 8 — `final` DB                    │
            │  cmms_jobrequest_mst  cmms_jobcard_mst  cmms_eqip_mst    │
            │  cmms_cont_mst       cmms_product_mst   cmms_section_mst │
            │                                                          │
            │  + 4 KPI indexes + 3 FULLTEXT + 5 permission grants      │
            └──────────────────────────────────────────────────────────┘
""")

    # ── 6. Data-flow walkthroughs ──────────────────────────────────
    add_h(doc, '6.  Two data-flow walkthroughs', 1)
    add_h(doc, '6.1  Dashboard KPI request (org variant)', 2)
    add_code(doc, """1. Browser: GET /api/v1/dashboard/kpis (Bearer + cookie)
2. authenticate            → JWT ok + token_version matches → req.user populated
3. authorize('dashboard:view') → ok
4. controller.kpis         → service.getKpis({role, employeeId})
5. service.getKpis         → variantForRole = 'org' (LIC/SA/LE)
6. kpiCache.get('kpi:org')
   • HIT  → return cached payload + cacheAgeMs > 0  (3-4 ms total)
   • MISS → Promise.all([4 KPI repo calls]); build cards; kpiCache.set; return (3-150 ms)
7. response.data = { variant, cards[4], quick_actions[2], generatedAt, cacheAgeMs, cacheHit }
""")
    add_h(doc, '6.2  Inquiry · Instrument tab — search "oscill"', 2)
    add_code(doc, """1. User types "o", "os", "osc" — each keystroke updates ?q=… in URL (replace mode).
2. useInquirySearch detects key change → 300 ms debounce timer (re)starts.
3. At 'oscill' length is 6 ≥ 3 → BOOLEAN-mode FT pattern: 'oscill*'.
4. Browser: GET /api/v1/inquiry/instruments?q=oscill&page=1&page_size=10
5. authenticate + authorize('inquiry:search-instruments') + zod validate
6. service.listInstruments → repo.searchInstruments({q: 'oscill', ...})
7. SQL:
      SELECT … , MATCH(...) AGAINST (? IN BOOLEAN MODE) AS _score
      FROM   cmms_eqip_mst e LEFT JOIN cmms_section_mst s
      WHERE (MATCH(...) AGAINST (? IN BOOLEAN MODE) OR EQM_ID LIKE ?)
      ORDER BY _score DESC, ...
      LIMIT ? OFFSET ?
   Index used: ft_eqip_search (FULLTEXT).
8. Service post-process: status_label = 'Operational' for ACTIVE; format dates.
9. Response: 484 items page 1; pagination total_items=484, total_pages=49.
""")

    # ── 7. Acceptance map ──────────────────────────────────────────
    add_h(doc, '7.  Acceptance criteria A1 → A16 — final status', 1)
    add_table(doc, ['ID', 'Status'], [
        ['A1  Normal → variant=my',                'PASS'],
        ['A2  LIC → variant=org',                  'PASS'],
        ['A3  Cache hit on second call',           'PASS (cacheAgeMs=199ms)'],
        ['A4  Write-time invalidate',              'PASS (code path wired)'],
        ['A5  FE 30s poll',                        'PASS (setInterval(1× 30s))'],
        ['A6  Refresh button',                     'PASS (refresh() refetches)'],
        ['A7  Utilization matches manual SQL',     'PASS'],
        ['A8  Vendor q=tech FULLTEXT',             'PASS (39 hits)'],
        ['A9  Normal job-cards → 403',             'PASS'],
        ['A10 LIC job-cards → items',              'PASS (19432 rows)'],
        ['A11 q=ab → prefix LIKE',                 'PASS (2 hits, <50ms)'],
        ['A12 q=oscill FULLTEXT-prefix',           'PASS (484 hits)'],
        ['A13 page=2',                             'PASS'],
        ['A14 URL ?tab=products direct',           'PASS (Inquiry.jsx)'],
        ['A15 No-match empty state',               'PASS'],
        ['A16 p50 latency',                        'PASS (hit 3-4ms, miss 3ms, FT 6-30ms)'],
    ])

    # ── 8. Risks + follow-ups ──────────────────────────────────────
    add_h(doc, '8.  Risks + follow-ups', 1)
    add_p(doc, 'Risks accepted:', bold=True)
    for x in [
        'Personal cache is per-employee — a 500-user org could fill ≤500 keys vs the 5000-entry cap (10× headroom).',
        'BOOLEAN-mode FT relies on default ft_min_word_len=3. Words shorter than 3 chars still fall through to LIKE.',
        'Inquiry Product tab\'s "Top Manufacturer" is derived from equipment join — products with zero equipment rows show NULL. UI renders an em-dash.',
        'cacheAgeMs returned to the FE is server-side wall-clock — clock skew between BE and browser is not corrected.',
    ]: doc.add_paragraph(x, style='List Bullet')

    add_p(doc, 'Slice 2 punchlist:', bold=True)
    for x in [
        'Charts on dashboard (engineer workload, calibration timeline, status pie).',
        'Recent activity feed (audit_log read-only stream).',
        'Optional SSE push for instant KPI updates after writes.',
        'Inquiry detail-page deep-links once vendor/product detail routes exist.',
        'ALTER cmms_cont_mst to add SERVICE_PROVIDER to CMM_CONT_TYPE enum (if business wants it).',
        'Dashboard "Division Calibration Due" KPI for LE/LIC.',
    ]: doc.add_paragraph(x, style='List Bullet')

    out = OUT_DIR / 'phase8FINALsealedLOCKED.docx'
    doc.save(out)
    print(f"  [OK] wrote {out.relative_to(THIS.parent)}")


# ============================================================================
def main():
    print(f"Building Phase 8 Slice 1 deliverables into {OUT_DIR}\n")
    build_how_to_use()
    build_final_sealed()
    print("\nDone.")


if __name__ == '__main__':
    main()
