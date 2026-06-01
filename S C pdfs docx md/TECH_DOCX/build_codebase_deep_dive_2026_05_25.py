from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "CMCMIS_END_TO_END_CODEBASE_DEEP_DIVE_2026-05-25.docx"
MD_PATH = OUT_DIR / "CMCMIS_END_TO_END_CODEBASE_DEEP_DIVE_2026-05-25.md"
MEMORY_PATH = OUT_DIR / "CODEBASE_MEMORY_2026-05-25.md"


COLORS = {
    "GREEN": "D9EAD3",
    "BLUE": "D9EAF7",
    "YELLOW": "FFF2CC",
    "RED": "F4CCCC",
    "PURPLE": "EADCF8",
    "GRAY": "E8EEF5",
    "WHITE": "FFFFFF",
}


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C9D3DF", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, size=None, bold=False, color=None, font=None):
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if font:
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_colored_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    if level == 1:
        style_run(run, size=16, bold=True, color="2E74B5")
    elif level == 2:
        style_run(run, size=13, bold=True, color="2E74B5")
    else:
        style_run(run, size=12, bold=True, color="1F4D78")
    return paragraph


def add_para(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = paragraph.add_run(bold_prefix)
        style_run(r1, bold=True)
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1
    set_table_borders(table, color="D6DEE8")
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FC")
    set_cell_margins(cell, 120, 160, 120, 160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.strip("\n"))
    style_run(run, size=8.5, font="Courier New", color="1F2937")
    return table


def add_table(doc, headers, rows, status_col=None, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = 1
    table.style = "Table Grid"
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        shade_cell(header_cells[i], COLORS["GRAY"])
        set_cell_margins(header_cells[i])
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                style_run(run, bold=True, color="1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[i])
            if status_col is not None and i == status_col:
                shade_cell(cells[i], COLORS.get(str(value).upper(), COLORS["WHITE"]))
                for paragraph in cells[i].paragraphs:
                    for run in paragraph.runs:
                        style_run(run, bold=True)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    return doc


def build_docx():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CMCMIS End-to-End Codebase Deep Dive")
    style_run(r, size=22, bold=True, color="1F4D78")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Frontend + Backend + Database + Workflow Memory Report")
    style_run(r, size=13, bold=True, color="2E74B5")
    add_para(doc, "Prepared on 2026-05-25 after deep local inspection of FE, BE, DATABASE, TECH_DOCX, PDFs, temp, and the live MariaDB/phpMyAdmin-backed database.")
    add_table(
        doc,
        ["Color", "Meaning", "How to read it"],
        [
            ("GREEN", "Verified / healthy", "Evidence confirms the behavior or connection is currently working."),
            ("BLUE", "Core architecture", "Stable design pattern or system structure to remember."),
            ("YELLOW", "Attention / improvement", "Not broken, but worth tracking before future expansion."),
            ("RED", "Risk / blocker", "Could break behavior or needs action before production confidence."),
            ("PURPLE", "Memory / style", "A rule, convention, or user preference I should preserve."),
        ],
        status_col=0,
        widths=[1.1, 1.8, 4.2],
    )
    add_code_block(
        doc,
        """
PROJECT VOCABULARY REMEMBERED
FE       = frontend folder
BE       = backend folder
DATABASE = database folder and SQL/migration assets
Style    = color-coded + tables + diagrams + beginner-to-advanced flow
DB       = MariaDB/MySQL through phpMyAdmin, active DB cmcmis_simplified
        """,
    )

    page_break(doc)
    add_colored_heading(doc, "Manual Contents", 1)
    contents = [
        "1. Executive Summary",
        "2. Project Identity and Folder Map",
        "3. Current Verification Snapshot",
        "4. End-to-End Architecture",
        "5. Frontend Deep Dive",
        "6. Backend Deep Dive",
        "7. Database Deep Dive",
        "8. Core Workflow Diagrams",
        "9. Security, RBAC, Audit, and Compliance",
        "10. Risks, Gaps, and Recommended Roadmap",
        "11. Long-Term Codebase Memory",
        "12. Appendices: inventories, counts, modules, routes",
    ]
    for item in contents:
        add_para(doc, item)

    page_break(doc)
    add_colored_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "CMCMIS is an industry-style Computerized Maintenance and Calibration Management Information System. "
        "The codebase is not only a CRUD app: it coordinates equipment records, job requests, job cards, approvals, role-based permissions, audit history, notifications, analytics, reports, schedule, procurement, and PDF generation."
    )
    add_table(
        doc,
        ["Status", "Area", "Insight", "Evidence / Description"],
        [
            ("GREEN", "Database", "Live MariaDB connectivity is confirmed.", "BE environment connects to cmcmis_simplified on localhost:3306; phpMyAdmin route responds for the same DB."),
            ("GREEN", "Frontend", "Production build completed.", "Vite build succeeded with 2546 transformed modules."),
            ("GREEN", "Backend", "DB pool boot check completed.", "mysql2 pool validated SELECT 1 and reported DB pool ready."),
            ("BLUE", "Architecture", "Clear FE -> BE -> repository -> MariaDB shape.", "React/Vite UI calls Express APIs; BE uses service/repository modules and raw SQL."),
            ("BLUE", "Domain", "The system is workflow/state driven.", "Job Requests and Job Cards have explicit state machines and audit trails."),
            ("YELLOW", "Maintainability", "Frontend bundle is large.", "Vite warns main JS chunk is over 500 kB; future lazy loading/manual chunks should be considered."),
            ("YELLOW", "API polish", "Some reserved/stub routes remain.", "Equipment detail/update/verify/condemn/delete stubs and one legacy job-card PDF route should be tracked."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "2. Project Identity", 1)
    add_table(
        doc,
        ["Color", "Item", "Current understanding"],
        [
            ("BLUE", "Project name", "CMCMIS simplified codebase for maintenance and calibration management."),
            ("BLUE", "Target environment", "On-prem/private infrastructure style; MySQL/MariaDB database accessible through phpMyAdmin."),
            ("BLUE", "User model", "Employee-id login, one primary role, RBAC through permission codes."),
            ("BLUE", "Primary modules", "Equipment, Job Requests, Conversion, Job Cards, Schedule, Procurement, Reports, Analytics, Inquiry, Notifications, Admin, Audit."),
            ("PURPLE", "User style", "Use color-coded tables, diagrams, and beginner-to-advanced explanations."),
            ("PURPLE", "Naming memory", "Always call FE the frontend folder, BE the backend folder, and DATABASE the database folder."),
        ],
        status_col=0,
    )
    add_para(
        doc,
        "Important mental model: the project has a legacy-plus-MVP database strategy. Legacy CMMS tables still exist and are read, while newer MVP tables, ALTERs, history tables, and permissions make the modern workflow safer."
    )

    page_break(doc)
    add_colored_heading(doc, "3. Folder Map and Inventory", 1)
    add_table(
        doc,
        ["Folder", "Files inspected", "Role in system", "Notes"],
        [
            ("BE", "151", "Backend API, middleware, modules, PDF generators, DB config.", "139 JavaScript files; modular Express architecture."),
            ("FE", "159", "Frontend React app, pages, hooks, API clients, UI shell.", "96 JSX files and 58 JS files; permission-routed app."),
            ("DATABASE", "53", "Migrations, runner, SQL assets, schema documentation.", "41 SQL migrations/assets plus runner scripts."),
            ("PDFs", "35", "Legacy/reference PDFs and SQL/document artifacts.", "Large historical assets; useful reference but not app runtime code."),
            ("TECH_DOCX", "21 before this report", "Project documentation and generated technical docs.", "This new report is placed here."),
            ("temp", "39", "Temporary/research/schema artifacts.", "Includes discovery docs and SQL/PDF temporary outputs."),
        ],
    )
    add_table(
        doc,
        ["Codebase slice", "File count", "Line count / scale", "Meaning"],
        [
            ("BE .js", "139", "25,039 lines", "Backend is the largest application code slice."),
            ("FE .jsx", "96", "15,822 lines", "UI and page workflows are significant and feature-rich."),
            ("FE .js", "58", "4,108 lines", "API clients, hooks, utilities, permissions, stores."),
            ("DATABASE .sql", "41", "832,281 lines", "Database folder contains substantial SQL/migration history."),
            ("temp .sql", "5", "2,037,923 lines", "Temporary SQL exports are very large and should be treated as artifacts."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "4. Current Verification Snapshot", 1)
    add_table(
        doc,
        ["Status", "Check", "Result", "Interpretation"],
        [
            ("GREEN", "MariaDB server", "Reachable at localhost:3306", "The database server is available from this environment."),
            ("GREEN", "Active DB", "cmcmis_simplified", "BE .env points to this database."),
            ("GREEN", "phpMyAdmin", "Database structure URL responded with title containing cmcmis_simplified", "Browser-level DB administration route is reachable."),
            ("GREEN", "Backend DB pool", "DB pool ready; poolLimit 15", "BE config and credentials are valid."),
            ("GREEN", "Frontend build", "npm run build succeeded", "FE compiles successfully."),
            ("YELLOW", "Frontend build warning", "Main JS chunk around 1,257.67 kB", "Consider lazy routes/chunk splitting later."),
            ("YELLOW", "DATABASE phase3 .env", "Points to final, while BE points to cmcmis_simplified", "Treat DATABASE env as runner/context-specific unless intentionally updated."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "5. End-to-End Architecture", 1)
    add_para(
        doc,
        "The simplest useful view is: user action starts in the frontend, moves through the Axios API client, reaches Express middleware, enters a module controller, passes business validation in a service, reads/writes through a repository, and persists in MariaDB."
    )
    add_code_block(
        doc,
        """
[User]
  |
  v
[FE React/Vite frontend]
  |  Axios + in-memory access token + CSRF header
  v
[BE Express backend]
  |  helmet/cors/compression/json/cookies/logger/auth/RBAC
  v
[Module Controller]
  v
[Service: validation + workflow + audit intent]
  v
[Repository: raw SQL via mysql2 pool]
  v
[MariaDB database: cmcmis_simplified]

[phpMyAdmin] ---------------------------> [MariaDB]
[BE PDFKit routes] ---------------------> [On-demand PDF response]
[BE multer storage] <-------------------- [Job Card document uploads]
        """,
    )
    add_table(
        doc,
        ["Layer", "Primary responsibility", "Remembered design rule"],
        [
            ("Frontend", "Workflow screens, permission routing, form UX, query hooks.", "Keep FE as React/Vite/Tailwind JavaScript, no TypeScript unless project direction changes."),
            ("Backend", "Security, routing, state transitions, audit, SQL access, PDF streaming.", "Use module folders with route/controller/service/repository split."),
            ("Database", "Legacy and MVP tables, permissions, audit logs, migrations.", "Use migrations and phpMyAdmin/MySQL inspection rather than guessing schema."),
            ("Documentation", "Phase docs, discovery docs, generated reports.", "TECH_DOCX is the correct place for polished generated deliverables."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "6. Frontend Deep Dive: Stack", 1)
    add_table(
        doc,
        ["Status", "Frontend item", "Description"],
        [
            ("BLUE", "Framework", "React 18 with Vite."),
            ("BLUE", "Styling", "Tailwind CSS v3 with component-level UI files."),
            ("BLUE", "Routing", "react-router-dom v6 with protected permission routes."),
            ("BLUE", "Data layer", "Axios clients plus React Query-style hooks for server state."),
            ("BLUE", "Forms", "react-hook-form and zod are available for validation-heavy flows."),
            ("BLUE", "UI libraries", "lucide-react for icons, sonner for toasts, recharts for charts."),
            ("GREEN", "Build", "Production build succeeded."),
            ("YELLOW", "Bundle size", "Large main bundle should be optimized after feature stabilization."),
        ],
        status_col=0,
    )
    add_para(
        doc,
        "The frontend is already organized by domain pages, API files, reusable hooks, layout components, auth handling, and permission utilities. The architecture is readable and maps closely to backend modules."
    )

    page_break(doc)
    add_colored_heading(doc, "7. Frontend Routes and Permissions", 1)
    add_table(
        doc,
        ["Route", "Feature", "Required permission"],
        [
            ("/dashboard", "Dashboard", "dashboard:view"),
            ("/equipment", "Equipment list", "equipment:read-list"),
            ("/equipment/new", "New equipment", "equipment:create"),
            ("/job-requests", "Job Request list", "job_request:read-own"),
            ("/job-requests/new", "New Job Request", "job_request:create"),
            ("/job-requests/:id", "Job Request detail", "job_request:read-own"),
            ("/conversion", "Request conversion/approval", "job_request:approve"),
            ("/job-cards", "Job Card list", "job_card:read-list"),
            ("/job-cards/:id", "Job Card detail", "job_card:read-detail"),
            ("/schedule", "Schedule", "schedule:read-list"),
            ("/procurement", "Procurement", "procurement:read-list"),
            ("/inquiry", "Inquiry/global search", "inquiry:search-vendors / search-instruments context"),
            ("/reports", "Reports", "reports:view-analytics"),
            ("/analytics", "Analytics", "analytics:view"),
            ("/notifications", "Notifications", "notifications:read-own"),
            ("/admin/users", "Admin users", "user:read-list"),
            ("/admin/employees", "Admin employees", "master:employees:manage"),
            ("/audit", "Audit log", "audit:read-list"),
        ],
    )
    add_para(
        doc,
        "This permission-first route layout matches the backend rule: UI visibility and API access should be based on permission codes, not direct role-name branching."
    )

    page_break(doc)
    add_colored_heading(doc, "8. Frontend Shell and UX Flow", 1)
    add_code_block(
        doc,
        """
[App.jsx]
  |
  +-- [ProtectedRoute]
  |      checks auth and required permission
  |
  +-- [Layout.jsx]
         |
         +-- [Sidebar.jsx] permission-filtered navigation
         +-- [TopBar.jsx] global search, notification bell, account menu
         +-- [TokenHost] centralized success/error capsules
         +-- [Outlet] active page
        """,
    )
    add_table(
        doc,
        ["Component / pattern", "Purpose", "Memory"],
        [
            ("Layout.jsx", "Main application shell.", "Sidebar collapsed state is persisted in localStorage key cmcmis.sidebar.collapsed."),
            ("Sidebar.jsx", "Permission-filtered navigation.", "Navigation is data-driven from permission utilities."),
            ("TopBar.jsx", "Search, notifications, account actions.", "Global search routes to /inquiry?q=..."),
            ("TokenHost / tokenInterceptor", "Friendly mutation result capsules.", "Good UX pattern, but route rules should be audited as endpoints evolve."),
            ("DataTable.jsx", "Reusable table shell.", "Project likes table-based views, matching the user's own style."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "9. Frontend Data and Auth Pattern", 1)
    add_code_block(
        doc,
        """
Page
  |
  v
Domain hook (useJobRequestList, useEquipmentList, useReport, ...)
  |
  v
API wrapper (src/api/*.js)
  |
  v
Axios client
  |
  +-- request: attach Bearer access token and CSRF for mutations
  +-- response: coalesce 401 refresh and retry once
  +-- revoked token: clear memory and redirect to /login?reason=session_revoked
        """,
    )
    add_table(
        doc,
        ["Status", "Pattern", "Description"],
        [
            ("GREEN", "Token storage", "Access token and CSRF are held in module memory, not browser localStorage/sessionStorage."),
            ("GREEN", "Refresh", "Refresh token is handled as httpOnly cookie on the backend side."),
            ("BLUE", "Hooks", "Feature hooks wrap API clients and keep page code cleaner."),
            ("YELLOW", "Token capsules", "Some endpoint pattern rules should be verified against newer POST routes and -rows paths."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "10. Frontend Page Area Inventory", 1)
    add_table(
        doc,
        ["Page area", "Files", "Lines", "Interpretation"],
        [
            ("admin", "3", "968", "User and employee administration."),
            ("analytics", "4", "924", "Charting and CSV analytics views."),
            ("audit", "2", "526", "Audit review/export interface."),
            ("conversion", "5", "986", "Approval/rejection/conversion bridge."),
            ("dashboard", "6", "593", "KPI cards and overview surface."),
            ("equipment", "4", "945", "Equipment list/create/reserved detail route."),
            ("inquiry", "7", "614", "Search across vendors/products/job cards/instruments."),
            ("jobCards", "13", "2783", "Largest FE workflow area."),
            ("jobRequests", "12", "1778", "Request creation/detail/history/conversion inputs."),
            ("notifications", "1", "169", "Notification list/read state."),
            ("procurement", "6", "1070", "Purchase orders and spare parts."),
            ("reports", "10", "1190", "Report pages and export actions."),
            ("schedule", "4", "770", "Schedule calendar/list/status workflows."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "11. Backend Deep Dive: Stack", 1)
    add_table(
        doc,
        ["Status", "Backend item", "Description"],
        [
            ("BLUE", "Runtime", "Node.js / Express."),
            ("BLUE", "Database driver", "mysql2/promise pool with raw SQL repositories."),
            ("BLUE", "Validation", "zod for request validation where used."),
            ("BLUE", "Security", "helmet, cors, cookie-parser, JWT, refresh tokens, token version checks."),
            ("BLUE", "Operational helpers", "compression, pino logging, rate limiting, PDFKit, multer uploads."),
            ("GREEN", "DB boot", "Backend pool check succeeded against cmcmis_simplified."),
        ],
        status_col=0,
    )
    add_code_block(
        doc,
        """
server.js
  |
  +-- env validation
  +-- DB pool initialization
  +-- security/performance middleware
  +-- public health/auth routes
  +-- authenticated domain routes
  +-- notFound + errorHandler
        """,
    )

    page_break(doc)
    add_colored_heading(doc, "12. Backend Route Map", 1)
    add_table(
        doc,
        ["Mount", "Module / purpose"],
        [
            ("/healthz", "Health check."),
            ("/api/v1/auth", "Login, refresh, logout."),
            ("/api/v1/me", "Current user profile."),
            ("/api/v1/equipment", "Equipment list/create/helpers/bulk calibration done/reserved detail actions."),
            ("/api/v1/job-requests", "Job Request list/create/detail/submit/convert/reject/cancel/history."),
            ("/api/v1/job-cards", "Job Card list/detail/lifecycle/tabs/documents/tasks/rows."),
            ("/api/v1/lookups", "Lookup/reference values."),
            ("/api/v1/admin/users", "Admin user management."),
            ("/api/v1/admin/employees", "Employee master management."),
            ("/api/v1/dashboard", "KPI data."),
            ("/api/v1/inquiry", "Search endpoints."),
            ("/api/v1/notifications", "Notification list/read/unread count."),
            ("/api/v1/reports", "Report JSON/PDF endpoints."),
            ("/api/v1/analytics", "Chart/CSV analytics endpoints."),
            ("/api/v1/schedules", "Schedule CRUD/status/ICS."),
            ("/api/v1/procurement", "Purchase orders and spare parts."),
            ("/api/v1/audit", "Audit list/filter/export/detail."),
            ("/api/v1/.../*.pdf", "PDFKit-generated documents."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "13. Backend Module Inventory", 1)
    add_table(
        doc,
        ["Module", "Files", "Lines", "Responsibility"],
        [
            ("adminUsers", "6", "1054", "Role/user lifecycle, force logout, history."),
            ("analytics", "5", "720", "Chart datasets and CSV exports."),
            ("audit", "5", "891", "Read-only audit log exploration."),
            ("auth", "7", "841", "Login, refresh, logout, JWT, password checks."),
            ("dashboard", "5", "991", "KPI aggregation."),
            ("employees", "5", "762", "Employee master and account creation."),
            ("equipment", "5", "836", "Equipment list/create/helpers and bulk calibration done."),
            ("inquiry", "5", "721", "Search vendors/products/cards/instruments."),
            ("jobCards", "22", "3122", "Largest backend workflow module."),
            ("jobRequests", "6", "2545", "Request lifecycle and conversion."),
            ("lookups", "3", "291", "Shared reference endpoints."),
            ("notifications", "7", "685", "User notifications and read state."),
            ("pdf", "9", "1969", "Job Card / Job Request PDFs."),
            ("procurement", "5", "1210", "Purchase orders and spares."),
            ("reports", "6", "1961", "Operational reports and PDF export."),
            ("schedule", "7", "1240", "Schedules and ICS."),
            ("users", "3", "129", "Current user / me endpoint."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "14. Backend Middleware and Security Flow", 1)
    add_code_block(
        doc,
        """
HTTP request
  |
  +-- helmet
  +-- cors
  +-- compression
  +-- express.json(1mb)
  +-- cookie-parser
  +-- pino-http logger
  +-- authenticate JWT
  +-- token_version check against DB/cache
  +-- authorize / authorizeAny permission code
  +-- controller -> service -> repository
  +-- errorHandler
        """,
    )
    add_table(
        doc,
        ["Status", "Security behavior", "Description"],
        [
            ("GREEN", "Permission-first authorization", "Backend gates by permission codes such as job_request:approve, not role names."),
            ("GREEN", "Revocation model", "JWT includes token_version; backend rejects stale tokens as SESSION_REVOKED."),
            ("GREEN", "Refresh token storage", "Refresh token cookie is httpOnly and checked by backend refresh route."),
            ("GREEN", "Rate limiting", "Login and refresh routes use separate limiters."),
            ("BLUE", "Audit intent", "Critical writes are designed to leave audit/history records."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "15. Database Deep Dive: Live Shape", 1)
    add_table(
        doc,
        ["Status", "DB item", "Observed value"],
        [
            ("GREEN", "Server", "MariaDB 10.4.32-MariaDB."),
            ("GREEN", "Database", "cmcmis_simplified."),
            ("GREEN", "Table count", "104 tables."),
            ("GREEN", "Roles", "5 roles."),
            ("GREEN", "Permissions", "73 permission rows."),
            ("GREEN", "Users", "61 users and 61 user_roles."),
            ("BLUE", "Architecture", "Legacy CMMS tables plus MVP/security/workflow tables."),
        ],
        status_col=0,
    )
    add_table(
        doc,
        ["Table", "Exact rows", "Meaning"],
        [
            ("cmms_eqip_mst", "5701", "Equipment master records."),
            ("cmms_jobrequest_mst", "21520", "Legacy/MVP job request records."),
            ("cmms_jobcard_mst", "19440", "Job card records."),
            ("job_request_status_history", "21614", "JR lifecycle/history trail."),
            ("job_card_status_history", "14", "JC lifecycle/history trail."),
            ("audit_log", "138", "Critical-write audit records."),
            ("notifications", "58", "User notification records."),
            ("schedules", "5", "Schedule entries."),
            ("task_library", "45", "Reusable job card tasks."),
            ("schema_migrations", "38", "Applied migrations recorded in the live DB."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "16. Database Roles and Permissions", 1)
    add_table(
        doc,
        ["Role code", "Permission count", "Meaning"],
        [
            ("SUPER_ADMIN", "73", "Full system capability."),
            ("LAB_IN_CHARGE", "56", "Lab-wide operational control."),
            ("LAB_ENGINEER", "40", "Assigned/queue work execution capability."),
            ("VIEW_ONLY", "28", "Read-oriented access."),
            ("NORMAL_USER", "23", "Own-request and limited read actions."),
        ],
    )
    add_table(
        doc,
        ["Resource", "Permission count"],
        [
            ("analytics", "1"),
            ("audit", "2"),
            ("audit_log", "1"),
            ("auth", "3"),
            ("dashboard", "1"),
            ("equipment", "8"),
            ("export", "1"),
            ("inquiry", "4"),
            ("job_card", "10"),
            ("job_request", "8"),
            ("master", "5"),
            ("me", "1"),
            ("notifications", "2"),
            ("procurement", "7"),
            ("reports", "8"),
            ("schedule", "5"),
            ("user", "6"),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "17. Database Migration Timeline", 1)
    add_para(
        doc,
        "DATABASE/phase3/migrations contains a phased migration story. The runner is idempotent and records SHA-256 checksums in schema_migrations. Later phases added the permissions, workflow, reports, analytics, notifications, schedule, procurement, and bulk-action features."
    )
    add_code_block(
        doc,
        """
001-010  Core auth/RBAC/MVP tables and legacy ALTERs
050/099  Compatibility and controlled cleanup
100-122  Phase 6-8 Job Request, dashboard, inquiry, search, permissions
200-304  Conversion and Job Card detail lifecycle expansion
400-431  Reports, PDFs, analytics, notifications
500-510  Schedule and procurement
600-620  Audit and Phase 15 bulk permissions
        """,
    )
    add_table(
        doc,
        ["Status", "Migration insight", "Description"],
        [
            ("GREEN", "Runner design", "Migration runner uses schema_migrations, checksums, dry-run/status/reset modes."),
            ("BLUE", "Two-universe model", "Legacy tables remain while MVP tables and new columns support modern workflows."),
            ("YELLOW", "Old verifier counts", "The phase3 runner verification expectations are phase-specific and no longer match the expanded DB counts."),
            ("YELLOW", "Env difference", "DATABASE/phase3/.env points to final while BE runtime uses cmcmis_simplified."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "18. Core Workflow: Authentication", 1)
    add_code_block(
        doc,
        """
[Login form: employee_id + password]
  |
  v
POST /api/v1/auth/login
  |
  +-- validate credentials
  +-- create access JWT with token_version
  +-- set httpOnly refresh cookie
  +-- return CSRF/access info to FE memory
  |
  v
[Protected app]
  |
401? -> POST /api/v1/auth/refresh -> retry original request
SESSION_REVOKED? -> clear tokens -> /login?reason=session_revoked
        """,
    )
    add_table(
        doc,
        ["Status", "Auth rule", "Meaning"],
        [
            ("GREEN", "Employee-id login", "The system is employee-master aligned, not public self-registration."),
            ("GREEN", "Access token memory", "Frontend avoids persistent token storage."),
            ("GREEN", "Refresh cookie", "Backend-controlled httpOnly refresh token reduces JS exposure."),
            ("GREEN", "Token version", "Force logout and revocation are possible by changing user token_version."),
            ("BLUE", "Audit", "Login attempts are part of the security trail."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "19. Core Workflow: Job Request", 1)
    add_code_block(
        doc,
        """
DRAFT
  | submit
  v
SUBMITTED
  | approve
  v
APPROVED (logical/transient in service/history)
  | assign engineer / convert
  v
ASSIGNED -> Job Card created/linked

DRAFT -- cancel --> logically CANCELLED
SUBMITTED -- reject --> REJECTED
        """,
    )
    add_table(
        doc,
        ["Status", "JR behavior", "Description"],
        [
            ("BLUE", "Draft editing", "Draft requests can be edited before submission."),
            ("BLUE", "Submit", "DRAFT moves to SUBMITTED."),
            ("BLUE", "Conversion", "Convert requires both approve and assign-engineer style permissions."),
            ("BLUE", "History", "job_request_status_history preserves lifecycle evidence."),
            ("YELLOW", "Logical states", "APPROVED and CANCELLED have service/history semantics that should be understood before schema changes."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "20. Core Workflow: Job Card", 1)
    add_code_block(
        doc,
        """
ASSIGNED
  | start-work
  v
IN_PROGRESS
  | mark-complete
  v
COMPLETED
  | verify-close by LIC/SA
  v
VERIFIED_CLOSED

COMPLETED or VERIFIED_CLOSED -- reopen with reason --> IN_PROGRESS
        """,
    )
    add_table(
        doc,
        ["Status", "JC behavior", "Description"],
        [
            ("BLUE", "Owner gates", "Assigned engineer or LIC/SA can work/save/complete where allowed."),
            ("BLUE", "Verification", "LIC/SA verifies closure."),
            ("BLUE", "Reopen", "Reopen requires a reason with meaningful length."),
            ("BLUE", "Tabs", "Job Cards include tasks, maintenance rows, spares rows, documents, readings/observations style data."),
            ("YELLOW", "Legacy PDF route", "GET /api/v1/job-cards/:id/pdf still looks like an older placeholder while .pdf endpoints exist."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "21. Core Workflow: Equipment", 1)
    add_code_block(
        doc,
        """
Equipment create
  |
  v
PENDING_VERIFICATION
  |
  +-- verify by LIC/SA --> ACTIVE
  +-- condemn/delete reserved in routes
  |
  v
Equipment participates in calibration/job workflows
        """,
    )
    add_table(
        doc,
        ["Status", "Equipment insight", "Description"],
        [
            ("GREEN", "List/create", "Equipment listing and creation are implemented."),
            ("GREEN", "Helpers", "Types, makes, and divisions endpoints support forms."),
            ("GREEN", "Bulk action", "Phase 15 bulk calibration done permission and route exist."),
            ("YELLOW", "Reserved stubs", "Detail/update/verify/condemn/delete routes are present but still return Phase-6 style 404 stubs."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "22. Core Workflow: Reports, Analytics, PDFs", 1)
    add_table(
        doc,
        ["Feature", "Endpoints / outputs", "Permission model"],
        [
            ("Reports", "Calibration due, pending jobs, equipment utilization, engineer summary, job-card summary, job-request summary.", "View permission plus export permission for PDF."),
            ("Analytics", "Monthly activity, equipment status, monthly jobs, division-wise, calibration completion, job type, engineer workload, lifecycle funnel, trends.", "analytics:view or reports:view-analytics; CSV needs reports:export."),
            ("PDFs", "Job Request details, Job Card details, Job Card certificate.", "PDF routes are mounted explicitly before routers for stable matching."),
            ("Audit export", "Audit list/filter/export/detail.", "Read/export permissions keep audit read-only."),
        ],
    )
    add_code_block(
        doc,
        """
Report page -> API JSON -> table/chart
            -> export action -> BE report/pdf module -> PDFKit stream

Analytics page -> chart endpoint -> Recharts view
               -> CSV endpoint -> export audit trail
        """,
    )

    page_break(doc)
    add_colored_heading(doc, "23. Core Workflow: Schedule and Procurement", 1)
    add_table(
        doc,
        ["Area", "Implemented surface", "State / behavior"],
        [
            ("Schedule", "List, create, detail, patch, status, delete, export.ics, single-event ICS.", "PLANNED/SCHEDULED/DUE/COMPLETED/CANCELLED with terminal completed/cancelled."),
            ("Procurement", "Purchase orders and spare parts list/export/create/detail/patch/order.", "Supports operational purchasing/spares workflows."),
            ("Notifications", "List, unread count, read-all, read single.", "Supports user-facing workflow feedback."),
        ],
    )
    add_code_block(
        doc,
        """
PLANNED -> SCHEDULED -> DUE -> COMPLETED
       \\          \\        \\-> CANCELLED
        \\----------\\--------------> CANCELLED

COMPLETED and CANCELLED are terminal schedule states.
        """,
    )

    page_break(doc)
    add_colored_heading(doc, "24. Audit and Compliance Model", 1)
    add_table(
        doc,
        ["Status", "Audit layer", "Description"],
        [
            ("BLUE", "History tables", "Job Request and Job Card state changes have dedicated status history tables."),
            ("BLUE", "audit_log", "General critical-write audit table contains actor/action/context evidence."),
            ("BLUE", "audit_log_changes", "Supports before/after change details where used."),
            ("BLUE", "export_audit", "Exports can be tracked separately."),
            ("GREEN", "Audit API", "Audit module is read-only and permission protected."),
            ("PURPLE", "Engineering memory", "When changing write workflows, update audit/history behavior at the same time."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "25. Risks and Attention Map", 1)
    add_table(
        doc,
        ["Color", "Risk / attention", "Why it matters", "Recommended action"],
        [
            ("YELLOW", "Frontend bundle size", "Large JS can slow initial load.", "Introduce route-level lazy loading and Vite manual chunks after workflow stability."),
            ("YELLOW", "Equipment reserved routes", "Users may expect detail/update/verify/condemn/delete to work because routes exist.", "Prioritize implementation or hide/guard UI entry points."),
            ("YELLOW", "Legacy job-card PDF stub", "Duplicate/old route can confuse future maintainers.", "Consolidate PDF routes around the .pdf endpoints."),
            ("YELLOW", "Token capsule endpoint patterns", "If frontend success messages miss newer endpoints, UX feedback becomes inconsistent.", "Audit pattern map against current BE routes, especially POST JR actions and -rows child paths."),
            ("YELLOW", "DATABASE env mismatch", "Migration runner context can be mistaken for app runtime DB.", "Document active DB or align env files intentionally."),
            ("YELLOW", "Old phase verifier counts", "Stale expected counts can look like a failure after later migrations.", "Update verifier or label as historical Phase 3 baseline."),
            ("GREEN", "No red blocker found", "Core build/connectivity checks are green.", "Continue with focused implementation/testing improvements."),
        ],
        status_col=0,
    )

    page_break(doc)
    add_colored_heading(doc, "26. Recommended Roadmap", 1)
    add_table(
        doc,
        ["Priority", "Work item", "Expected benefit"],
        [
            ("P1", "Implement/finish equipment detail lifecycle routes and matching FE pages.", "Closes visible workflow gaps in a core module."),
            ("P1", "Add focused backend tests for JR/JC state machines and permission gates.", "Protects the riskiest business rules."),
            ("P2", "Audit frontend token capsule route map.", "Improves UX consistency after mutations."),
            ("P2", "Add route-level lazy loading in FE.", "Reduces initial bundle and resolves Vite warning."),
            ("P2", "Refresh migration runner verification expectations.", "Prevents confusion between early phase baseline and current schema."),
            ("P3", "Create a compact architecture README at root.", "Helps future developers onboard faster."),
            ("P3", "Add API contract notes per module.", "Makes FE/BE alignment easier during later phases."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "27. Long-Term Codebase Memory", 1)
    add_table(
        doc,
        ["Memory type", "Remembered rule"],
        [
            ("Vocabulary", "FE means frontend folder; BE means backend folder; DATABASE means database folder."),
            ("Style", "User prefers color-coded, table-based, flowchart/diagram heavy explanations."),
            ("Language", "Project code style is JavaScript, not TypeScript."),
            ("Architecture", "Prefer existing module/service/repository patterns over new abstractions."),
            ("Security", "Permissions are codes, not role-name checks."),
            ("Database", "Always verify live DB schema through MariaDB/phpMyAdmin context before assuming."),
            ("Workflow", "JR and JC state machines are the heart of the application."),
            ("Docs", "TECH_DOCX is the home for polished generated technical documents."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "28. Appendix: FE API and Hook Inventory", 1)
    add_table(
        doc,
        ["FE API file", "Lines"],
        [
            ("adminUsers.js", "36"),
            ("audit.js", "55"),
            ("dashboard.js", "30"),
            ("employees.js", "30"),
            ("equipment.js", "66"),
            ("inquiry.js", "40"),
            ("jobCards.js", "182"),
            ("jobRequests.js", "134"),
            ("lookups.js", "58"),
            ("notifications.js", "45"),
            ("pdf.js", "94"),
            ("procurement.js", "92"),
            ("reports.js", "127"),
            ("schedule.js", "68"),
        ],
    )
    add_table(
        doc,
        ["Representative hook", "Lines", "Purpose"],
        [
            ("useAutoSave.js", "81", "Auto-save pattern."),
            ("useDashboardKpis.js", "76", "Dashboard query."),
            ("useEquipmentList.js", "67", "Equipment listing."),
            ("useJobCardDetail.js", "61", "Job Card detail query."),
            ("useJobRequestDetail.js", "64", "Job Request detail query."),
            ("useNotifications.js", "112", "Notification read/unread workflow."),
            ("useProcurement.js", "84", "Procurement data workflow."),
            ("useSchedule.js", "65", "Schedule query/mutation workflow."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "29. Appendix: Backend Route Highlights", 1)
    add_table(
        doc,
        ["Module", "Important routes / behavior"],
        [
            ("auth", "POST /login, POST /refresh, POST /logout."),
            ("jobRequests", "GET/POST list/create; submit; bulk-verify-all; convert; reject; cancel; history."),
            ("jobCards", "List/detail/history/edit; start-work; mark-complete; verify-close; reopen; tasks/docs/maintenance/spares child routers."),
            ("equipment", "List/create/helpers/bulk-cal-done plus reserved detail lifecycle stubs."),
            ("reports", "Six JSON reports and matching PDF exports with reports:export checks."),
            ("analytics", "Twelve JSON chart endpoints and CSV exports."),
            ("schedule", "CRUD/status/ICS export."),
            ("procurement", "Purchase order and spare part workflows."),
            ("audit", "Read-only list/filter/export/detail."),
        ],
    )

    page_break(doc)
    add_colored_heading(doc, "30. Appendix: Final Understanding", 1)
    add_para(
        doc,
        "The project is best understood as a permissioned maintenance/calibration workflow platform. The frontend is the operational cockpit, the backend is the rule and audit engine, and the database is a combined legacy-and-modern record system. The most important future work should protect workflows first: authorization, state transitions, audit/history, and DB correctness."
    )
    add_table(
        doc,
        ["Final label", "Meaning"],
        [
            ("GREEN", "The project builds and connects."),
            ("BLUE", "The architecture is coherent and domain-driven."),
            ("YELLOW", "There are manageable improvement areas, mainly bundle size, stale stubs, and documentation alignment."),
            ("PURPLE", "I will keep using the user's preferred color/table/diagram explanation style for this project."),
        ],
        status_col=0,
    )

    doc.save(DOCX_PATH)


def build_markdown():
    md = []
    md.append("# CMCMIS End-to-End Codebase Deep Dive\n")
    md.append("Prepared on 2026-05-25 after inspecting FE, BE, DATABASE, TECH_DOCX, PDFs, temp, and the live MariaDB/phpMyAdmin-backed database.\n")
    md.append("## Color Legend\n")
    md.append("| Color | Meaning | Description |")
    md.append("|---|---|---|")
    md.append("| GREEN | Verified / healthy | Evidence confirms the behavior or connection is currently working. |")
    md.append("| BLUE | Core architecture | Stable design pattern or system structure to remember. |")
    md.append("| YELLOW | Attention / improvement | Not broken, but worth tracking before future expansion. |")
    md.append("| RED | Risk / blocker | Could break behavior or needs action before production confidence. |")
    md.append("| PURPLE | Memory / style | A rule, convention, or user preference I should preserve. |\n")

    md.append("## Executive Summary\n")
    md.append("| Status | Area | Insight |")
    md.append("|---|---|---|")
    md.append("| GREEN | Database | MariaDB/phpMyAdmin and BE DB pool are connected to `cmcmis_simplified`. |")
    md.append("| GREEN | Frontend | `npm run build` succeeded; Vite transformed 2546 modules. |")
    md.append("| GREEN | Backend | DB pool boot check succeeded with poolLimit 15. |")
    md.append("| BLUE | Architecture | React/Vite frontend -> Express backend -> service/repository -> MariaDB. |")
    md.append("| YELLOW | Maintainability | Frontend bundle is large and should later be split. |")
    md.append("| YELLOW | Gaps | Some reserved/stub routes and stale verifier expectations remain. |\n")

    md.append("## Project Architecture\n")
    md.append("```mermaid")
    md.append("flowchart TD")
    md.append('  U["User"] --> FE["FE frontend: React/Vite"]')
    md.append('  FE --> AX["Axios client: access token memory + CSRF"]')
    md.append('  AX --> BE["BE backend: Express middleware"]')
    md.append('  BE --> C["Controller"]')
    md.append('  C --> S["Service: validation + workflow"]')
    md.append('  S --> R["Repository: raw SQL"]')
    md.append('  R --> DB["MariaDB: cmcmis_simplified"]')
    md.append('  PMA["phpMyAdmin"] --> DB')
    md.append('  BE --> PDF["PDFKit on-demand PDFs"]')
    md.append('  BE --> ST["BE/storage job card documents"]')
    md.append("```\n")

    md.append("## Folder Inventory\n")
    md.append("| Folder | Files | Role |")
    md.append("|---|---:|---|")
    md.append("| BE | 151 | Backend API, modules, middleware, PDF generators, DB config. |")
    md.append("| FE | 159 | React pages, hooks, API clients, layout, permissions. |")
    md.append("| DATABASE | 53 | Migrations, runner, SQL/schema documentation. |")
    md.append("| PDFs | 35 | Historical/reference PDF and SQL artifacts. |")
    md.append("| TECH_DOCX | 21 before this report | Generated technical documentation. |")
    md.append("| temp | 39 | Temporary discovery/schema/export artifacts. |\n")

    md.append("## Frontend Route Memory\n")
    md.append("| Route | Feature | Permission |")
    md.append("|---|---|---|")
    for route, feature, perm in [
        ("/dashboard", "Dashboard", "dashboard:view"),
        ("/equipment", "Equipment list", "equipment:read-list"),
        ("/equipment/new", "New equipment", "equipment:create"),
        ("/job-requests", "Job Request list", "job_request:read-own"),
        ("/job-requests/new", "New Job Request", "job_request:create"),
        ("/conversion", "Conversion", "job_request:approve"),
        ("/job-cards", "Job Card list", "job_card:read-list"),
        ("/job-cards/:id", "Job Card detail", "job_card:read-detail"),
        ("/schedule", "Schedule", "schedule:read-list"),
        ("/procurement", "Procurement", "procurement:read-list"),
        ("/inquiry", "Inquiry", "inquiry search permissions"),
        ("/reports", "Reports", "reports:view-analytics"),
        ("/analytics", "Analytics", "analytics:view"),
        ("/admin/users", "Admin users", "user:read-list"),
        ("/admin/employees", "Admin employees", "master:employees:manage"),
        ("/audit", "Audit", "audit:read-list"),
    ]:
        md.append(f"| `{route}` | {feature} | `{perm}` |")
    md.append("")

    md.append("## Backend Module Inventory\n")
    md.append("| Module | Files | Lines | Responsibility |")
    md.append("|---|---:|---:|---|")
    for module, files, lines, resp in [
        ("adminUsers", 6, 1054, "User administration."),
        ("analytics", 5, 720, "Charts and CSV exports."),
        ("audit", 5, 891, "Audit exploration."),
        ("auth", 7, 841, "Login/refresh/logout."),
        ("dashboard", 5, 991, "KPI aggregation."),
        ("employees", 5, 762, "Employee master."),
        ("equipment", 5, 836, "Equipment list/create/helpers."),
        ("inquiry", 5, 721, "Search."),
        ("jobCards", 22, 3122, "Largest workflow module."),
        ("jobRequests", 6, 2545, "Request lifecycle."),
        ("pdf", 9, 1969, "PDF generation."),
        ("reports", 6, 1961, "Reports/PDF export."),
        ("schedule", 7, 1240, "Schedule/ICS."),
        ("procurement", 5, 1210, "Purchase/spares."),
    ]:
        md.append(f"| {module} | {files} | {lines} | {resp} |")
    md.append("")

    md.append("## Job Request Flow\n")
    md.append("```mermaid")
    md.append("flowchart LR")
    md.append('  D["DRAFT"] -->|submit| S["SUBMITTED"]')
    md.append('  S -->|approve| A["APPROVED logical"]')
    md.append('  A -->|assign / convert| AS["ASSIGNED + Job Card"]')
    md.append('  S -->|reject| R["REJECTED"]')
    md.append('  D -->|cancel| C["CANCELLED logical"]')
    md.append("```\n")

    md.append("## Job Card Flow\n")
    md.append("```mermaid")
    md.append("flowchart LR")
    md.append('  AS["ASSIGNED"] -->|start-work| IP["IN_PROGRESS"]')
    md.append('  IP -->|mark-complete| CO["COMPLETED"]')
    md.append('  CO -->|verify-close| VC["VERIFIED_CLOSED"]')
    md.append('  CO -->|reopen with reason| IP')
    md.append('  VC -->|reopen with reason| IP')
    md.append("```\n")

    md.append("## Live Database Snapshot\n")
    md.append("| Item | Value |")
    md.append("|---|---:|")
    md.append("| Tables | 104 |")
    md.append("| Roles | 5 |")
    md.append("| Permissions | 73 |")
    md.append("| Users | 61 |")
    md.append("| Equipment records | 5701 |")
    md.append("| Job requests | 21520 |")
    md.append("| Job cards | 19440 |")
    md.append("| Job request history rows | 21614 |")
    md.append("| Audit log rows | 138 |")
    md.append("| Schema migration rows | 38 |\n")

    md.append("## Risks and Attention Map\n")
    md.append("| Color | Attention item | Action |")
    md.append("|---|---|---|")
    md.append("| YELLOW | Large FE bundle | Add route-level lazy loading/manual chunks later. |")
    md.append("| YELLOW | Equipment stubs | Finish detail/update/verify/condemn/delete or hide UI entry points. |")
    md.append("| YELLOW | Legacy job-card PDF stub | Consolidate around current `.pdf` endpoints. |")
    md.append("| YELLOW | Token capsule route patterns | Compare pattern map against current BE routes. |")
    md.append("| YELLOW | DATABASE env mismatch | Document or align `final` vs `cmcmis_simplified`. |")
    md.append("| GREEN | No red blocker found | Build/connectivity checks are green. |\n")

    md.append("## Long-Term Memory\n")
    md.append("- FE means frontend folder.")
    md.append("- BE means backend folder.")
    md.append("- DATABASE means database folder.")
    md.append("- User style is color-coded, table-based, flowchart/diagram heavy, beginner-to-advanced.")
    md.append("- Code style is JavaScript-first.")
    md.append("- RBAC should use permission codes, not role-name checks.")
    md.append("- Job Request and Job Card state machines are the heart of the project.")
    md.append("- Always verify live DB schema before assuming table shape.")

    MD_PATH.write_text("\n".join(md), encoding="utf-8")


def build_memory_note():
    memory = """# CMCMIS Codebase Memory - 2026-05-25

## Naming and user style
- FE means the frontend folder.
- BE means the backend folder.
- DATABASE means the database folder and schema/migration assets.
- Preferred explanation style: color-coded, table-based, flowchart/diagram heavy, beginner-to-advanced.

## Project identity
- CMCMIS is a Computerized Maintenance and Calibration MIS.
- The app is an industry-style workflow system, not a simple CRUD demo.
- Runtime stack: React/Vite frontend, Express/mysql2 backend, MariaDB/MySQL database, phpMyAdmin administration.
- Current active runtime DB: `cmcmis_simplified`.

## Verified current state
- MariaDB is reachable at localhost:3306.
- phpMyAdmin opens the `cmcmis_simplified` structure route.
- Backend DB pool check succeeds.
- Frontend production build succeeds.
- Live DB has 104 tables, 5 roles, 73 permissions, 61 users.

## Architecture memory
- FE routes are permission-protected.
- BE modules follow route/controller/service/repository patterns.
- Database uses legacy CMMS tables plus newer MVP/security/workflow tables.
- RBAC must be permission-code based, not role-name based.
- Job Request and Job Card state machines are the core business rules.

## Attention items
- FE bundle is large; consider lazy loading/manual chunks later.
- Equipment detail/update/verify/condemn/delete backend routes are still reserved/stub style.
- One legacy Job Card PDF route remains while `.pdf` endpoints exist.
- FE token capsule route patterns should be checked against current BE routes.
- DATABASE/phase3 `.env` differs from BE runtime DB name.
"""
    MEMORY_PATH.write_text(memory, encoding="utf-8")


def main():
    build_docx()
    build_markdown()
    build_memory_note()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {MEMORY_PATH}")


if __name__ == "__main__":
    main()
